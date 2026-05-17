"""
DOCX Translator Pro — Flask backend v3
- gemini-3.1-flash-lite làm model chính, fallback chain đầy đủ
- Dịch song song ThreadPoolExecutor (MAX_WORKERS=4)
- Role detection bằng rule (không gọi AI classify cho heading/bullet/toc/table_cell)
- TOC được dịch (chỉ skip header/footer)
- build_document_context() 1 lần, truyền vào mọi chunk
- Prompt ngắn, nhanh
- SSE streaming progress realtime
- Inline edit + rebuild DOCX
"""
import os, io, json, uuid, traceback, logging, time
from concurrent.futures import ThreadPoolExecutor, as_completed

from flask import Flask, request, jsonify, render_template, send_file, Response, stream_with_context
from docx import Document
from google import genai
from google.genai import types

# ── Config ─────────────────────────────────────────────────────────────────────
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "AQ.Ab8RN6JJSM-jr9mamVH1q40AdSU9a5H0Km--epkgKdIEdh3_oA")

GEMINI_MODELS = [
    "gemini-3.1-flash-lite",   # model chính — nhanh nhất
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
]

CHUNK_SIZE   = 25    # đoạn/chunk
MAX_WORKERS  = 4     # parallel calls — giảm xuống 3 nếu bị rate limit
MAX_TOKENS   = 65536

# Chỉ 2 role này KHÔNG dịch — header/footer giữ nguyên tiếng Anh
NO_TRANSLATE_ROLES = {"header", "footer"}

# ── App init ───────────────────────────────────────────────────────────────────
app = Flask(__name__, template_folder=os.path.dirname(os.path.abspath(__file__)) or ".")
app.secret_key = "docx-translator-local"
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("docx-translator")

client = genai.Client(api_key=GEMINI_API_KEY)
STORE: dict = {}
_working_model: list = [None]   # cache model đang hoạt động


# ══════════════════════════════════════════════════════════════════════════════
# Paragraph iteration — SAME iterator cho extract VÀ apply
# ══════════════════════════════════════════════════════════════════════════════

def _iter_paragraphs_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            for nested in cell.tables:
                yield from _iter_paragraphs_in_table(nested)


def iter_all_paragraphs(doc):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        yield from _iter_paragraphs_in_table(table)
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                yield para
            for tbl in section.header.tables:
                yield from _iter_paragraphs_in_table(tbl)
        except Exception:
            pass
        try:
            for para in section.footer.paragraphs:
                yield para
            for tbl in section.footer.tables:
                yield from _iter_paragraphs_in_table(tbl)
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# Format-safe text replacement — không rebuild XML, không xóa run
# ══════════════════════════════════════════════════════════════════════════════

def replace_paragraph_text_keep_format(paragraph, new_text: str):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    first_text_run = None
    for run in runs:
        if run.text:
            first_text_run = run
            break
    if first_text_run is None:
        runs[0].text = new_text
        return
    first_text_run.text = new_text
    after = False
    for run in runs:
        if run is first_text_run:
            after = True
            continue
        if after and run.text:
            run.text = ""


# ══════════════════════════════════════════════════════════════════════════════
# Gemini helpers
# ══════════════════════════════════════════════════════════════════════════════

def _call_gemini(prompt: str, temperature: float = 0.0) -> str:
    """Gọi Gemini, ưu tiên model đang hoạt động, fallback qua danh sách."""
    models = ([_working_model[0]] if _working_model[0] else []) + \
             [m for m in GEMINI_MODELS if m != _working_model[0]]
    last_err = "no models"
    for model in models:
        try:
            resp = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=temperature,
                    max_output_tokens=MAX_TOKENS,
                ),
            )
            text = (resp.text or "").strip()
            if not text:
                last_err = "empty response"
                continue
            if _working_model[0] != model:
                log.info(f"✓ Active model: {model}")
                _working_model[0] = model
            return text
        except Exception as e:
            log.warning(f"Model {model}: {str(e)[:120]}")
            last_err = str(e)
    raise RuntimeError(f"All models failed. Last: {last_err}")


def _parse_json_list(raw: str) -> list:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()
    s, e = raw.find("["), raw.rfind("]")
    if s != -1 and e > s:
        candidate = raw[s:e+1]
        try:
            data = json.loads(candidate)
            if isinstance(data, list):
                return data
        except json.JSONDecodeError:
            last = candidate.rfind("},")
            if last == -1: last = candidate.rfind("}")
            if last > 0:
                try:
                    data = json.loads(candidate[:last+1].rstrip(",") + "]")
                    if isinstance(data, list):
                        log.warning("JSON truncated — partial recovery")
                        return data
                except Exception:
                    pass
    log.error(f"Cannot parse JSON list. Raw[:300]: {raw[:300]}")
    return []


# ══════════════════════════════════════════════════════════════════════════════
# Role detection hoàn toàn bằng rule — KHÔNG gọi AI
# ══════════════════════════════════════════════════════════════════════════════

def _detect_role_by_rule(para, in_table: bool, in_header: bool, in_footer: bool) -> str:
    """
    Detect role bằng Word metadata — nhanh, không tốn API call.
    Trả về: header | footer | toc | section_heading | bullet | table_cell | paragraph
    """
    if in_header:
        return "header"
    if in_footer:
        return "footer"

    style_name = (para.style.name or "").lower()

    # TOC — Word style "TOC 1", "TOC 2", "toc heading", etc.
    if "toc" in style_name:
        return "toc"

    # Heading
    if "heading" in style_name:
        return "section_heading"

    # Table cell
    if in_table:
        return "table_cell"

    # Bullet — check numPr in XML
    try:
        pPr = para._p.pPr
        if pPr is not None and pPr.numPr is not None:
            return "bullet"
    except Exception:
        pass

    # Title style
    if style_name in ("title", "subtitle"):
        return "title"

    # Note / caption
    if any(k in style_name for k in ("caption", "note", "warning", "caution")):
        return "note"

    return "paragraph"


# ══════════════════════════════════════════════════════════════════════════════
# Extract blocks — dùng iter_all_paragraphs, rule-based role
# ══════════════════════════════════════════════════════════════════════════════

def extract_docx_blocks(docx_bytes: bytes) -> list:
    doc    = Document(io.BytesIO(docx_bytes))
    blocks = []

    # Build set of paragraphs inside tables
    table_para_set: set = set()
    for tbl in doc.tables:
        for para in _iter_paragraphs_in_table(tbl):
            table_para_set.add(id(para._element))

    # Build set of paragraphs inside header/footer
    header_para_set: set = set()
    footer_para_set: set = set()
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                header_para_set.add(id(para._element))
            for tbl in section.header.tables:
                for para in _iter_paragraphs_in_table(tbl):
                    header_para_set.add(id(para._element))
        except Exception:
            pass
        try:
            for para in section.footer.paragraphs:
                footer_para_set.add(id(para._element))
            for tbl in section.footer.tables:
                for para in _iter_paragraphs_in_table(tbl):
                    footer_para_set.add(id(para._element))
        except Exception:
            pass

    for idx, para in enumerate(iter_all_paragraphs(doc)):
        text = para.text.strip()
        if not text:
            continue

        pid       = id(para._element)
        in_table  = pid in table_para_set
        in_header = pid in header_para_set
        in_footer = pid in footer_para_set

        role = _detect_role_by_rule(para, in_table, in_header, in_footer)

        blocks.append({
            "id":       f"p{idx}",
            "text":     text,
            "role":     role,
            "para_idx": idx,
        })

    translatable = sum(1 for b in blocks if b["role"] not in NO_TRANSLATE_ROLES)
    hf_count     = sum(1 for b in blocks if b["role"] in NO_TRANSLATE_ROLES)
    toc_count    = sum(1 for b in blocks if b["role"] == "toc")
    log.info(f"Extracted {len(blocks)} blocks: {translatable} translatable "
             f"({toc_count} toc), {hf_count} header/footer")
    return blocks


# ══════════════════════════════════════════════════════════════════════════════
# Build document context — 1 lần, dùng cho mọi chunk
# ══════════════════════════════════════════════════════════════════════════════

def build_document_context(blocks: list) -> str:
    """
    Tạo context ngắn gọn từ title + headings + toc để AI hiểu cấu trúc tài liệu.
    Tối đa ~300 từ để không làm nặng prompt.
    """
    title_blocks   = [b["text"] for b in blocks if b["role"] == "title"][:2]
    heading_blocks = [b["text"] for b in blocks if b["role"] == "section_heading"][:10]
    toc_blocks     = [b["text"] for b in blocks if b["role"] == "toc"][:8]

    lines = []
    if title_blocks:
        lines.append("Document title: " + " | ".join(title_blocks))
    if heading_blocks:
        lines.append("Main sections: " + " | ".join(heading_blocks[:6]))
    if toc_blocks:
        lines.append("TOC structure (first entries): " + " | ".join(toc_blocks[:5]))

    ctx = "\n".join(lines) if lines else "Technical document."

    # Heuristic: detect domain
    all_text = " ".join(b["text"] for b in blocks[:30]).lower()
    if any(k in all_text for k in ["elevator", "lift", "hoistway", "schindler", "inventio"]):
        ctx += "\nDomain: elevator/lift engineering."
    elif any(k in all_text for k in ["safety", "standard", "regulation", "iso", "en 81"]):
        ctx += "\nDomain: safety standards / technical regulation."
    elif any(k in all_text for k in ["software", "api", "code", "function", "module"]):
        ctx += "\nDomain: software/IT."

    return ctx


# ══════════════════════════════════════════════════════════════════════════════
# Translate one chunk — prompt ngắn, format-faithful
# ══════════════════════════════════════════════════════════════════════════════

def translate_chunk(chunk: list, target_lang: str, doc_context: str) -> dict:
    payload = json.dumps(
        [{"id": b["id"], "text": b["text"], "role": b.get("role", "paragraph")}
         for b in chunk],
        ensure_ascii=False
    )

    prompt = f"""Translate these Word document blocks into {target_lang}.

Document context:
{doc_context}

Rules:
- Return ONLY valid JSON array. No markdown, no explanation.
- Keep "id" exactly as given.
- Translate ALL English words. Leave nothing in English.
- Preserve: numbers, units, product codes, model names, document IDs, revision codes.
- Do NOT translate company names (e.g. INVENTIO AG, Schindler, Otis).
- Preserve ALL punctuation, whitespace, tabs (\\t), line breaks (\\n).
- Preserve CAPS/Title Case/lowercase exactly as in source.
- Match source sentence count — do NOT merge or split sentences.
- For toc: translate heading text only; preserve numbering, dot leaders (...), page numbers, tabs.
- For bullet: concise imperative, keep bullet symbol.
- For heading/section_heading: concise technical, match source brevity.
- For table_cell: concise, consistent across row.
- For note: keep WARNING/NOTE/CAUTION label verbatim.

Format:
[{{"id":"...","text":"translated text"}}]

Blocks:
{payload}"""

    try:
        raw    = _call_gemini(prompt, temperature=0.1)
        parsed = _parse_json_list(raw)
        result = {}
        for item in parsed:
            if isinstance(item, dict) and item.get("id") and item.get("text") is not None:
                result[item["id"]] = str(item["text"])
        # Fallback giữ text gốc cho block bị miss
        for b in chunk:
            if b["id"] not in result:
                log.warning(f"Block {b['id']} not in response — keeping original")
                result[b["id"]] = b["text"]
        return result
    except Exception as e:
        log.error(f"translate_chunk error: {e}")
        return {b["id"]: b["text"] for b in chunk}


def translate_chunk_with_retry(chunk: list, target_lang: str, doc_context: str,
                                retries: int = 3) -> dict:
    last_err = None
    for attempt in range(retries):
        try:
            return translate_chunk(chunk, target_lang, doc_context)
        except Exception as e:
            last_err = e
            wait = 2 ** attempt   # 1s, 2s, 4s
            log.warning(f"Chunk retry {attempt+1}/{retries} in {wait}s: {e}")
            time.sleep(wait)
    log.error(f"Chunk failed after {retries} retries: {last_err}")
    return {b["id"]: b["text"] for b in chunk}


# ══════════════════════════════════════════════════════════════════════════════
# Apply translations back into DOCX
# ══════════════════════════════════════════════════════════════════════════════

def apply_translations(original_docx_bytes: bytes, blocks: list, translations: dict) -> bytes:
    doc = Document(io.BytesIO(original_docx_bytes))
    idx_to_para: dict = {}
    for idx, para in enumerate(iter_all_paragraphs(doc)):
        idx_to_para[idx] = para

    applied = skipped = 0
    for block in blocks:
        # Chỉ skip header/footer — toc và mọi role khác đều apply
        if block.get("role") in NO_TRANSLATE_ROLES:
            skipped += 1
            continue
        tr = translations.get(block["id"])
        if not tr:
            continue
        para_idx = block.get("para_idx")
        if para_idx is None:
            continue
        para = idx_to_para.get(para_idx)
        if para is None:
            log.warning(f"Block {block['id']}: para_idx {para_idx} not found")
            continue
        try:
            replace_paragraph_text_keep_format(para, tr)
            applied += 1
        except Exception as e:
            log.warning(f"Block {block['id']} write failed: {e}")

    log.info(f"Applied {applied} translations, skipped {skipped} header/footer")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# Routes
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    f = request.files.get("file")
    if not f or not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "Cần file .docx (Microsoft Word)"}), 400

    sid        = str(uuid.uuid4())
    docx_bytes = f.read()

    try:
        blocks = extract_docx_blocks(docx_bytes)
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Không đọc được DOCX: {e}"}), 500

    STORE[sid] = {
        "docx_bytes":      docx_bytes,
        "filename":        f.filename,
        "blocks":          blocks,
        "translations":    {},
        "translated_docx": None,
    }

    translatable = [b for b in blocks if b["role"] not in NO_TRANSLATE_ROLES]
    hf_count     = len(blocks) - len(translatable)
    toc_count    = sum(1 for b in blocks if b["role"] == "toc")
    preview      = [
        {"id": b["id"], "text": b["text"], "role": b.get("role", "paragraph")}
        for b in blocks[:20]
    ]

    return jsonify({
        "sid":           sid,
        "filename":      f.filename,
        "total_blocks":  len(blocks),
        "translatable":  len(translatable),
        "header_footer": hf_count,
        "toc_count":     toc_count,
        "preview":       preview,
    })


@app.route("/translate_stream")
def translate_stream():
    """SSE endpoint — dịch song song, stream progress realtime."""
    sid      = request.args.get("sid")
    lang     = request.args.get("lang", "Vietnamese")
    skip_raw = request.args.get("skip_ids", "[]")
    try:
        skip_ids = set(json.loads(skip_raw))
    except Exception:
        skip_ids = set()
    s = STORE.get(sid)

    if not s:
        def err_gen():
            yield f"data: {json.dumps({'type':'error','message':'Phiên không tồn tại'})}\n\n"
        return Response(stream_with_context(err_gen()), mimetype="text/event-stream")

    def generate():
        try:
            blocks = s["blocks"]

            # Build document context — 1 lần
            yield f"data: {json.dumps({'type':'status','message':'Đang phân tích tài liệu...'})}\n\n"
            doc_context = build_document_context(blocks)
            log.info(f"Document context:\n{doc_context}")

            all_translatable = [b for b in blocks if b.get("role") not in NO_TRANSLATE_ROLES]
            # Rescan mode: chỉ dịch các block chưa được skip (chưa dịch)
            if skip_ids:
                translatable = [b for b in all_translatable if b["id"] not in skip_ids]
                log.info(f"Rescan: {len(translatable)} chưa dịch / {len(all_translatable)} tổng")
            else:
                translatable = all_translatable
            total = len(translatable)

            if total == 0:
                yield f"data: {json.dumps({'type':'done','translated_count':0,'total':0,'message':'Tất cả đoạn đã được dịch rồi!'})}\n\n"
                return

            chunks     = [translatable[i:i+CHUNK_SIZE] for i in range(0, total, CHUNK_SIZE)]
            num_chunks = len(chunks)
            # Rescan: khởi đầu từ bản dịch đã có để apply_translations giữ cả cũ lẫn mới
            translations = dict(s.get("translations", {})) if skip_ids else {}
            done_count   = 0

            mode_label = "Quét bỏ sót" if skip_ids else "Dịch"
            yield f"data: {json.dumps({'type':'status','message':f'{mode_label}: {total} đoạn — {num_chunks} chunk — {MAX_WORKERS} luồng song song...'})}\n\n"

            # Song song với ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                future_map = {
                    executor.submit(
                        translate_chunk_with_retry,
                        chunk, lang, doc_context
                    ): (ci, chunk)
                    for ci, chunk in enumerate(chunks)
                }

                for future in as_completed(future_map):
                    ci, chunk = future_map[future]
                    try:
                        chunk_result = future.result()
                    except Exception as ex:
                        log.error(f"Chunk {ci+1} future error: {ex}")
                        chunk_result = {b["id"]: b["text"] for b in chunk}

                    translations.update(chunk_result)
                    done_count += len(chunk)

                    chunk_rows = [
                        {
                            "id":         b["id"],
                            "role":       b.get("role", "paragraph"),
                            "english":    b["text"],
                            "vietnamese": chunk_result.get(b["id"], b["text"]),
                        }
                        for b in chunk
                    ]

                    pct = round(done_count / total * 100)
                    event = {
                        "type":         "progress",
                        "chunk":        ci + 1,
                        "total_chunks": num_chunks,
                        "done":         done_count,
                        "total":        total,
                        "pct":          pct,
                        "rows":         chunk_rows,
                    }
                    yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

            # Build final docx
            yield f"data: {json.dumps({'type':'status','message':'Đang tạo file DOCX...'})}\n\n"
            s["translations"]    = translations
            s["translated_docx"] = apply_translations(s["docx_bytes"], blocks, translations)

            yield f"data: {json.dumps({'type':'done','translated_count':len(translations),'total':total}, ensure_ascii=False)}\n\n"

        except Exception as e:
            traceback.print_exc()
            yield f"data: {json.dumps({'type':'error','message':str(e)})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/save_edit", methods=["POST"])
def save_edit():
    """Lưu bản dịch đã sửa inline và rebuild DOCX."""
    data = request.json or {}
    sid  = data.get("sid")
    bid  = data.get("id")
    text = data.get("text", "")
    s    = STORE.get(sid)
    if not s:
        return jsonify({"error": "Phiên không tồn tại"}), 404
    if not bid:
        return jsonify({"error": "Thiếu block id"}), 400

    s["translations"][bid] = text
    try:
        s["translated_docx"] = apply_translations(
            s["docx_bytes"], s["blocks"], s["translations"]
        )
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/download/<sid>")
def download(sid):
    s = STORE.get(sid)
    if not s or not s["translated_docx"]:
        return "Không tìm thấy", 404
    orig = s.get("filename", "document.docx")
    name = orig.rsplit(".", 1)[0] + "_translated.docx"
    return send_file(
        io.BytesIO(s["translated_docx"]),
        as_attachment=True,
        download_name=name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    import threading, webbrowser
    port = int(os.environ.get("PORT", 5000))
    threading.Timer(1.0, lambda: webbrowser.open(f"http://localhost:{port}")).start()
    print(f"\n✅ DOCX Translator v3: http://localhost:{port}")
    print(f"🚀 Model: {GEMINI_MODELS[0]} (fallback: {', '.join(GEMINI_MODELS[1:])})")
    print(f"⚡ CHUNK_SIZE={CHUNK_SIZE}, MAX_WORKERS={MAX_WORKERS}")
    print("🛑 Ctrl+C để dừng\n")
    app.run(host="0.0.0.0", port=port, debug=False, use_reloader=False)

"""
So sánh bản dịch DOCX với bản gốc → đánh giá chất lượng + đề xuất.

Flow:
  extract original + translated → align by position → chunk pairs →
  LLM review (parallel) → list issues with severity/suggestion/reason.
"""
import io
import json
from concurrent.futures import ThreadPoolExecutor, as_completed

from google.genai import types as gtypes

from word_backend import get_working_model
from config import NO_TRANSLATE_ROLES, WORD_MODELS

REVIEW_CHUNK_SIZE = 12       # pairs per LLM call
MAX_REVIEW_WORKERS = 6


def align_blocks(orig_blocks: list[dict],
                 trans_blocks: list[dict]) -> tuple[list[tuple], list[str]]:
    """
    Align original ↔ translated blocks by position (after filtering empty/NO_TRANSLATE).
    Returns (pairs, warnings).
    """
    def keep(b):
        return (b["text"].strip()
                and b["role"] not in NO_TRANSLATE_ROLES
                and b["role"] != "image_alt")  # alt-text alignment is brittle

    o = [b for b in orig_blocks if keep(b)]
    t = [b for b in trans_blocks if keep(b)]
    warnings = []
    if len(o) != len(t):
        warnings.append(
            f"⚠️ Số đoạn không khớp: gốc={len(o)} vs dịch={len(t)}. "
            f"Sẽ align theo thứ tự đến hết phần ngắn hơn ({min(len(o), len(t))} đoạn)."
        )
    n = min(len(o), len(t))
    pairs = []
    for i in range(n):
        ob, tb = o[i], t[i]
        if ob["role"] != tb["role"] and i < 5:
            warnings.append(
                f"⚠️ Pair #{i}: role gốc={ob['role']} ≠ dịch={tb['role']}"
            )
        pairs.append((ob, tb))
    return pairs, warnings


def _build_review_prompt(pairs_chunk: list[tuple], target_lang: str,
                         doc_context: str, strictness: str) -> str:
    strictness_desc = {
        "strict":  "Strict — flag every imperfection, even minor style issues.",
        "balanced": "Balanced — flag accuracy, completeness, and notable style issues. Ignore trivial wording preferences.",
        "permissive": "Permissive — only flag clear errors: wrong meaning, missing content, severe style/register problems.",
    }.get(strictness, "Balanced")

    pairs_text = []
    for i, (ob, tb) in enumerate(pairs_chunk):
        pairs_text.append(
            f"[{i}] id={ob['id']} role={ob['role']}\n"
            f"  ORIGINAL: {ob['text']}\n"
            f"  TRANSLATED: {tb['text']}"
        )

    return (
        f"You are a professional translation reviewer. The target language is: {target_lang}.\n"
        f"Document context: {doc_context}\n\n"
        f"Strictness: {strictness_desc}\n\n"
        f"For EACH pair below, evaluate:\n"
        f"1. ACCURACY — does the translation convey the original meaning faithfully?\n"
        f"2. COMPLETENESS — is any content missing or hallucinated?\n"
        f"3. STYLE — tone, register, fluency, naturalness in {target_lang}?\n"
        f"4. TERMINOLOGY — consistency and correctness of domain terms?\n\n"
        f"Return ONLY a JSON object:\n"
        f"{{\"reviews\": [\n"
        f"  {{\"id\": \"<id>\", \"severity\": \"critical|major|minor|ok\",\n"
        f"   \"issues\": [\"short label\", ...], \"suggested\": \"improved translation\",\n"
        f"   \"reason\": \"concise explanation\"}}, ...\n"
        f"]}}\n\n"
        f"Rules:\n"
        f"- severity=ok → omit \"suggested\" and \"reason\" (or leave empty)\n"
        f"- severity=critical → wrong meaning, missing key content, hallucination\n"
        f"- severity=major → significant style/terminology issues\n"
        f"- severity=minor → small improvements only\n"
        f"- Include EVERY pair in your response, even if severity=ok\n"
        f"- \"suggested\" must be in {target_lang}\n"
        f"- \"reason\" must be in Vietnamese\n\n"
        f"Pairs:\n" + "\n\n".join(pairs_text)
    )


def review_chunk(client, pairs_chunk, target_lang, doc_context, strictness,
                 chunk_idx: int) -> tuple[list[dict], int, int]:
    """Return (reviews, in_tokens, out_tokens)."""
    prompt = _build_review_prompt(pairs_chunk, target_lang, doc_context, strictness)
    model = get_working_model() or WORD_MODELS[0]
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=gtypes.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.2,
        ),
    )
    meta  = getattr(response, "usage_metadata", None)
    in_t  = getattr(meta, "prompt_token_count",     0) or 0
    out_t = getattr(meta, "candidates_token_count", 0) or 0
    try:
        data = json.loads(response.text)
        reviews = data.get("reviews", [])
        if not isinstance(reviews, list):
            reviews = []
    except Exception:
        reviews = []
    return reviews, in_t, out_t


def chunk_pairs(pairs: list[tuple], size: int = REVIEW_CHUNK_SIZE) -> list[list[tuple]]:
    return [pairs[i:i+size] for i in range(0, len(pairs), size)]


def review_parallel(holder, client, chunks, target_lang, doc_context, strictness,
                    max_workers: int = MAX_REVIEW_WORKERS):
    """Background runner. holder is shared dict updated as chunks finish."""
    try:
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {
                ex.submit(review_chunk, client, c, target_lang, doc_context,
                          strictness, i): i
                for i, c in enumerate(chunks)
            }
            for fut in as_completed(futures):
                idx = futures[fut]
                try:
                    reviews, in_t, out_t = fut.result()
                    holder["reviews"].extend(reviews)
                    holder["tok_in"]  += in_t
                    holder["tok_out"] += out_t
                    holder["chunk_done"] += 1
                    holder["chunk_log"].append({
                        "idx": idx, "size": len(chunks[idx]),
                        "in_t": in_t, "out_t": out_t, "error": None,
                    })
                except Exception as e:
                    holder["chunk_done"] += 1
                    holder["chunk_log"].append({
                        "idx": idx, "size": len(chunks[idx]),
                        "in_t": 0, "out_t": 0, "error": str(e)[:200],
                    })
    except Exception as e:
        holder["error"] = str(e)
    finally:
        holder["done"] = True


def summarize_reviews(reviews: list[dict]) -> dict:
    """Count by severity."""
    counts = {"critical": 0, "major": 0, "minor": 0, "ok": 0}
    for r in reviews:
        sev = r.get("severity", "ok")
        if sev in counts:
            counts[sev] += 1
    return counts


def build_review_report_docx(pairs_by_id: dict, reviews: list[dict]) -> bytes:
    """Generate DOCX report: only entries with severity != ok."""
    from docx import Document
    from docx.shared import RGBColor, Inches

    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)

    doc.add_heading("Báo cáo đánh giá bản dịch", level=0)

    counts = summarize_reviews(reviews)
    p = doc.add_paragraph()
    p.add_run(f"Tổng: {len(reviews)} đoạn  •  ").bold = True
    p.add_run(f"🔴 Critical: {counts['critical']}  ").font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(f"🟠 Major: {counts['major']}  ").font.color.rgb = RGBColor(0xD0, 0x70, 0x00)
    p.add_run(f"🟡 Minor: {counts['minor']}  ").font.color.rgb = RGBColor(0x80, 0x80, 0x00)
    p.add_run(f"✅ OK: {counts['ok']}").font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    sev_order = {"critical": 0, "major": 1, "minor": 2, "ok": 3}
    sorted_revs = sorted(reviews, key=lambda r: sev_order.get(r.get("severity"), 4))

    for r in sorted_revs:
        sev = r.get("severity", "ok")
        if sev == "ok":
            continue
        pair = pairs_by_id.get(r.get("id"))
        if not pair:
            continue
        ob, tb = pair

        h = doc.add_heading(f"[{sev.upper()}] {r['id']}", level=2)
        if sev == "critical":
            h.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif sev == "major":
            h.runs[0].font.color.rgb = RGBColor(0xD0, 0x70, 0x00)

        issues = r.get("issues") or []
        if issues:
            doc.add_paragraph("Vấn đề: " + ", ".join(issues)).runs[0].italic = True

        doc.add_paragraph().add_run("Gốc:").bold = True
        doc.add_paragraph(ob["text"])
        doc.add_paragraph().add_run("Bản dịch hiện tại:").bold = True
        doc.add_paragraph(tb["text"])
        if r.get("suggested"):
            doc.add_paragraph().add_run("Đề xuất:").bold = True
            doc.add_paragraph(r["suggested"])
        if r.get("reason"):
            doc.add_paragraph().add_run("Lý do:").bold = True
            doc.add_paragraph(r["reason"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

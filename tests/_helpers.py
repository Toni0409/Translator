"""Shared helpers cho smoke scripts trong tests/.

Smoke scripts dùng stdlib + project deps (KHÔNG dùng pytest).
Chạy:   python tests/smoke_<name>.py

Helpers:
- `setup_env()`: stub `gemini` module + tạo `.streamlit/secrets.toml` tạm thời
  để config import OK. Trả về cleanup callable.
- `make_docx_with_image(...)`: build DOCX có ảnh.
- `make_docx_with_text(...)`: build DOCX text only.
"""
from __future__ import annotations

import io
import os
import sys
import types


REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def setup_env():
    """Stub gemini module + secrets.toml. Trả về cleanup callable."""
    if REPO_ROOT not in sys.path:
        sys.path.insert(0, REPO_ROOT)

    if "gemini" not in sys.modules:
        fake = types.ModuleType("gemini")
        fake.generate         = lambda *a, **kw: ("{}", 0, 0)
        fake.usage_tokens     = lambda *a, **kw: (0, 0)
        fake.parse_json_loose = lambda x: None
        fake.get_client       = lambda: None
        sys.modules["gemini"] = fake

    secrets_path = os.path.join(REPO_ROOT, ".streamlit", "secrets.toml")
    created_secrets = False
    if not os.path.exists(secrets_path):
        os.makedirs(os.path.dirname(secrets_path), exist_ok=True)
        with open(secrets_path, "w") as f:
            f.write('GEMINI_API_KEY = ""\nAPP_PASSWORD = ""\n')
        created_secrets = True

    def cleanup():
        if created_secrets and os.path.exists(secrets_path):
            os.remove(secrets_path)

    return cleanup


def make_docx_with_image() -> bytes:
    """DOCX với text paragraph + image-only paragraph + closing paragraph.

    Yêu cầu: python-docx + Pillow.
    """
    from docx import Document
    from docx.shared import Inches
    import PIL.Image as Image

    img_buf = io.BytesIO()
    Image.new("RGB", (50, 50), (128, 200, 80)).save(img_buf, "PNG")
    img_buf.seek(0)

    doc = Document()
    doc.add_paragraph("Hello world")
    p_img = doc.add_paragraph()
    p_img.add_run().add_picture(img_buf, width=Inches(0.5))
    doc.add_paragraph("Final text")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def make_docx_with_inline_image() -> bytes:
    """DOCX paragraph có bold text + image inline + plain text."""
    from docx import Document
    from docx.shared import Inches
    import PIL.Image as Image

    img_buf = io.BytesIO()
    Image.new("RGB", (40, 40), (255, 0, 0)).save(img_buf, "PNG")
    img_buf.seek(0)

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Bold here")
    r1.bold = True
    r_img = p.add_run()
    r_img.add_picture(img_buf, width=Inches(0.4))
    p.add_run(" after image")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def make_docx_plain() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ── Simple assertion / report ──────────────────────────────────────────────
class Report:
    def __init__(self, name: str):
        self.name = name
        self.passed = 0
        self.failed = 0

    def check(self, label: str, ok: bool, detail: str = ""):
        if ok:
            self.passed += 1
            print(f"  ✅ {label}")
            if detail:
                print(f"     {detail}")
        else:
            self.failed += 1
            print(f"  ❌ {label}")
            if detail:
                print(f"     {detail}")

    def summary(self) -> int:
        total = self.passed + self.failed
        print(f"\n[{self.name}] {self.passed}/{total} passed", end="")
        if self.failed:
            print(f", {self.failed} FAILED")
            return 1
        print()
        return 0

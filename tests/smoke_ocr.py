"""OCR smoke checks — cost helpers, occurrence extraction, caption selection.

Chạy: python tests/smoke_ocr.py
"""
from __future__ import annotations

import io
import os
import sys

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from tests._helpers import setup_env, _PNG_1X1, Report


cleanup = setup_env()
try:
    from word_backend import (
        extract_image_occurrences,
        estimate_image_input_tokens, estimate_ocr_cost,
        _build_ocr_prompt,
        insert_ocr_captions_into_docx,
    )

    report = Report("smoke_ocr")

    # ── T1: cost helpers ────────────────────────────────────────────────
    report.check("small img (200x100) = 258 tokens",
                 estimate_image_input_tokens(200, 100) == 258)
    report.check("medium img (1024x768) = 2 tiles × 258 = 516",
                 estimate_image_input_tokens(1024, 768) == 516)
    report.check("big img (2048x1536) = 6 tiles × 258 = 1548",
                 estimate_image_input_tokens(2048, 1536) == 1548)
    report.check("unknown dims → fallback 258",
                 estimate_image_input_tokens(None, None) == 258)

    # estimate_ocr_cost with sample occurrences
    occs_sample = [
        {"id": "OCC_0", "data": _PNG_1X1 * 100, "width_px": 1024, "height_px": 768},
        {"id": "OCC_1", "data": b"x" * 100,   "width_px": 50,   "height_px": 50},   # < 5KB → skip
    ]
    est = estimate_ocr_cost(occs_sample, skip_under_bytes=5000)
    report.check("estimate skips small image",
                 est["n_skipped"] == 1 and est["n_to_ocr"] == 1,
                 f"skipped={est['n_skipped']}, to_ocr={est['n_to_ocr']}")
    report.check("estimate has USD/VND > 0",
                 est["usd"] > 0 and est["vnd"] > 0,
                 f"${est['usd']:.6f} / {est['vnd']:.2f} VND")

    # ── T2: prompt builder ──────────────────────────────────────────────
    p_simple  = _build_ocr_prompt(None, "Vietnamese", None, None)
    p_full    = _build_ocr_prompt("English", "Vietnamese",
                                  {"car": "cabin"}, {"elevator"})
    report.check("prompt without source_lang has no 'from ...'",
                 "from " not in p_simple.split("Translate")[1][:30],
                 "OK simple prompt")
    report.check("prompt with source_lang says 'from English to Vietnamese'",
                 "from English to Vietnamese" in p_full)
    report.check("prompt with elevator subdomain has domain block",
                 "elevator/escalator engineering" in p_full)
    report.check("prompt with glossary includes term",
                 "car → cabin" in p_full)
    report.check("prompt always asks for regions[bbox]",
                 "regions" in p_simple and "bbox" in p_simple)

    # ── T3: occurrence extraction (same-media multi occurrence) ────────
    # Build a DOCX with the SAME image inserted 3 times
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_paragraph("Para 1")
    p1 = doc.add_paragraph()
    p1.add_run().add_picture(io.BytesIO(_PNG_1X1), width=Inches(0.4))
    doc.add_paragraph("Para 2")
    p2 = doc.add_paragraph()
    p2.add_run().add_picture(io.BytesIO(_PNG_1X1), width=Inches(0.4))
    doc.add_paragraph("Para 3")
    p3 = doc.add_paragraph()
    p3.add_run().add_picture(io.BytesIO(_PNG_1X1), width=Inches(0.4))
    out = io.BytesIO(); doc.save(out)
    src_docx = out.getvalue()

    occs = extract_image_occurrences(src_docx)
    report.check("3 occurrences extracted from 3 inserts",
                 len(occs) == 3,
                 f"got {len(occs)}")
    report.check("occurrence_index counts up per (part, rId)",
                 sorted(o["occurrence_index"] for o in occs) == [0, 1, 2] or
                 sorted(o["occurrence_index"] for o in occs) == [0, 0, 0],
                 # python-docx tạo 3 rId riêng → mỗi occ là 0; nếu chung rId thì 0,1,2
                 f"occurrence_indexes = {[o['occurrence_index'] for o in occs]}")
    report.check("each occurrence has rId, doc_part, paragraph_index",
                 all(o.get("rId") and o.get("doc_part") and "paragraph_index" in o
                     for o in occs))

    # ── T4: caption mode chọn từng ảnh ──────────────────────────────────
    # No selection → 0 captions inserted, DOCX unchanged
    ocr_results = {
        o["id"]: {"has_text": True, "ocr": f"text {i}", "translation": f"dịch {i}",
                  "regions": [], "confidence": 0.9}
        for i, o in enumerate(occs)
    }

    out_none = insert_ocr_captions_into_docx(src_docx, occs, ocr_results,
                                              selected_ids=[])
    report.check("selected_ids=[] → DOCX unchanged",
                 out_none == src_docx)

    # Pick only first → exactly 1 caption inserted
    only_first = [occs[0]["id"]]
    out_one = insert_ocr_captions_into_docx(src_docx, occs, ocr_results,
                                             selected_ids=only_first)
    # Count [OCR] occurrences in document.xml
    import zipfile
    with zipfile.ZipFile(io.BytesIO(out_one)) as zf:
        doc_xml = zf.read("word/document.xml")
    n_captions = doc_xml.count(b"[OCR] ")
    report.check("1 selection → 1 caption inserted",
                 n_captions == 1,
                 f"found {n_captions} [OCR] markers")

    # Pick all 3
    out_all = insert_ocr_captions_into_docx(src_docx, occs, ocr_results,
                                             selected_ids=[o["id"] for o in occs])
    with zipfile.ZipFile(io.BytesIO(out_all)) as zf:
        doc_xml = zf.read("word/document.xml")
    n_captions = doc_xml.count(b"[OCR] ")
    report.check("3 selections → 3 captions inserted",
                 n_captions == 3,
                 f"found {n_captions} [OCR] markers")

    # Re-apply same selection → dedupe (no duplicate captions)
    out_again = insert_ocr_captions_into_docx(out_all, occs, ocr_results,
                                               selected_ids=[o["id"] for o in occs])
    with zipfile.ZipFile(io.BytesIO(out_again)) as zf:
        doc_xml = zf.read("word/document.xml")
    n_captions = doc_xml.count(b"[OCR] ")
    report.check("Re-applying same selection does NOT duplicate captions",
                 n_captions == 3,
                 f"after re-apply: {n_captions} [OCR] markers")

    # Edited translations override
    edited = {occs[0]["id"]: "EDITED TEXT"}
    out_edit = insert_ocr_captions_into_docx(src_docx, occs, ocr_results,
                                              selected_ids=[occs[0]["id"]],
                                              edited_translations=edited)
    with zipfile.ZipFile(io.BytesIO(out_edit)) as zf:
        doc_xml = zf.read("word/document.xml")
    report.check("edited_translations override base translation",
                 b"[OCR] EDITED TEXT" in doc_xml)

finally:
    cleanup()


sys.exit(report.summary())

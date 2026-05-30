"""Smoke check: coverage cho 2 vùng trước đây bị bỏ sót khi dịch.

1. Auto-TOC nằm trong content control <w:sdt> (python-docx `doc.paragraphs`
   không đi vào sdt → mục lục không được dịch).
2. Header/Footer first-page + even-page (python-docx `section.header/footer`
   chỉ trả bản mặc định → bảng title-page trang đầu + footer trang đầu/chẵn
   không được dịch).

Đồng thời kiểm tra TÍNH TẤT ĐỊNH của `iter_all_paragraphs` (extract ↔ apply
phải cùng thứ tự, nếu lệch → apply ghi nhầm paragraph).

Chạy: python tests/smoke_sdt_hf.py
"""
from __future__ import annotations

import io
import os
import sys

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from tests._helpers import setup_env, Report


def _make_docx_with_sdt_and_special_hf() -> bytes:
    """DOCX có: body paragraphs + 1 sdt chứa TOC-like paragraph +
    header/footer mặc định, first-page và even-page khác nhau."""
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn, nsdecls

    doc = Document()
    doc.add_paragraph("Body paragraph one.")
    doc.add_paragraph("Body paragraph two.")

    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    try:
        doc.settings.odd_and_even_pages_header_footer = True
    except Exception:
        pass

    sec.header.add_paragraph("DEFAULT HEADER")
    sec.first_page_header.add_paragraph("FIRSTPAGE HEADER")
    sec.even_page_header.add_paragraph("EVENPAGE HEADER")
    sec.footer.add_paragraph("DEFAULT FOOTER")
    sec.first_page_footer.add_paragraph("FIRSTPAGE FOOTER")
    sec.even_page_footer.add_paragraph("EVENPAGE FOOTER")

    # Inject block-level sdt (mô phỏng auto-TOC) trước sectPr của body.
    # Dùng hyperlink anchor "_Toc..." (như auto-TOC thật) để `_is_toc_paragraph`
    # nhận diện — style "TOC1" không định nghĩa trong styles.xml synthetic nên
    # không dựa vào style được.
    sdt = parse_xml(
        f'<w:sdt {nsdecls("w")}>'
        f'<w:sdtPr/><w:sdtEndPr/>'
        f'<w:sdtContent>'
        f'<w:p><w:pPr><w:pStyle w:val="TOC1"/></w:pPr>'
        f'<w:hyperlink w:anchor="_Toc123456">'
        f'<w:r><w:t>TOCENTRY HEADING</w:t></w:r></w:hyperlink></w:p>'
        f'</w:sdtContent>'
        f'</w:sdt>'
    )
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr.addprevious(sdt)
    else:
        body.append(sdt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


cleanup = setup_env()
try:
    from word_backend import (
        extract_docx_blocks, apply_translations, iter_all_paragraphs,
        validate_docx_output,
    )
    from docx import Document

    report = Report("smoke_sdt_hf")

    src = _make_docx_with_sdt_and_special_hf()
    blocks = extract_docx_blocks(src)
    all_text = {b["text"] for b in blocks}

    # ── SDT / TOC coverage ───────────────────────────────────────────────
    report.check("sdt (TOC) paragraph extracted",
                 "TOCENTRY HEADING" in all_text)
    report.check("sdt paragraph detected as role 'toc' (TOC* style)",
                 any(b["text"] == "TOCENTRY HEADING" and b["role"] == "toc"
                     for b in blocks))

    # ── First-page / even-page header+footer coverage ────────────────────
    for label in ("FIRSTPAGE HEADER", "EVENPAGE HEADER",
                  "FIRSTPAGE FOOTER", "EVENPAGE FOOTER", "DEFAULT HEADER",
                  "DEFAULT FOOTER"):
        report.check(f"{label!r} extracted",
                     label in all_text)

    # ── Determinism: hai walk phải cho thứ tự y hệt ──────────────────────
    w1 = [p.text for p in iter_all_paragraphs(Document(io.BytesIO(src)))]
    w2 = [p.text for p in iter_all_paragraphs(Document(io.BytesIO(src)))]
    report.check("iter_all_paragraphs deterministic across walks",
                 w1 == w2,
                 f"len {len(w1)} vs {len(w2)}")

    # ── Apply: bản dịch phải ghi đúng vào mọi vùng ───────────────────────
    trans = {b["id"]: "VI_" + b["text"]
             for b in blocks if b["text"] and b["role"] != "image_alt"}
    out = apply_translations(src, blocks, trans)

    import zipfile

    def part_text(data: bytes, name: str) -> str:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            if name not in zf.namelist():
                return ""
            return zf.read(name).decode("utf-8", "ignore")

    doc_xml = part_text(out, "word/document.xml")
    report.check("sdt/TOC translated in document.xml",
                 "VI_TOCENTRY HEADING" in doc_xml)

    # Gom toàn bộ header*/footer* parts để kiểm các vùng H/F đã dịch
    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        hf_names = [n for n in zf.namelist()
                    if n.startswith("word/header") or n.startswith("word/footer")]
    hf_blob = "".join(part_text(out, n) for n in hf_names)
    for label in ("FIRSTPAGE HEADER", "EVENPAGE HEADER",
                  "FIRSTPAGE FOOTER", "EVENPAGE FOOTER"):
        report.check(f"{label!r} translated in output H/F part",
                     "VI_" + label in hf_blob)

    # ── Output vẫn hợp lệ ────────────────────────────────────────────────
    val = validate_docx_output(out, original_bytes=src)
    report.check("validate passes", val["valid"], f"errors={val['errors']}")

finally:
    cleanup()


sys.exit(report.summary())

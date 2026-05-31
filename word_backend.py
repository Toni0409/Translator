"""
Word backend: extract DOCX paragraphs, gọi Gemini (fallback chain + retry +
parallel chunks), ghi lại DOCX giữ nguyên format runs.

3 nhóm logic:
1. Extract  : `extract_docx_blocks`, `iter_all_paragraphs`, `_detect_role`,
              `chunk_blocks` (adaptive theo char count)
2. Translate: `translate_chunk` (raw, raises), `translate_chunk_with_retry`,
              `translate_parallel` (background worker entry)
3. Render   : `apply_translations`, `replace_paragraph_text_keep_format`
"""
import hashlib
import io
import json
import re
import time
import threading
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from config import (
    WORD_MODELS, MAX_WORD_TOKENS, MAX_WORD_WORKERS, NO_TRANSLATE_ROLES,
    CHUNK_RETRIES, TARGET_CHUNK_CHARS, MIN_CHUNK_BLOCKS, MAX_CHUNK_BLOCKS,
    HF_REPEAT_THRESHOLD, HF_REPEAT_MIN_CHARS,
    COVERAGE_RETRY_ROUNDS, RETRY_SUBCHUNK_BLOCKS,
)
from gemini import generate, usage_tokens, parse_json_loose


# Module-level cache cho model đang work (persist giữa các Streamlit rerun + thread).
# Cần lock vì 4 worker thread cùng đọc/ghi.
_working_model: list = [None]
_model_lock = threading.Lock()


# ══════════════════════════════════════════════════════════════════════════════
# PARAGRAPH ITERATION
# ══════════════════════════════════════════════════════════════════════════════
def _iter_paragraphs_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            for nested in cell.tables:
                yield from _iter_paragraphs_in_table(nested)


def _iter_footnote_endnote_paragraphs(doc):
    """
    Yield (Paragraph, hint) cho footnotes và endnotes.
    hint = "footnote" | "endnote".
    Bỏ qua separator/continuation (id = -1, 0).
    """
    SKIP_IDS = {"-1", "0"}
    for rel in doc.part.rels.values():
        rt = rel.reltype
        if "footnotes" not in rt and "endnotes" not in rt:
            continue
        hint = "footnote" if "footnotes" in rt else "endnote"
        try:
            root = rel.target_part._element
            for fn in root:
                if fn.get(qn("w:id"), "") in SKIP_IDS:
                    continue
                for p_elem in fn.iter(qn("w:p")):
                    yield Paragraph(p_elem, parent=doc), hint
        except Exception:
            pass


def _iter_comment_paragraphs(doc):
    """Yield (Paragraph, "comment") for all comment body paragraphs."""
    for rel in doc.part.rels.values():
        if "comments" not in rel.reltype:
            continue
        try:
            root = rel.target_part._element
            for comment in root:
                for p_elem in comment.iter(qn("w:p")):
                    yield Paragraph(p_elem, parent=doc), "comment"
        except Exception:
            pass


def _iter_blocks_with_meta(doc):
    """
    Single-walk generator → yield (paragraph, meta_dict) theo thứ tự ổn định.

    Meta = {
      "role_hint":  "body" | "header" | "footer" | "textbox"
                    | "footnote" | "endnote",
      "in_table":   bool,
      "table_cell": (T#, R#, C#) tuple 1-based, hoặc None,
    }

    Lý do single-walk: lxml element proxy có Python `id()` KHÔNG ổn định
    (cùng XML element được wrap nhiều proxy khác nhau ở các thời điểm khác nhau,
    và id có thể collide). → KHÔNG dùng `id(elem)` làm dict key cross-walk.
    → Dùng position index (idx) khi cần map extract ↔ apply.
    """
    t_counter = [0]

    def _walk_table(table, t_idx, role_hint):
        for r_idx, row in enumerate(table.rows, 1):
            for c_idx, cell in enumerate(row.cells, 1):
                for para in cell.paragraphs:
                    yield para, {
                        "role_hint": role_hint,
                        "in_table":  True,
                        "table_cell": (t_idx, r_idx, c_idx),
                    }
                for nested in cell.tables:
                    yield from _walk_table(nested, t_idx, role_hint)

    # Body paragraphs
    for para in doc.paragraphs:
        hint = "toc" if _is_toc_paragraph(para) else "body"
        yield para, {"role_hint": hint, "in_table": False, "table_cell": None}
    # Body tables
    for tbl in doc.tables:
        t_counter[0] += 1
        yield from _walk_table(tbl, t_counter[0], "body")
    # Block-level content controls (w:sdt) trong BODY — python-docx KHÔNG đi vào
    # <w:sdt> nên auto-TOC (Word bọc Table of Contents trong sdtContent) và
    # rich-text content control bị bỏ sót hoàn toàn (→ "mục lục chưa dịch").
    # Chỉ duyệt sdt TOP-LEVEL (không lồng trong sdt khác) rồi yield mọi <w:p> bên
    # trong → mỗi đoạn yield đúng 1 lần. Dedup bằng quan hệ ANCESTOR, KHÔNG dùng
    # id() proxy lxml: id() bất ổn/đụng độ cross-walk (xem chú thích đầu hàm) →
    # nếu dùng sẽ yield số đoạn khác nhau giữa extract↔apply, lệch hết para_idx.
    _SDT, _SDTC, _TXBX = qn("w:sdt"), qn("w:sdtContent"), qn("w:txbxContent")

    def _has_ancestor(el, tag, stop) -> bool:
        a = el.getparent()
        while a is not None and a is not stop:
            if a.tag == tag:
                return True
            a = a.getparent()
        return False

    body_el = doc.element.body
    for sdt in body_el.iter(_SDT):
        if _has_ancestor(sdt, _SDT, body_el):          # nested sdt → để sdt cha xử lý
            continue
        content = sdt.find(_SDTC)
        if content is None:
            continue
        for p_elem in content.iter(qn("w:p")):
            if _has_ancestor(p_elem, _TXBX, content):  # text-box xử lý ở sweep riêng
                continue
            para = Paragraph(p_elem, parent=doc)
            hint = "toc" if _is_toc_paragraph(para) else "body"
            yield para, {"role_hint": hint, "in_table": False, "table_cell": None}
    # Headers + footers — TẤT CẢ loại: default + first-page + even-page.
    # python-docx `section.header/footer` CHỈ trả header/footer mặc định →
    # bỏ sót bảng title-page (first-page header) và footer trang đầu/chẵn khi doc
    # bật "different first page" (titlePg). `is_linked_to_previous` = phần này kế
    # thừa từ section trước (đã xử lý rồi) → skip để khỏi dịch trùng.
    for section in doc.sections:
        for hf, hint in (
            (section.header,            "header"),
            (section.first_page_header, "header"),
            (section.even_page_header,  "header"),
            (section.footer,            "footer"),
            (section.first_page_footer, "footer"),
            (section.even_page_footer,  "footer"),
        ):
            try:
                if hf.is_linked_to_previous:
                    continue
                for para in hf.paragraphs:
                    yield para, {"role_hint": hint, "in_table": False, "table_cell": None}
                for tbl in hf.tables:
                    t_counter[0] += 1
                    yield from _walk_table(tbl, t_counter[0], hint)
            except Exception:
                pass
    # Text-boxes / shapes trong body document part
    # (text-boxes trong header/footer parts chưa cover ở iteration này)
    # parent=doc để `para.style` resolve được qua doc.part.styles
    for txbx in doc.element.iter(qn("w:txbxContent")):
        for p_elem in txbx.iter(qn("w:p")):
            yield Paragraph(p_elem, parent=doc), {
                "role_hint": "textbox", "in_table": False, "table_cell": None,
            }
    # Footnotes + endnotes
    for para, hint in _iter_footnote_endnote_paragraphs(doc):
        yield para, {"role_hint": hint, "in_table": False, "table_cell": None}
    # VML text-boxes (Word 2003/older format) — w:txbxContent usually already caught above;
    # this handles the rare case where VML v:textbox has direct w:p without w:txbxContent
    _VML_TB = "{urn:schemas-microsoft-com:vml}textbox"
    _seen_p = {id(p_elem) for txbx in doc.element.iter(qn("w:txbxContent"))
               for p_elem in txbx.iter(qn("w:p"))}
    for vml_tb in doc.element.iter(_VML_TB):
        for p_elem in vml_tb.iter(qn("w:p")):
            if id(p_elem) not in _seen_p:
                yield Paragraph(p_elem, parent=doc), {
                    "role_hint": "textbox", "in_table": False, "table_cell": None,
                }
    # Comments (word/comments.xml)
    for para, hint in _iter_comment_paragraphs(doc):
        yield para, {"role_hint": hint, "in_table": False, "table_cell": None}


def iter_all_paragraphs(doc):
    """
    Lặp qua toàn bộ paragraph: body → tables → headers → footers → text-boxes.
    Thin wrapper around `_iter_blocks_with_meta` — guarantee cùng thứ tự với extract.
    """
    for para, _meta in _iter_blocks_with_meta(doc):
        yield para


# ══════════════════════════════════════════════════════════════════════════════
# INLINE FORMAT — encode/decode bold/italic/underline qua tag để AI giữ được
# ══════════════════════════════════════════════════════════════════════════════
_TAG_RE      = re.compile(r"</?[biu]>")
_TAG_STRIP   = re.compile(r"</?[biu]>")
_ESC_MAP     = (("&", "&amp;"), ("<", "&lt;"), (">", "&gt;"))
_UNESC_MAP   = (("&lt;", "<"), ("&gt;", ">"), ("&amp;", "&"))


def _esc(s: str) -> str:
    for a, b in _ESC_MAP:
        s = s.replace(a, b)
    return s


def _unesc(s: str) -> str:
    for a, b in _UNESC_MAP:
        s = s.replace(a, b)
    return s


def runs_to_tagged_text(paragraph) -> str:
    """
    Convert paragraph runs → tagged text. VD: '<b>Important</b>: read <i>carefully</i>'.
    Escape & < > trong text gốc để không xung đột với tag.
    Bao gồm text bên trong w:hyperlink (URL được giữ nguyên khi apply).
    Math equations (m:oMath) → <MATH/> placeholder.
    Field runs (w:fldChar, w:instrText) are skipped — kept as-is.
    """
    from docx.text.run import Run
    _MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    parts = []
    for child in paragraph._p.iterchildren():
        # Math equations — preserve verbatim with placeholder
        if child.tag == f"{{{_MATH_NS}}}oMath":
            parts.append("<MATH/>")
            continue
        tag = child.tag.split("}", 1)[-1]
        if tag == "r":
            # Skip field instruction/delimiter runs
            if child.find(qn("w:instrText")) is not None:
                continue
            fld = child.find(qn("w:fldChar"))
            if fld is not None:
                continue
            run = Run(child, paragraph)
            text = run.text
            if not text:
                continue
            text = _esc(text)
            if run.bold:      text = f"<b>{text}</b>"
            if run.italic:    text = f"<i>{text}</i>"
            if run.underline: text = f"<u>{text}</u>"
            parts.append(text)
        elif tag == "hyperlink":
            # Include hyperlink inner text — URL stays in XML, only text is translated
            for r_elem in child.iterchildren(qn("w:r")):
                if r_elem.find(qn("w:instrText")) is not None:
                    continue
                if r_elem.find(qn("w:fldChar")) is not None:
                    continue
                run = Run(r_elem, paragraph)
                text = run.text
                if not text:
                    continue
                text = _esc(text)
                if run.bold:      text = f"<b>{text}</b>"
                if run.italic:    text = f"<i>{text}</i>"
                if run.underline: text = f"<u>{text}</u>"
                parts.append(text)
        elif tag == "ins":
            # Track changes: inserted text — include for translation
            for r_elem in child.iterchildren(qn("w:r")):
                if r_elem.find(qn("w:instrText")) is not None:
                    continue
                if r_elem.find(qn("w:fldChar")) is not None:
                    continue
                run = Run(r_elem, paragraph)
                text = run.text
                if not text:
                    continue
                text = _esc(text)
                if run.bold:      text = f"<b>{text}</b>"
                if run.italic:    text = f"<i>{text}</i>"
                if run.underline: text = f"<u>{text}</u>"
                parts.append(text)
        # w:del: intentionally skipped — deleted text not translated
    return "".join(parts).strip()


def has_inline_format(tagged: str) -> bool:
    return _TAG_RE.search(tagged) is not None


def strip_tags(tagged: str) -> str:
    """Bỏ tag, unescape — dùng khi fallback về plain text."""
    return _unesc(_TAG_STRIP.sub("", tagged))


def _parse_tagged(tagged: str) -> list[tuple[str, bool, bool, bool]]:
    """
    Parse tagged text → list of (text, bold, italic, underline).
    Tolerant với tag mismatch (ignore).
    """
    segments: list = []
    stack:    list = []
    last           = 0

    def flush(text):
        if text:
            segments.append((_unesc(text), "b" in stack, "i" in stack, "u" in stack))

    for m in _TAG_RE.finditer(tagged):
        if m.start() > last:
            flush(tagged[last:m.start()])
        tag = m.group()
        if tag.startswith("</"):
            if stack and stack[-1] == tag[2]:
                stack.pop()
        else:
            stack.append(tag[1])
        last = m.end()
    if last < len(tagged):
        flush(tagged[last:])
    return segments


def _paragraph_has_non_run_children(paragraph) -> bool:
    """True nếu paragraph có hyperlink/field/... — không rebuild runs an toàn được.

    Cũng trả True nếu paragraph chứa drawing/pict/object/AlternateContent ở bất kỳ
    cấp độ nào (kể cả trong run). Lý do: rebuild runs sẽ xoá những element này → mất ảnh.
    """
    SAFE = {"pPr", "r", "hyperlink", "ins", "del",
            "bookmarkStart", "bookmarkEnd", "proofErr", "fldSimple",
            "rPr", "pPrChange", "oMath"}
    DANGER_DESC = {"drawing", "pict", "object", "AlternateContent"}
    for child in paragraph._p.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag not in SAFE:
            return True
    # Quét sâu — drawing/pict thường nằm trong w:r, vẫn nguy hiểm khi rebuild
    for descendant in paragraph._p.iter():
        d_tag = descendant.tag.split("}", 1)[-1]
        if d_tag in DANGER_DESC:
            return True
    return False


def replace_paragraph_text_keep_format(paragraph, new_text: str):
    """
    Replace paragraph text keeping format of first run.
    Handles hyperlinks: collects ALL w:r elements (regular + inside w:hyperlink)
    so translated text is written correctly even in hyperlink-only paragraphs.
    Hyperlink URLs (relationships) are untouched; only inner text changes.
    """
    # Collect all w:r elements in document order, including inside w:hyperlink
    # Field runs (w:fldChar, w:instrText) are excluded — keep original XML
    all_r = []
    for child in paragraph._p.iterchildren():
        ctag = child.tag.split("}", 1)[-1]
        if ctag == "r":
            # Field runs — skip from translation
            has_fld = (child.find(qn("w:fldChar")) is not None or
                       child.find(qn("w:instrText")) is not None)
            if has_fld:
                continue
            all_r.append(child)
        elif ctag == "hyperlink":
            for r_elem in child.iterchildren(qn("w:r")):
                has_fld = (r_elem.find(qn("w:fldChar")) is not None or
                           r_elem.find(qn("w:instrText")) is not None)
                if not has_fld:
                    all_r.append(r_elem)
        elif ctag == "ins":
            # Track changes insertion: include runs for translation
            for r_elem in child.iterchildren(qn("w:r")):
                has_fld = (r_elem.find(qn("w:fldChar")) is not None or
                           r_elem.find(qn("w:instrText")) is not None)
                if not has_fld:
                    all_r.append(r_elem)
        # del: skip — deleted text untouched

    if not all_r:
        paragraph.add_run(new_text)
        return

    # Find first run that has non-empty text (use as formatting template)
    first_r = next(
        (r for r in all_r
         if any((t.text or "").strip() for t in r.findall(qn("w:t")))),
        all_r[0],
    )

    # Write new_text into first_r
    t_elems = first_r.findall(qn("w:t"))
    if t_elems:
        t_elems[0].text = new_text
        if new_text and (new_text[0] == " " or new_text[-1] == " "):
            t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        for t in t_elems[1:]:
            t.text = ""
    else:
        from lxml import etree
        t = etree.SubElement(first_r, qn("w:t"))
        t.text = new_text
        if new_text and (new_text[0] == " " or new_text[-1] == " "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Clear all other runs (including hyperlink runs)
    for r_elem in all_r:
        if r_elem is first_r:
            continue
        for t in r_elem.findall(qn("w:t")):
            t.text = ""


def _any_run_has_media(paragraph) -> bool:
    """True nếu bất kỳ <w:r> nào trong paragraph chứa drawing/pict/object."""
    for r in paragraph._p.iter(qn("w:r")):
        for tag in ("w:drawing", "w:pict", "w:object"):
            if r.find(f".//{qn(tag)}") is not None:
                return True
    return False


def replace_paragraph_with_tagged(paragraph, tagged_text: str):
    """
    Path 2 (new): parse <b><i><u> tags trong tagged_text và rebuild runs để
    giữ inline format. Inherit toàn bộ run-properties (rPr) qua deepcopy của
    template's w:rPr — fallback name/size/color nếu deepcopy fail.

    Fallback về `replace_paragraph_text_keep_format` nếu paragraph có
    hyperlink/field/drawing, hoặc không parse được tag.
    """
    # P2.3/P2.4: bất kỳ run nào chứa media → KHÔNG rebuild runs (mất ảnh)
    if _paragraph_has_non_run_children(paragraph) or _any_run_has_media(paragraph):
        replace_paragraph_text_keep_format(paragraph, strip_tags(tagged_text))
        return

    segments = _parse_tagged(tagged_text)
    if not segments:
        replace_paragraph_text_keep_format(paragraph, strip_tags(tagged_text))
        return

    # Template format: ưu tiên deepcopy w:rPr; fallback name/size/color
    from copy import deepcopy
    runs     = paragraph.runs
    template = next((r for r in runs if r.text), runs[0] if runs else None)
    tpl_rpr = None
    tpl_name = tpl_size = tpl_color = None
    if template is not None:
        try:
            rpr_elem = template._r.find(qn("w:rPr"))
            if rpr_elem is not None:
                tpl_rpr = deepcopy(rpr_elem)
        except Exception:
            tpl_rpr = None
        try: tpl_name  = template.font.name
        except Exception: pass
        try: tpl_size  = template.font.size
        except Exception: pass
        try: tpl_color = template.font.color.rgb
        except Exception: pass

    # Xóa hết runs cũ (an toàn vì đã guard media ở trên)
    p_elem = paragraph._p
    for run in list(runs):
        try: p_elem.remove(run._r)
        except Exception: pass

    # Thêm runs mới theo segments
    for text, bold, italic, underline in segments:
        if not text:
            continue
        run = paragraph.add_run(text)
        # Apply deepcopy rPr trước, sau đó override bold/italic/underline
        if tpl_rpr is not None:
            try:
                # Xoá rPr mặc định run mới tạo nếu có, rồi insert template rPr
                existing = run._r.find(qn("w:rPr"))
                if existing is not None:
                    run._r.remove(existing)
                run._r.insert(0, deepcopy(tpl_rpr))
            except Exception:
                # Fallback: dùng font name/size/color
                if tpl_name:  run.font.name = tpl_name
                if tpl_size:  run.font.size = tpl_size
                if tpl_color:
                    try: run.font.color.rgb = tpl_color
                    except Exception: pass
        else:
            if tpl_name:  run.font.name = tpl_name
            if tpl_size:  run.font.size = tpl_size
            if tpl_color:
                try: run.font.color.rgb = tpl_color
                except Exception: pass
        # Toggle bold/italic/underline theo tagged text — set sau deepcopy để override
        if bold:      run.bold      = True
        if italic:    run.italic    = True
        if underline: run.underline = True


# ══════════════════════════════════════════════════════════════════════════════
# ROLE DETECTION (rule-based, không gọi AI)
# ══════════════════════════════════════════════════════════════════════════════
def _detect_role(para, in_table: bool, in_header: bool, in_footer: bool) -> str:
    if in_header: return "header"
    if in_footer: return "footer"
    style = (para.style.name or "").lower()
    if "toc" in style:     return "toc"
    if "heading" in style: return "section_heading"
    if in_table:           return "table_cell"
    try:
        pPr = para._p.pPr
        if pPr is not None and pPr.numPr is not None:
            return "bullet"
    except Exception:
        pass
    if style in ("title", "subtitle"):
        return "title"
    if any(k in style for k in ("caption", "note", "warning", "caution")):
        return "note"
    return "paragraph"


# ══════════════════════════════════════════════════════════════════════════════
# TOC DETECTION
# ══════════════════════════════════════════════════════════════════════════════
def _is_toc_paragraph(paragraph) -> bool:
    """Detect if paragraph is part of a Table of Contents."""
    try:
        style = (paragraph.style.name or "").lower()
        if style.startswith("toc"):
            return True
    except Exception:
        pass
    # Check for hyperlink to _Toc bookmark
    try:
        for h in paragraph._p.iter(qn("w:hyperlink")):
            anchor = h.get(qn("w:anchor"), "")
            if anchor.startswith("_Toc"):
                return True
    except Exception:
        pass
    return False


def _link_toc_to_headings(blocks: list[dict]) -> int:
    """Set _toc_mirror on TOC blocks pointing to matching heading block id."""
    headings = {}
    for b in blocks:
        if b["role"] == "section_heading":
            key = b["text"].strip().lower()
            if key and key not in headings:
                headings[key] = b["id"]
    linked = 0
    for b in blocks:
        if b["role"] != "toc":
            continue
        raw = b["text"]
        # Strip trailing tab+digits or dots+digits (page number suffix)
        m = re.match(r"^(.*?)(?:[\t\s\.]+\d+)?$", raw)
        head_text = (m.group(1) if m else raw).strip().lower()
        if head_text in headings:
            b["_toc_mirror"] = headings[head_text]
            linked += 1
    return linked


# ══════════════════════════════════════════════════════════════════════════════
# EXTRACT
# ══════════════════════════════════════════════════════════════════════════════
def _mark_repeating_as_hf(blocks: list[dict],
                          threshold: int = HF_REPEAT_THRESHOLD,
                          min_chars: int = HF_REPEAT_MIN_CHARS) -> int:
    """
    Phát hiện block trong body có text lặp lại ≥ threshold lần
    (text dài ≥ min_chars) → đánh dấu role = 'body_repeated'.

    Đây là heuristic phổ biến để bắt header/footer ẩn trong body
    (vd: tiêu đề chapter chạy trên mỗi trang, watermark "Confidential",
    section banner...).

    Trả về số block bị đánh dấu.
    """
    counter: Counter = Counter()
    for b in blocks:
        if b["role"] in ("header", "footer"):
            continue   # đã là H/F thật
        if len(b["text"]) < min_chars:
            continue   # text quá ngắn, dễ false positive
        counter[b["text"]] += 1

    repeating_texts = {t for t, c in counter.items() if c >= threshold}
    marked = 0
    for b in blocks:
        if b["role"] in ("header", "footer"):
            continue
        if b["text"] in repeating_texts:
            b["role"] = "body_repeated"
            marked += 1
    return marked


def _iter_image_alt_texts(doc):
    """
    Yield (element, attr_name, text) for every image alt-text/title in the doc.
    Caller can later set element.set(attr_name, translated_text).

    Covers:
    - drawingML wp:docPr (descr, title)
    - drawingML pic:cNvPr (descr, title)
    - VML v:shape (alt)
    """
    WP_NS  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    VML_NS = "urn:schemas-microsoft-com:vml"

    for tag in (f"{{{WP_NS}}}docPr", f"{{{PIC_NS}}}cNvPr"):
        for el in doc.element.iter(tag):
            for attr in ("descr", "title"):
                val = (el.get(attr) or "").strip()
                if val:
                    yield el, attr, val

    for el in doc.element.iter(f"{{{VML_NS}}}shape"):
        val = (el.get("alt") or "").strip()
        if val:
            yield el, "alt", val


def _paragraph_has_media(paragraph) -> bool:
    """True nếu paragraph chứa drawing/picture/object (ảnh, shape, embed)."""
    for tag in ("w:drawing", "w:pict", "w:object"):
        if paragraph._p.find(f".//{qn(tag)}") is not None:
            return True
    return False


def extract_docx_blocks(docx_bytes: bytes) -> list[dict]:
    """
    Đọc DOCX → list block {id, text, role, para_idx, table_cell, ...}.

    Single-walk extraction (xem `_iter_blocks_with_meta` để biết tại sao).
    Role:
    - header / footer / textbox : từ meta hint
    - body_repeated             : text lặp lại ≥ 3 lần trong body (heuristic H/F)
    - media_only                : paragraph rỗng text nhưng có ảnh/drawing → giữ nguyên,
                                  không dịch (P2: trước đây bị skip → mất ảnh)
    - section_heading / title / bullet / table_cell / note / toc / paragraph
    """
    doc = Document(io.BytesIO(docx_bytes))
    blocks = []
    for idx, (para, meta) in enumerate(_iter_blocks_with_meta(doc)):
        text = para.text.strip()
        if not text:
            # Empty paragraph: nếu có media → vẫn track để apply_translations
            # không xoá nhầm. Skip nếu hoàn toàn trống.
            if _paragraph_has_media(para):
                blocks.append({
                    "id":          f"p{idx}",
                    "text":        "",
                    "text_tagged": "",
                    "has_format":  False,
                    "role":        "media_only",
                    "para_idx":    idx,
                    "table_cell":  meta["table_cell"],
                })
            continue
        role = _detect_role(
            para,
            meta["in_table"],
            meta["role_hint"] == "header",
            meta["role_hint"] == "footer",
        )
        # Override special roles (chỉ khi không phải H/F thật)
        if meta["role_hint"] == "textbox" and role not in ("header", "footer"):
            role = "textbox"
        if meta["role_hint"] in ("footnote", "endnote") and role not in ("header", "footer"):
            role = meta["role_hint"]
        if meta["role_hint"] == "comment" and role not in ("header", "footer"):
            role = "comment"
        if meta["role_hint"] == "toc" and role not in ("header", "footer"):
            role = "toc"
        tagged = runs_to_tagged_text(para)
        blocks.append({
            "id":          f"p{idx}",
            "text":        text,
            "text_tagged": tagged,
            "has_format":  has_inline_format(tagged),
            "role":        role,
            "para_idx":    idx,
            "table_cell":  meta["table_cell"],   # (T,R,C) tuple hoặc None
        })
    _mark_repeating_as_hf(blocks)
    _link_toc_to_headings(blocks)

    # Image alt-texts (drawingML + VML). Appended AFTER body blocks so
    # `apply_translations` can re-iterate `_iter_image_alt_texts` and match by order.
    for i, (_, _, alt_text) in enumerate(_iter_image_alt_texts(doc)):
        blocks.append({
            "id":          f"IMG_ALT_{i}",
            "text":        alt_text,
            "text_tagged": alt_text,
            "has_format":  False,
            "role":        "image_alt",
            "para_idx":    -1,
            "table_cell":  None,
        })

    return blocks


def count_by_role(blocks: list[dict]) -> dict:
    """Stat helper: đếm số block theo từng role + tổng theo nhóm."""
    counter = Counter(b["role"] for b in blocks)
    return {
        "total":         len(blocks),
        "body":          sum(c for r, c in counter.items() if r not in NO_TRANSLATE_ROLES),
        "header":        counter.get("header", 0),
        "footer":        counter.get("footer", 0),
        "body_repeated": counter.get("body_repeated", 0),
        "footnote":      counter.get("footnote", 0),
        "endnote":       counter.get("endnote", 0),
        "comment":       counter.get("comment", 0),
        "media_only":    counter.get("media_only", 0),
        "by_role":       dict(counter),
    }


_NOUN_PHRASE_RE = re.compile(r"\b[A-Z][a-zA-Z\-]+(?:\s+[A-Z]?[a-zA-Z\-]+){0,3}\b")
_TECH_TERM_RE   = re.compile(r"\b[A-Z][a-zA-Z]{4,}\b")


def build_glossary(client, blocks: list[dict], target_lang: str,
                   source_lang: str | None = None,
                   seed: dict | None = None,
                   top_n: int = 30, min_repeat: int = 3) -> dict:
    """
    One-shot AI call để dịch top-N thuật ngữ lặp lại trong doc.
    Inject vào mọi chunk prompt để dịch consistent cross-chunk.

    Precedence (P4.5): seed → AI-extract. AI KHÔNG override seed entry đã có.
    Frontend layer sau đó merge user-imported terms với precedence cao nhất.

    Heuristic candidates:
    - Noun phrases (capitalized, 1-4 words) — "Hoistway Door", "Control Panel"
    - Single technical terms (CamelCase, ≥5 chars) — "Inverter", "Calibration"

    Trả về (ít nhất) seed dict nếu AI fail.
    """
    seed_dict = dict(seed or {})

    text_all = " ".join(b["text"] for b in blocks
                        if b["role"] not in NO_TRANSLATE_ROLES)
    if not text_all:
        return seed_dict

    candidates = _NOUN_PHRASE_RE.findall(text_all) + _TECH_TERM_RE.findall(text_all)
    counter    = Counter(candidates)
    top_terms  = [w for w, c in counter.most_common(top_n * 2) if c >= min_repeat][:top_n]
    # Bỏ những term đã có trong seed để không tốn token AI dịch lại
    top_terms  = [t for t in top_terms if t not in seed_dict]
    if not top_terms:
        return seed_dict

    src_clause = f"from {source_lang} " if source_lang else ""
    prompt = f"""Translate these technical terms {src_clause}to {target_lang}.
Return ONLY a JSON object: {{"term": "translation", ...}}

Rules:
- Use formal, consistent terminology suitable for technical documents.
- Keep proper nouns / company names / product codes / brand names UNCHANGED.
- Keep numbers, units, model codes UNCHANGED.
- Each term gets exactly ONE canonical translation.

Terms ({len(top_terms)}):
{json.dumps(top_terms, ensure_ascii=False)}"""

    try:
        raw, _, _ = call_gemini(client, prompt)
    except Exception:
        return seed_dict
    parsed = parse_json_loose(raw)
    if not isinstance(parsed, dict):
        return seed_dict
    # Filter — chỉ giữ entry hợp lệ; seed NOT overridden
    out = dict(seed_dict)
    for k, v in parsed.items():
        if not (k and v and isinstance(v, str)):
            continue
        ks = str(k).strip()
        vs = str(v).strip()
        if not ks or not vs or vs == ks:
            continue
        if ks in out:
            continue   # seed có rồi → giữ
        out[ks] = vs
    return out


_DOMAIN_STYLE_BLOCK = """\
Domain: elevator/escalator engineering.
Use formal technical register.
Preserve units verbatim: mm, m, m/s, kg, kN, V, Hz, A, W, kW, deg C, °C, dB.
Preserve standard references verbatim: EN 81-20, EN 81-50, ISO 22201, ISO 14798,
ASME A17.1, GB 7588, TCVN 6395, TCVN 6396, JIS A 4302.
Preserve part numbers, drawing numbers, revision codes (e.g. "Rev. A", "DWG-1234").
Avoid colloquialisms.
Translate terminology consistently across the document.
"""


def build_doc_context(blocks: list[dict], source_lang: str | None = None,
                      subdomains: set[str] | None = None) -> str:
    """Tóm tắt cấu trúc tài liệu (title + headings + TOC + domain hint).

    `source_lang` được nhận để consistent signature; chưa inject vì AI đã có
    `source_lang` ở header prompt (P1.8).
    `subdomains`: nếu chứa `elevator`/`escalator` → thêm domain style guide (P4.6).
    """
    titles   = [b["text"] for b in blocks if b["role"] == "title"][:2]
    headings = [b["text"] for b in blocks if b["role"] == "section_heading"][:10]
    tocs     = [b["text"] for b in blocks if b["role"] == "toc"][:8]

    lines = []
    if titles:   lines.append("Document title: " + " | ".join(titles))
    if headings: lines.append("Main sections: "  + " | ".join(headings[:6]))
    if tocs:     lines.append("TOC: "            + " | ".join(tocs[:5]))
    ctx = "\n".join(lines) if lines else "Technical document."

    # Domain style — ưu tiên subdomains explicit từ caller (P4.4); fallback heuristic
    # text-scan khi không có subdomain set.
    if subdomains and (subdomains & {"elevator", "escalator"}):
        ctx += "\n" + _DOMAIN_STYLE_BLOCK
    else:
        all_text = " ".join(b["text"] for b in blocks[:30]).lower()
        if any(k in all_text for k in ("elevator", "lift", "hoistway", "schindler", "inventio")):
            ctx += "\n" + _DOMAIN_STYLE_BLOCK
        elif any(k in all_text for k in ("safety", "standard", "regulation", "iso", "en 81")):
            ctx += "\nDomain: safety standards."
        elif any(k in all_text for k in ("software", "api", "code", "function")):
            ctx += "\nDomain: software/IT."
    return ctx


# ══════════════════════════════════════════════════════════════════════════════
# ADAPTIVE CHUNKING — chia theo char count, có min/max paragraph count
# ══════════════════════════════════════════════════════════════════════════════
def chunk_blocks(blocks: list[dict],
                 target_chars: int = TARGET_CHUNK_CHARS,
                 min_count: int   = MIN_CHUNK_BLOCKS,
                 max_count: int   = MAX_CHUNK_BLOCKS) -> list[list[dict]]:
    """
    Chia list block thành chunks sao cho:
    - Mỗi chunk có tổng char ~ target_chars
    - Tối thiểu min_count, tối đa max_count paragraph / chunk
    """
    chunks: list[list[dict]] = []
    current: list[dict]      = []
    cur_chars = 0
    for b in blocks:
        size = len(b["text"]) + 30   # +overhead JSON wrapping
        full = (cur_chars + size > target_chars and len(current) >= min_count) \
               or len(current) >= max_count
        if current and full:
            chunks.append(current)
            current, cur_chars = [], 0
        current.append(b)
        cur_chars += size
    if current:
        chunks.append(current)
    return chunks


# ══════════════════════════════════════════════════════════════════════════════
# GEMINI CALL với fallback chain + token tracking
# ══════════════════════════════════════════════════════════════════════════════
def call_gemini(client, prompt: str) -> tuple[str, int, int]:
    """
    Gọi Gemini với fallback chain. Trả về (text, in_tokens, out_tokens).
    Raise RuntimeError nếu tất cả model đều fail. Thread-safe.
    """
    with _model_lock:
        cur = _working_model[0]
    models = ([cur] if cur else []) + [m for m in WORD_MODELS if m != cur]
    last_err = "no models"
    for model in models:
        try:
            resp = generate(client, model, prompt,
                            max_output_tokens=MAX_WORD_TOKENS, temperature=0.1)
            text = (resp.text or "").strip()
            if not text:
                last_err = "empty response"
                continue
            in_t, out_t = usage_tokens(resp)
            with _model_lock:
                _working_model[0] = model
            return text, in_t, out_t
        except Exception as e:
            last_err = str(e)
    raise RuntimeError(f"All Word models failed: {last_err}")


def get_working_model() -> str | None:
    with _model_lock:
        return _working_model[0]


# ══════════════════════════════════════════════════════════════════════════════
# TRANSLATE chunk (raw + retry)
# ══════════════════════════════════════════════════════════════════════════════
_TABLE_PREFIX_RE = re.compile(r"^\s*\(T\d+\s+R\d+\s+C\d+\)\s*")


def _format_block_text(b: dict) -> str:
    """Prefix (T#R#C#) cho table cell, dùng tagged text nếu có inline format."""
    text = b.get("text_tagged") or b["text"]
    tc   = b.get("table_cell")
    if tc:
        return f"(T{tc[0]} R{tc[1]} C{tc[2]}) {text}"
    return text


def _build_chunk_prompt(chunk: list[dict], target_lang: str, doc_context: str,
                        glossary: dict | None = None,
                        custom_rules: dict | None = None,
                        source_lang: str | None = None) -> str:
    """Build prompt. Dùng `text_tagged` nếu có (preserve inline format)."""
    payload = json.dumps(
        [{
            "id":   b["id"],
            "text": _format_block_text(b),
            "role": b.get("role", "paragraph"),
        } for b in chunk],
        ensure_ascii=False,
    )

    has_format = any(b.get("has_format") for b in chunk)
    has_tables = any(b.get("table_cell") for b in chunk)

    format_rule = (
        "- Some texts contain HTML-like tags <b>, <i>, <u> marking bold/italic/underline. "
        "PRESERVE these tags EXACTLY in the translation, wrapping the equivalent translated words. "
        "Do NOT add new tags where source has none. Do NOT use markdown (** or _).\n"
        if has_format else ""
    )
    table_rule = (
        "- Table cells are prefixed with (T# R# C#) — table/row/column context. "
        "DO NOT include this prefix in the translation output. "
        "Cells in the same column (same C#) MUST use the same terminology. "
        "Row 1 (R1) is the header — keep concise. "
        "Preserve numbers, units, model codes in cells verbatim.\n"
        if has_tables else ""
    )

    glossary_section = ""
    if glossary:
        # Limit glossary size để không bloat prompt — tăng 50 → 80 (P4.7).
        # Python ≥3.7 dict bảo toàn insertion order: seed entries (push trước
        # AI-extract trong build_glossary) sẽ xuất hiện đầu list → ưu tiên cao.
        items = list(glossary.items())[:80]
        glossary_str = "\n".join(f"  {en} → {vi}" for en, vi in items)
        glossary_section = (
            f"\nGlossary (USE THESE EXACT translations for consistency across the document):\n"
            f"{glossary_str}\n"
        )

    custom_rules_section = ""
    if custom_rules:
        extra = (custom_rules.get("_extra") or "").strip()
        if extra:
            custom_rules_section += f"\nAdditional instruction: {extra}\n"
        lines = [f"- For {role}: {rule.strip()}"
                 for role, rule in custom_rules.items()
                 if role != "_extra" and isinstance(rule, str) and rule.strip()]
        if lines:
            custom_rules_section += (
                "\nCustom rules (override defaults for these roles):\n"
                + "\n".join(lines) + "\n"
            )

    translate_header = (
        f"Translate these Word document blocks from {source_lang} into {target_lang}."
        if source_lang else
        f"Translate these Word document blocks into {target_lang}."
    )

    return f"""{translate_header}

Document context:
{doc_context}{glossary_section}{custom_rules_section}

Rules:
- Return ONLY valid JSON array. No markdown, no explanation.
- Keep "id" exactly as given.
- Translate ALL words into {target_lang}. Leave nothing in the source language.
- Preserve: numbers, units, product codes, model names, document IDs, revision codes.
- Do NOT translate company names (e.g. INVENTIO AG, Schindler, Otis).
- Preserve ALL punctuation, whitespace, tabs (\\t), line breaks (\\n).
- Match source sentence count — do NOT merge or split sentences.
{format_rule}{table_rule}- Preserve <MATH/> placeholders exactly as-is (they represent math equations).
- Preserve <FIELD>...</FIELD> content unchanged.
- For toc: translate heading text only; preserve numbering, dot leaders (...), page numbers.
- For bullet: concise imperative, keep bullet symbol.
- For heading/section_heading: concise technical, match source brevity.
- For table_cell: concise, consistent across row.
- For note: keep WARNING/NOTE/CAUTION label verbatim.
- For header/footer/body_repeated: concise like a page header/footer; preserve page numbers ("Page 1 of 10"), dates, document IDs verbatim.
- For footnote/endnote: translate fully as body text; preserve footnote reference numbers (¹²³ or superscript digits) verbatim.
- For comment: translate the comment text naturally as body text.
- For image_alt: concise description (max ~100 chars), suitable as screen-reader alt-text.

Format:
[{{"id":"...","text":"translated text"}}]

Blocks:
{payload}"""


def translate_chunk(client, chunk: list[dict], target_lang: str,
                    doc_context: str,
                    glossary: dict | None = None,
                    custom_rules: dict | None = None,
                    source_lang: str | None = None) -> tuple[dict, int, int]:
    """
    Dịch 1 chunk → (translations_dict, in_tokens, out_tokens).
    RAISE exception nếu API fail.

    `translations_dict` CHỈ chứa các id mà model THỰC SỰ trả về (text non-empty,
    đã strip prefix). Block bị model bỏ sót / trả rỗng sẽ KHÔNG có trong dict —
    `translate_chunk_with_retry` chịu trách nhiệm dịch bù hoặc backfill text gốc.
    (Trước đây hàm này tự `setdefault` text gốc → response rỗng/hỏng bị "nuốt"
    thành công giả, khiến cả chunk giữ nguyên tiếng nguồn mà không ai phát hiện.)
    """
    prompt = _build_chunk_prompt(
        chunk, target_lang, doc_context,
        glossary=glossary, custom_rules=custom_rules,
        source_lang=source_lang,
    )
    raw, in_t, out_t = call_gemini(client, prompt)
    parsed = parse_json_loose(raw) or []
    if not isinstance(parsed, list):
        parsed = []
    result = {}
    for item in parsed:
        if isinstance(item, dict) and item.get("id") and item.get("text") is not None:
            text = _TABLE_PREFIX_RE.sub("", str(item["text"]))   # AI lỡ giữ prefix → strip
            if text.strip():
                result[item["id"]] = text
    return result, in_t, out_t


def _subchunks(blocks: list[dict], size: int) -> list[list[dict]]:
    """Chia list block thành các sub-chunk ≤ size phần tử."""
    size = max(1, size)
    return [blocks[i:i + size] for i in range(0, len(blocks), size)]


def _pending_blocks(chunk: list[dict], result: dict,
                    source_lang: str | None, target_lang: str | None) -> list[dict]:
    """Block cần dịch (lại): model trả thiếu/rỗng, hoặc bản dịch CÒN tiếng nguồn."""
    pending = []
    for b in chunk:
        tr = result.get(b["id"])
        if not tr or not str(tr).strip():
            pending.append(b)                                  # model bỏ sót
        elif has_source_language_residue(tr, source_lang, target_lang):
            pending.append(b)                                  # còn tiếng nguồn
    return pending


def translate_chunk_with_retry(client, chunk: list[dict], target_lang: str,
                               doc_context: str,
                               retries: int = CHUNK_RETRIES,
                               glossary: dict | None = None,
                               custom_rules: dict | None = None,
                               fallback_individually: bool = True,
                               source_lang: str | None = None,
                               ) -> tuple[dict, int, int, str | None]:
    """
    Dịch 1 chunk, đảm bảo COVERAGE. Trả về (translations, in_t, out_t, error).

    1. Dịch cả chunk, retry exponential backoff khi API raise.
    2. Dịch bù: block nào model bỏ sót / trả về còn tiếng nguồn → dịch lại theo
       sub-chunk nhỏ (RETRY_SUBCHUNK_BLOCKS), tối đa COVERAGE_RETRY_ROUNDS vòng.
       Sub-chunk nhỏ tránh lặp lại lỗi "cả chunk lớn trả rỗng/hỏng" — đây là
       lưới an toàn chính chống dịch sót.
    3. Last resort: block vẫn chưa dịch được → giữ text gốc (preserve format).
    """
    result: dict = {}
    last_err = None
    tok_in_total = tok_out_total = 0

    # Phase 1 — dịch cả chunk, retry khi API lỗi
    for attempt in range(retries):
        try:
            t, in_t, out_t = translate_chunk(
                client, chunk, target_lang, doc_context,
                glossary=glossary, custom_rules=custom_rules,
                source_lang=source_lang,
            )
            result.update(t)
            tok_in_total  += in_t
            tok_out_total += out_t
            last_err = None
            break
        except Exception as e:
            last_err = str(e)
            if attempt < retries - 1:
                time.sleep(2 ** attempt)   # 1s, 2s, 4s

    # Phase 2 — dịch bù theo sub-chunk nhỏ cho block bị sót / còn tiếng nguồn
    if fallback_individually and COVERAGE_RETRY_ROUNDS > 0:
        rounds = 0
        pending = _pending_blocks(chunk, result, source_lang, target_lang)
        while pending and rounds < COVERAGE_RETRY_ROUNDS:
            rounds += 1
            for sub in _subchunks(pending, RETRY_SUBCHUNK_BLOCKS):
                try:
                    t, in_t, out_t = translate_chunk(
                        client, sub, target_lang, doc_context,
                        glossary=glossary, custom_rules=custom_rules,
                        source_lang=source_lang,
                    )
                    result.update(t)
                    tok_in_total  += in_t
                    tok_out_total += out_t
                except Exception as e:
                    last_err = str(e)
            pending = _pending_blocks(chunk, result, source_lang, target_lang)

    # Phase 3 — last resort: giữ text gốc cho block vẫn chưa dịch được
    n_missing = 0
    for b in chunk:
        if b["id"] not in result:
            result[b["id"]] = b.get("text_tagged") or b["text"]
            n_missing += 1
    if n_missing and not last_err:
        last_err = f"{n_missing} đoạn không dịch được (giữ nguyên gốc)"

    return result, tok_in_total, tok_out_total, last_err


# ══════════════════════════════════════════════════════════════════════════════
# PARALLEL TRANSLATION — entry point cho background thread
# ══════════════════════════════════════════════════════════════════════════════
def translate_parallel(holder: dict, client,
                       chunks: list[list[dict]],
                       target_lang: str, doc_context: str,
                       max_workers: int,
                       glossary: dict | None = None,
                       custom_rules: dict | None = None,
                       source_lang: str | None = None):
    """
    Background worker: dịch nhiều chunk song song với ThreadPoolExecutor.

    `holder` (dict shared với main thread) sẽ được update dưới lock:
    - translations  : dict {block_id -> translated_text}
    - tok_in/tok_out: cumulative tokens
    - chunk_done    : số chunk đã xong
    - chunk_log     : list per-chunk result (đẩy thêm khi mỗi chunk xong)
    - done          : True khi xong tất cả
    - error         : str nếu fatal error
    """
    lock = threading.Lock()
    try:
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {
                ex.submit(translate_chunk_with_retry,
                          client, c, target_lang, doc_context,
                          CHUNK_RETRIES, glossary, custom_rules,
                          True, source_lang): (i, c)
                for i, c in enumerate(chunks)
            }
            for future in as_completed(futures):
                idx, chunk = futures[future]
                try:
                    result, in_t, out_t, err = future.result()
                except Exception as e:
                    result, in_t, out_t, err = (
                        {b["id"]: b["text"] for b in chunk}, 0, 0, str(e)
                    )
                with lock:
                    holder["translations"].update(result)
                    holder["tok_in"]  += in_t
                    holder["tok_out"] += out_t
                    holder["chunk_log"].append({
                        "idx":   idx,
                        "size":  len(chunk),
                        "in_t":  in_t,
                        "out_t": out_t,
                        "error": err,
                    })
                    holder["chunk_done"] += 1
        holder["done"] = True
    except Exception as e:
        holder["error"] = str(e)
        holder["done"]  = True


# ══════════════════════════════════════════════════════════════════════════════
# RENDER
# ══════════════════════════════════════════════════════════════════════════════
def _resolve_translation(block: dict, translations: dict) -> str | None:
    """Resolve translation for a block, honoring _toc_mirror inheritance."""
    mirror_id = block.get("_toc_mirror")
    if mirror_id and translations.get(mirror_id):
        return translations[mirror_id]
    return translations.get(block["id"])


def apply_translations(original_bytes: bytes, blocks: list[dict],
                       translations: dict) -> bytes:
    """
    Áp dụng translations vào DOCX gốc → trả về bytes DOCX mới.
    Block nào có translation trong dict thì apply, không phân biệt role —
    để H/F translation cũng được apply khi user bấm nút "Dịch H/F".
    """
    doc         = Document(io.BytesIO(original_bytes))
    idx_to_para = {idx: para for idx, para in enumerate(iter_all_paragraphs(doc))}

    for block in blocks:
        if block.get("role") == "image_alt":
            continue
        tr   = _resolve_translation(block, translations)
        para = idx_to_para.get(block.get("para_idx"))
        if not tr or para is None:
            continue
        try:
            if block.get("has_format"):
                replace_paragraph_with_tagged(para, tr)
            else:
                replace_paragraph_text_keep_format(para, tr)
        except Exception:
            pass

    # Image alt-texts: re-iterate in same order as extract; match to blocks by sequence.
    alt_iter   = list(_iter_image_alt_texts(doc))
    alt_blocks = [b for b in blocks if b.get("role") == "image_alt"]
    for (el, attr, _), block in zip(alt_iter, alt_blocks):
        tr = _resolve_translation(block, translations)
        if tr:
            try:
                el.set(attr, tr)
            except Exception:
                pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _count_docx_media(docx_bytes: bytes) -> dict:
    """Đếm media trong DOCX → dict {media_files, drawing, pict, object}."""
    import zipfile
    counts = {"media_files": 0, "drawing": 0, "pict": 0, "object": 0}
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            names = zf.namelist()
            counts["media_files"] = sum(1 for n in names if n.startswith("word/media/"))
            # XML count: drawing/pict/object trong document.xml + header/footer
            from lxml import etree
            for n in names:
                if not n.endswith(".xml"):
                    continue
                if not (n.startswith("word/") and
                        ("document" in n or "header" in n or "footer" in n
                         or "footnote" in n or "endnote" in n)):
                    continue
                try:
                    root = etree.fromstring(zf.read(n))
                    for tag in ("drawing", "pict", "object"):
                        counts[tag] += sum(1 for _ in root.iter(qn(f"w:{tag}")))
                except Exception:
                    pass
    except Exception:
        pass
    return counts


def validate_docx_output(docx_bytes: bytes,
                         original_bytes: bytes | None = None) -> dict:
    """
    Validate translated DOCX output. Returns:
    {
      "valid": bool,
      "block_count": int,
      "image_count": int,
      "warnings": [str, ...],
      "errors": [str, ...],
    }

    Khi `original_bytes` được cấp → so sánh số media file, drawing, pict với
    bản gốc; nếu media bị mất → đánh `valid=False`.
    """
    import zipfile
    result = {"valid": True, "block_count": 0, "image_count": 0,
              "warnings": [], "errors": []}
    try:
        # 1. ZIP integrity
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            bad = zf.testzip()
            if bad:
                result["errors"].append(f"ZIP corrupt: {bad}")
                result["valid"] = False
                return result
            names = zf.namelist()
            if "word/document.xml" not in names:
                result["errors"].append("Missing word/document.xml")
                result["valid"] = False
                return result
            result["image_count"] = sum(1 for n in names if n.startswith("word/media/"))
        # 2. python-docx opens cleanly
        doc = Document(io.BytesIO(docx_bytes))
        result["block_count"] = sum(1 for _ in doc.element.iter(qn("w:p")))
        # 3. XML well-formed for all .xml parts
        from lxml import etree
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            for name in zf.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    try:
                        etree.fromstring(zf.read(name))
                    except etree.XMLSyntaxError as e:
                        result["errors"].append(f"{name}: {str(e)[:120]}")
                        result["valid"] = False
        # 4. Sanity checks
        if result["block_count"] == 0:
            result["warnings"].append("Document has 0 paragraphs (suspicious)")

        # 5. Media preservation (chỉ chạy khi có original)
        if original_bytes is not None:
            orig = _count_docx_media(original_bytes)
            out  = _count_docx_media(docx_bytes)
            if out["media_files"] < orig["media_files"]:
                result["errors"].append(
                    f"Media file count giảm: {orig['media_files']} → {out['media_files']}"
                )
                result["valid"] = False
            if out["drawing"] < orig["drawing"]:
                result["errors"].append(
                    f"<w:drawing> count giảm: {orig['drawing']} → {out['drawing']}"
                )
                result["valid"] = False
            if out["pict"] < orig["pict"]:
                result["errors"].append(
                    f"<w:pict> count giảm: {orig['pict']} → {out['pict']}"
                )
                result["valid"] = False
            if out["object"] < orig["object"]:
                result["warnings"].append(
                    f"<w:object> count giảm: {orig['object']} → {out['object']}"
                )
    except Exception as e:
        result["valid"] = False
        result["errors"].append(f"Validation crashed: {str(e)[:200]}")
    return result


_ASCII_WORD_RE = re.compile(r"\b[A-Za-z][A-Za-z']*\b")
_VI_DIACRITIC_RE = re.compile(
    r"[ăâêôơưđáàảãạắằẳẵặấầẩẫậéèẻẽẹếềểễệ"
    r"íìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ]",
    re.IGNORECASE,
)

_EN_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "been", "before", "by",
    "for", "from", "has", "have", "if", "in", "into", "is", "it", "must",
    "of", "on", "or", "shall", "should", "that", "the", "this", "to",
    "was", "were", "when", "where", "which", "with",
}
_EN_TECH_TERMS = {
    "assembly", "brake", "buffer", "button", "cabin", "cable", "call",
    "car", "chain", "comb", "commissioning", "control", "controller",
    "current", "device", "dimension", "door", "drawing", "drive",
    "elevator", "escalator", "floor", "frequency", "governor", "guide",
    "handrail", "hoistway", "inspection", "installation", "landing",
    "load", "machine", "maintenance", "motor", "operation", "panel",
    "pit", "power", "pulley", "rail", "requirement", "revision", "room",
    "rope", "safety", "shaft", "speed", "standard", "step", "supply",
    "switch", "test", "voltage", "warning",
}
_PROTECTED_EN_TOKENS = {
    "ac", "ag", "api", "asme", "dc", "din", "en", "iec", "iso", "jis",
    "led", "pdf", "qr", "tcvn", "ui", "url", "usb", "vnd",
}
_VI_COMMON_WORDS = {
    "cua", "va", "la", "cho", "duoc", "khong", "trong", "voi", "tu",
    "den", "khi", "neu", "thi", "phai", "can", "cac", "mot", "nhung",
    "chua", "kiem", "tra", "van", "hanh", "thang", "may",
}


def _norm_text(text: str) -> str:
    """Normalize for untranslated comparison."""
    text = strip_tags(str(text or ""))
    text = _TABLE_PREFIX_RE.sub("", text)
    return re.sub(r"\s+", " ", text).strip().lower()


def _strip_protected_english_spans(text: str) -> str:
    """Remove codes/standards that are allowed to stay in English."""
    text = re.sub(
        r"\b(?:EN|ISO|ASME|TCVN|IEC|DIN|JIS)\s*[-:/A-Z0-9.]*\d[-:/A-Z0-9.]*\b",
        " ",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\b[A-Z]{2,}[-_/]?\d[A-Z0-9._/-]*\b", " ", text)
    text = re.sub(r"\b\d+(?:[.,]\d+)?\s*(?:mm|cm|m|kg|kn|v|hz|kw|rpm|m/s|c)\b", " ", text, flags=re.IGNORECASE)
    return text


def _english_residue(text: str) -> bool:
    cleaned = _strip_protected_english_spans(strip_tags(str(text or "")))
    tokens_raw = _ASCII_WORD_RE.findall(cleaned)
    candidates: list[str] = []

    for raw in tokens_raw:
        token = raw.strip("'")
        lower = token.lower()
        if len(lower) <= 2:
            continue
        if lower in _PROTECTED_EN_TOKENS:
            continue
        if token.isupper() and len(token) <= 6:
            continue
        candidates.append(lower)

    if not candidates:
        return False

    stop_hits = sum(1 for t in candidates if t in _EN_STOPWORDS)
    tech_hits = sum(1 for t in candidates if t in _EN_TECH_TERMS)

    if stop_hits >= 2:
        return True
    if stop_hits >= 1 and len(candidates) >= 3:
        return True
    if tech_hits >= 2:
        return True

    return False


def _vietnamese_residue(text: str) -> bool:
    raw = str(text or "")
    if _VI_DIACRITIC_RE.search(raw):
        return True
    tokens = [t.lower() for t in _ASCII_WORD_RE.findall(raw)]
    return sum(1 for t in tokens if t in _VI_COMMON_WORDS) >= 2


def has_source_language_residue(text: str,
                                source_lang: str | None = None,
                                target_lang: str | None = None) -> bool:
    """Heuristic guard for translations that still contain source language."""
    source = (source_lang or "").lower()
    target = (target_lang or "").lower()
    if source == "english" and target == "vietnamese":
        return _english_residue(text)
    if source == "vietnamese" and target == "english":
        return _vietnamese_residue(text)
    return False


def _is_untranslated(b: dict, translations: dict,
                     source_lang: str | None = None,
                     target_lang: str | None = None,
                     output_block: dict | None = None) -> bool:
    tr = _resolve_translation(b, translations) or ""
    if not str(tr).strip():
        return True

    if _norm_text(tr) == _norm_text(b["text"]):
        return True

    if has_source_language_residue(tr, source_lang, target_lang):
        return True

    if output_block is not None:
        out_text = output_block.get("text", "")
        if _norm_text(out_text) == _norm_text(b["text"]):
            return True
        if has_source_language_residue(out_text, source_lang, target_lang):
            return True

    return False


# ══════════════════════════════════════════════════════════════════════════════
# CHECKPOINT — persist partial translations across browser refresh
# ══════════════════════════════════════════════════════════════════════════════
import os, pickle as _pickle, tempfile as _tempfile


def _checkpoint_path(docx_bytes: bytes, target_lang: str) -> str:
    # Hash TOÀN BỘ file để tránh collision khi 2 docx khác nhau cùng prefix 8KB.
    # Dùng `tempfile.gettempdir()` để portable: Linux=/tmp, macOS=/var/folders/...,
    # Windows=%TEMP%\... (P3.5 — trước đây hard-code /tmp gây mất checkpoint trên Win).
    h = hashlib.md5(docx_bytes).hexdigest()[:16]
    slug = target_lang.replace(" ", "_")[:12]
    return os.path.join(_tempfile.gettempdir(), f"tr_ckpt_{h}_{slug}.pkl")


def checkpoint_save(docx_bytes: bytes, target_lang: str, translations: dict) -> None:
    try:
        with open(_checkpoint_path(docx_bytes, target_lang), "wb") as f:
            _pickle.dump(translations, f)
    except Exception:
        pass


def checkpoint_load(docx_bytes: bytes, target_lang: str) -> dict | None:
    path = _checkpoint_path(docx_bytes, target_lang)
    try:
        if os.path.exists(path):
            with open(path, "rb") as f:
                return _pickle.load(f)
    except Exception:
        pass
    return None


def checkpoint_clear(docx_bytes: bytes, target_lang: str) -> None:
    try:
        os.unlink(_checkpoint_path(docx_bytes, target_lang))
    except Exception:
        pass


def find_missed(blocks: list[dict], translations: dict,
                source_lang: str | None = None,
                target_lang: str | None = None,
                output_blocks: list[dict] | None = None) -> list[dict]:
    """Return translatable blocks that still look untranslated."""
    output_by_id = {b["id"]: b for b in (output_blocks or [])}
    return [
        b for b in blocks
        if b["role"] not in NO_TRANSLATE_ROLES
        and _is_untranslated(
            b, translations,
            source_lang=source_lang,
            target_lang=target_lang,
            output_block=output_by_id.get(b["id"]),
        )
    ]


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE OCR + TRANSLATE (Gemini Vision)
# ══════════════════════════════════════════════════════════════════════════════
# Image occurrence model (OCR_DEV P0.1):
#   {
#     "id":               "OCC_0", ..., unique per occurrence
#     "filename":         "image1.png" (basename in word/media/)
#     "content_type":     "image/png"
#     "data":             bytes
#     "doc_part":         "word/document.xml" | "word/header1.xml" | ...
#     "rels_path":        "word/_rels/document.xml.rels" | ...
#     "rId":              "rId7" (relationship id pointing to media)
#     "paragraph_index":  position of <w:p> ancestor in part (0-based)
#     "occurrence_index": 0-based among same (doc_part, rId) usages
#     "width_px":         int | None  (decoded from image header if Pillow available)
#     "height_px":        int | None
#   }
#
# Image dùng nhiều nơi (same rId trong cùng part hoặc cùng filename khác part)
# vẫn cho ra nhiều occurrence — user chọn occurrence nào, replace chỉ occurrence đó.

_R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_NS = "http://www.w3.org/XML/1998/namespace"

# Doc parts có thể chứa ảnh
_IMAGE_PARTS = (
    "word/document.xml",
    # header/footer/footnotes/endnotes/comments được scan động bên dưới
)

_IMG_EXTS = ("png", "jpg", "jpeg", "gif", "bmp", "webp")


def _img_dimensions(data: bytes) -> tuple[int | None, int | None]:
    """Return (width, height) in pixels. None khi không decode được."""
    try:
        from PIL import Image
        with Image.open(io.BytesIO(data)) as im:
            return im.size  # (w, h)
    except Exception:
        return None, None


def _list_doc_parts(docx_bytes: bytes) -> list[tuple[str, str]]:
    """Liệt kê các DOCX XML part có thể chứa ảnh + .rels tương ứng.

    Trả [(part_path, rels_path), ...].
    """
    import zipfile
    parts = []
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        names = set(zf.namelist())
        if "word/document.xml" in names:
            parts.append(("word/document.xml", "word/_rels/document.xml.rels"))
        for n in sorted(names):
            if not n.startswith("word/"):
                continue
            base = n.rsplit("/", 1)[-1]
            if base.endswith(".xml") and (base.startswith("header") or base.startswith("footer")
                                          or base in ("footnotes.xml", "endnotes.xml", "comments.xml")):
                rels = f"word/_rels/{base}.rels"
                if rels in names:
                    parts.append((n, rels))
    return parts


def _read_media(docx_bytes: bytes) -> dict[str, tuple[str, bytes]]:
    """Đọc tất cả file `word/media/*` → {basename: (content_type, bytes)}."""
    import zipfile
    media: dict[str, tuple[str, bytes]] = {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        for name in zf.namelist():
            if not name.startswith("word/media/") or "." not in name:
                continue
            ext = name.rsplit(".", 1)[-1].lower()
            if ext not in _IMG_EXTS:
                continue
            mime_ext = "jpeg" if ext == "jpg" else ext
            media[name.rsplit("/", 1)[-1]] = (f"image/{mime_ext}", zf.read(name))
    return media


def extract_image_occurrences(docx_bytes: bytes) -> list[dict]:
    """B2: liệt kê mọi lần ảnh xuất hiện theo vị trí paragraph.

    Trả list ImageOccurrence (shape mô tả ở đầu module). Mỗi `r:embed` trong 1
    `<w:p>` (qua w:drawing hoặc w:pict) là một occurrence.
    """
    import zipfile
    from lxml import etree

    media = _read_media(docx_bytes)
    if not media:
        return []

    occurrences: list[dict] = []
    rid_used_counter: dict[tuple[str, str], int] = {}

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        for part_path, rels_path in _list_doc_parts(docx_bytes):
            try:
                part_xml = zf.read(part_path)
                rels_xml = zf.read(rels_path)
            except KeyError:
                continue

            # rId → media filename (chỉ relationship trỏ tới media/)
            rid_to_fname: dict[str, str] = {}
            for rel in etree.fromstring(rels_xml):
                target = rel.get("Target", "")
                rid    = rel.get("Id", "")
                if "media/" in target:
                    rid_to_fname[rid] = target.split("/")[-1]

            if not rid_to_fname:
                continue

            root = etree.fromstring(part_xml)
            para_qn = f"{{{_W_NS}}}p"
            embed_attr = f"{{{_R_NS}}}embed"

            for para_idx, para in enumerate(root.iter(para_qn)):
                for elem in para.iter():
                    rid = elem.get(embed_attr)
                    if not rid or rid not in rid_to_fname:
                        continue
                    fname = rid_to_fname[rid]
                    if fname not in media:
                        continue
                    content_type, data = media[fname]
                    occ_key = (part_path, rid)
                    occ_idx = rid_used_counter.get(occ_key, 0)
                    rid_used_counter[occ_key] = occ_idx + 1
                    w, h = _img_dimensions(data)
                    occurrences.append({
                        "id":               f"OCC_{len(occurrences)}",
                        "filename":         fname,
                        "content_type":     content_type,
                        "data":             data,
                        "doc_part":         part_path,
                        "rels_path":        rels_path,
                        "rId":              rid,
                        "paragraph_index":  para_idx,
                        "occurrence_index": occ_idx,
                        "width_px":         w,
                        "height_px":        h,
                    })

    return occurrences


# ── Cost helpers ───────────────────────────────────────────────────────────
# Gemini Vision token model (snapshot 2026-05-26 — xem OCR_DEV.md):
#   ≤ 384px mỗi chiều: 258 tokens
#   > 384px: chia tile 768×768, mỗi tile 258 tokens
_TILE_PX     = 768
_SMALL_PX    = 384
_TOK_PER_IMG = 258


def estimate_image_input_tokens(width_px: int | None, height_px: int | None) -> int:
    """Estimate input tokens cho một ảnh theo Gemini Vision model."""
    if not width_px or not height_px:
        # Không biết kích thước → giả định 1 tile
        return _TOK_PER_IMG
    if width_px <= _SMALL_PX and height_px <= _SMALL_PX:
        return _TOK_PER_IMG
    tiles_w = (width_px + _TILE_PX - 1) // _TILE_PX
    tiles_h = (height_px + _TILE_PX - 1) // _TILE_PX
    return max(1, tiles_w * tiles_h) * _TOK_PER_IMG


def estimate_ocr_cost(occurrences: list[dict],
                      skip_under_bytes: int = 5_000,
                      avg_output_tokens: int = 300,
                      prompt_overhead_tokens: int = 250) -> dict:
    """P1.1: preflight estimate. Trả dict:
      {
        "n_total": N, "n_skipped": N, "n_to_ocr": N,
        "input_tokens": N, "output_tokens": N,
        "usd": float, "vnd": float,
      }
    """
    from ui_common import calc_cost

    n_total = len(occurrences)
    to_ocr = [o for o in occurrences if len(o.get("data", b"")) >= skip_under_bytes]
    n_skipped = n_total - len(to_ocr)

    img_tokens = sum(
        estimate_image_input_tokens(o.get("width_px"), o.get("height_px"))
        for o in to_ocr
    )
    in_tokens  = img_tokens + prompt_overhead_tokens * len(to_ocr)
    out_tokens = avg_output_tokens * len(to_ocr)
    usd, vnd = calc_cost(in_tokens, out_tokens)
    return {
        "n_total":       n_total,
        "n_skipped":     n_skipped,
        "n_to_ocr":      len(to_ocr),
        "input_tokens":  in_tokens,
        "output_tokens": out_tokens,
        "usd":           usd,
        "vnd":           vnd,
    }


# ── OCR call ───────────────────────────────────────────────────────────────
_OCR_RESPONSE_KEYS = {"has_text", "ocr", "translation", "regions", "confidence"}


def _build_ocr_prompt(source_lang: str | None, target_lang: str,
                      glossary: dict | None,
                      subdomains: set[str] | None) -> str:
    """B3: prompt OCR có direction + domain technical style + regions[bbox]."""
    src_clause = f"from {source_lang} " if source_lang else ""
    glossary_section = ""
    if glossary:
        items = list(glossary.items())[:40]
        glossary_str = "\n".join(f"  {k} → {v}" for k, v in items)
        glossary_section = (
            "\nGlossary (USE THESE EXACT translations for consistency):\n"
            f"{glossary_str}\n"
        )
    domain_section = ""
    if subdomains and (subdomains & {"elevator", "escalator"}):
        domain_section = (
            "\nDomain: elevator/escalator engineering. Use formal technical register. "
            "Preserve units (mm, m, m/s, kg, kN, V, Hz, °C, dB), standards "
            "(EN 81-20, EN 81-50, ISO 22201, ASME A17.1, GB 7588, TCVN 6395), "
            "part/drawing/revision codes (e.g. 'Rev. A', 'DWG-1234').\n"
        )

    return (
        f"You are an expert OCR + translator. Analyze this image carefully.\n\n"
        f"TASK:\n"
        f"1. Detect if image contains readable text (labels, captions, table cells, "
        f"diagram annotations, equations).\n"
        f"2. If text exists:\n"
        f"   a. Extract ALL text verbatim — use \\n for line breaks, '|' to separate "
        f"table columns.\n"
        f"   b. Keep numbers, units, punctuation, special characters EXACTLY.\n"
        f"   c. Translate {src_clause}to {target_lang}, preserving structure and order.\n"
        f"   d. For each text region (label, paragraph, cell), also return bounding box "
        f"in normalized coords (x, y, w, h are floats in [0, 1] relative to image size, "
        f"x/y is top-left corner).\n"
        f"   e. confidence: float [0, 1] — your confidence that bboxes are accurate.\n"
        f"{domain_section}{glossary_section}"
        f"\nRespond ONLY as JSON (no extra keys, no markdown):\n"
        f"  if text found : {{\"has_text\": true, \"ocr\": \"...\", \"translation\": \"...\", "
        f"\"regions\": [{{\"bbox\": [x, y, w, h], \"ocr\": \"...\", \"translation\": \"...\"}}], "
        f"\"confidence\": 0.0..1.0}}\n"
        f"  if no text    : {{\"has_text\": false, \"ocr\": \"\", \"translation\": \"\", "
        f"\"regions\": [], \"confidence\": 1.0}}"
    )


def _ocr_single_image(client, occ: dict, source_lang: str | None,
                      target_lang: str, model: str,
                      glossary: dict | None = None,
                      subdomains: set[str] | None = None,
                      retries: int = CHUNK_RETRIES) -> dict:
    """B3+B4: gọi Gemini Vision, retry, trả về result + cost.

    Result shape (mọi field luôn có):
      {
        "ocr":         str,
        "translation": str,
        "has_text":    bool,
        "regions":     [{"bbox": [x,y,w,h], "ocr": str, "translation": str}, ...],
        "confidence":  float,
        "tok_in":      int, "tok_out": int, "total_tokens": int,
        "usd":         float, "vnd":    float,
        "model":       str,
        "attempts":    int,
        "error":       str | None,
      }
    """
    from google.genai import types as gtypes
    from ui_common import calc_cost

    prompt = _build_ocr_prompt(source_lang, target_lang, glossary, subdomains)
    last_err = None
    attempts = 0
    for attempt in range(retries):
        attempts = attempt + 1
        try:
            response = client.models.generate_content(
                model=model,
                contents=[
                    gtypes.Part.from_bytes(data=occ["data"], mime_type=occ["content_type"]),
                    prompt,
                ],
                config=gtypes.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0.1,
                ),
            )
            data = parse_json_loose(getattr(response, "text", "")) or {}
            if not isinstance(data, dict):
                data = {}
            tok_in, tok_out = usage_tokens(response)
            usd, vnd = calc_cost(tok_in, tok_out)
            # Filter regions: only well-formed bbox
            regions_raw = data.get("regions") or []
            regions: list[dict] = []
            if isinstance(regions_raw, list):
                for r in regions_raw:
                    if not isinstance(r, dict):
                        continue
                    bbox = r.get("bbox")
                    if not (isinstance(bbox, list) and len(bbox) == 4
                            and all(isinstance(x, (int, float)) for x in bbox)):
                        continue
                    regions.append({
                        "bbox":        [float(x) for x in bbox],
                        "ocr":         str(r.get("ocr") or "").strip(),
                        "translation": str(r.get("translation") or "").strip(),
                    })
            return {
                "ocr":          str(data.get("ocr", "") or "").strip(),
                "translation":  str(data.get("translation", "") or "").strip(),
                "has_text":     bool(data.get("has_text", False)),
                "regions":      regions,
                "confidence":   float(data.get("confidence", 0.0) or 0.0),
                "tok_in":       tok_in,
                "tok_out":      tok_out,
                "total_tokens": tok_in + tok_out,
                "usd":          usd,
                "vnd":          vnd,
                "model":        model,
                "attempts":     attempts,
                "error":        None,
            }
        except Exception as e:
            last_err = str(e)
            if attempt < retries - 1:
                time.sleep(2 ** attempt)

    return {
        "ocr": "", "translation": "", "has_text": False,
        "regions": [], "confidence": 0.0,
        "tok_in": 0, "tok_out": 0, "total_tokens": 0,
        "usd": 0.0, "vnd": 0.0, "model": model,
        "attempts": attempts, "error": last_err,
    }


def ocr_and_translate_images(
    client,
    occurrences: list[dict],
    target_lang: str,
    source_lang: str | None = None,
    glossary: dict | None = None,
    subdomains: set[str] | None = None,
    progress_callback=None,
    max_workers: int = MAX_WORD_WORKERS,
    skip_under_bytes: int = 5_000,
) -> dict:
    """OCR + translate parallel. Trả dict {occ_id: result_dict_per_image} + key
    "_total" chứa aggregate cost.
    """
    from ui_common import calc_cost
    results: dict = {}
    model = get_working_model() or WORD_MODELS[0]

    to_process = []
    for occ in occurrences:
        if len(occ["data"]) < skip_under_bytes:
            results[occ["id"]] = {
                "ocr": "", "translation": "", "has_text": False,
                "regions": [], "confidence": 0.0,
                "tok_in": 0, "tok_out": 0, "total_tokens": 0,
                "usd": 0.0, "vnd": 0.0, "model": model,
                "attempts": 0, "error": None, "skipped": True,
            }
        else:
            to_process.append(occ)

    total_for_progress = len(occurrences)

    if to_process:
        done_count = 0
        lock = threading.Lock()
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {
                ex.submit(_ocr_single_image, client, occ, source_lang,
                          target_lang, model, glossary, subdomains): occ
                for occ in to_process
            }
            for future in as_completed(futures):
                occ = futures[future]
                try:
                    results[occ["id"]] = future.result()
                except Exception as e:
                    results[occ["id"]] = {
                        "ocr": "", "translation": "", "has_text": False,
                        "regions": [], "confidence": 0.0,
                        "tok_in": 0, "tok_out": 0, "total_tokens": 0,
                        "usd": 0.0, "vnd": 0.0, "model": model,
                        "attempts": 0, "error": str(e),
                    }
                with lock:
                    done_count += 1
                    current = done_count
                if progress_callback:
                    progress_callback(current, len(to_process))

    # Aggregate
    tok_in  = sum(r.get("tok_in", 0)  for r in results.values() if isinstance(r, dict))
    tok_out = sum(r.get("tok_out", 0) for r in results.values() if isinstance(r, dict))
    usd, vnd = calc_cost(tok_in, tok_out)
    results["_total"] = {
        "tok_in": tok_in, "tok_out": tok_out,
        "usd": usd, "vnd": vnd, "model": model,
        "n_called": len(to_process),
        "n_skipped": len(occurrences) - len(to_process),
        "n_total":   len(occurrences),
    }
    return results


# ══════════════════════════════════════════════════════════════════════════════
# CAPTION INSERTION (selectable) — OCR mode "đưa text dưới ảnh"
# ══════════════════════════════════════════════════════════════════════════════
def _find_occurrence_paragraph(part_root, rid: str, occ_idx: int):
    """Trả về phần tử <w:p> chứa occurrence thứ `occ_idx` của `rid` trong part root."""
    embed_attr = f"{{{_R_NS}}}embed"
    para_qn    = f"{{{_W_NS}}}p"
    seen = 0
    for para in part_root.iter(para_qn):
        # Check first <w:p> ancestor — find any descendant with this rId
        for elem in para.iter():
            if elem.get(embed_attr) == rid:
                if seen == occ_idx:
                    return para
                seen += 1
                break   # mỗi <w:p> đếm 1 occurrence của rId này
    return None


def insert_ocr_captions_into_docx(
    docx_bytes: bytes,
    occurrences: list[dict],
    ocr_results: dict,
    selected_ids: list[str] | set[str] | None = None,
    edited_translations: dict[str, str] | None = None,
    remove_original_ids: list[str] | set[str] | None = None,
) -> bytes:
    """B5: chèn caption (text dịch) DƯỚI mỗi occurrence ảnh được chọn.

    - `selected_ids`: chỉ chèn cho occurrence id có trong set. None = tất cả ảnh có text.
    - `edited_translations`: {occ_id: text} — override bản dịch trước khi chèn.
    - `remove_original_ids`: occurrence sẽ bị xoá ảnh gốc, chỉ giữ caption.

    Không chèn duplicate: nếu paragraph TIẾP SAU ảnh đã có dạng `[OCR] ...` thì update
    thay vì insert mới.
    """
    import zipfile
    from lxml import etree

    if selected_ids is not None:
        selected_set = set(selected_ids)
    else:
        selected_set = {occ["id"] for occ in occurrences
                        if ocr_results.get(occ["id"], {}).get("has_text")
                        and ocr_results.get(occ["id"], {}).get("translation")}
    edited_translations = edited_translations or {}
    remove_set = set(remove_original_ids or ())

    if not selected_set and not remove_set:
        return docx_bytes

    # Group occurrences by doc_part
    by_part: dict[str, list[dict]] = {}
    for occ in occurrences:
        if occ["id"] not in selected_set and occ["id"] not in remove_set:
            continue
        by_part.setdefault(occ["doc_part"], []).append(occ)

    # Read & modify each part XML, then write back
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        parts_data: dict[str, bytes] = {n: zf.read(n) for n in zf.namelist()}

    for part_path, occs in by_part.items():
        if part_path not in parts_data:
            continue
        root = etree.fromstring(parts_data[part_path])

        # Reverse-order so paragraph-index doesn't shift earlier occurrences.
        # Sort by paragraph_index desc, then occurrence_index desc.
        occs_sorted = sorted(occs,
                             key=lambda o: (o["paragraph_index"], o["occurrence_index"]),
                             reverse=True)
        for occ in occs_sorted:
            para = _find_occurrence_paragraph(root, occ["rId"], occ["occurrence_index"])
            if para is None:
                continue

            # 1) Insert/replace caption (nếu occ trong selected_set)
            if occ["id"] in selected_set:
                tr = (edited_translations.get(occ["id"])
                      or ocr_results.get(occ["id"], {}).get("translation", ""))
                tr = (tr or "").strip()
                if tr:
                    caption_xml = _make_caption_paragraph(tr)
                    parent = para.getparent()
                    idx = list(parent).index(para)
                    # Nếu paragraph kế tiếp là caption cũ (text bắt đầu "[OCR] ") → replace
                    next_para = parent[idx + 1] if idx + 1 < len(parent) else None
                    if next_para is not None and _is_ocr_caption(next_para):
                        parent.remove(next_para)
                    parent.insert(idx + 1, caption_xml)

            # 2) Remove ảnh gốc (nếu occ trong remove_set) — xoá w:drawing/w:pict
            #    chứa rId của occ. Nếu paragraph chỉ có ảnh → xoá luôn paragraph.
            if occ["id"] in remove_set:
                _strip_drawing_with_rid(para, occ["rId"])
                # Nếu paragraph rỗng (no text + no remaining drawing) → xoá
                has_text = any((t.text or "").strip()
                               for t in para.iter(f"{{{_W_NS}}}t"))
                has_media = (para.find(f".//{{{_W_NS}}}drawing") is not None
                             or para.find(f".//{{{_W_NS}}}pict") is not None)
                if not has_text and not has_media:
                    parent = para.getparent()
                    if parent is not None:
                        parent.remove(para)

        parts_data[part_path] = etree.tostring(root, xml_declaration=True,
                                               encoding="UTF-8", standalone=True)

    # Write back DOCX
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts_data.items():
            zout.writestr(name, data)
    return buf.getvalue()


def _make_caption_paragraph(text: str):
    """Trả <w:p> styled caption: căn giữa, italic, size 9, prefix `[OCR] `."""
    from lxml import etree
    p   = etree.Element(f"{{{_W_NS}}}p")
    pPr = etree.SubElement(p, f"{{{_W_NS}}}pPr")
    jc  = etree.SubElement(pPr, f"{{{_W_NS}}}jc")
    jc.set(f"{{{_W_NS}}}val", "center")

    run = etree.SubElement(p, f"{{{_W_NS}}}r")
    rPr = etree.SubElement(run, f"{{{_W_NS}}}rPr")
    etree.SubElement(rPr, f"{{{_W_NS}}}i")
    etree.SubElement(rPr, f"{{{_W_NS}}}iCs")
    color = etree.SubElement(rPr, f"{{{_W_NS}}}color")
    color.set(f"{{{_W_NS}}}val", "808080")
    sz = etree.SubElement(rPr, f"{{{_W_NS}}}sz")
    sz.set(f"{{{_W_NS}}}val", "18")
    szCs = etree.SubElement(rPr, f"{{{_W_NS}}}szCs")
    szCs.set(f"{{{_W_NS}}}val", "18")

    t = etree.SubElement(run, f"{{{_W_NS}}}t")
    t.text = f"[OCR] {text}"
    t.set(f"{{{_XML_NS}}}space", "preserve")
    return p


def _is_ocr_caption(para_elem) -> bool:
    """True nếu paragraph là caption [OCR] do chính app sinh ra (dedupe insertion)."""
    for t in para_elem.iter(f"{{{_W_NS}}}t"):
        if (t.text or "").startswith("[OCR] "):
            return True
    return False


def _strip_drawing_with_rid(para_elem, rid: str) -> int:
    """Xoá mọi w:drawing / w:pict / w:object trong `para_elem` mà tham chiếu `rid`.
    Trả số element bị xoá.
    """
    embed_attr = f"{{{_R_NS}}}embed"
    removed = 0
    for tag in ("drawing", "pict", "object"):
        full = f"{{{_W_NS}}}{tag}"
        for elem in list(para_elem.iter(full)):
            # check if any descendant has this rId
            hit = any(d.get(embed_attr) == rid for d in elem.iter())
            if hit:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
                    removed += 1
    return removed

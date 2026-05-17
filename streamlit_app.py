"""
Translator App — PDF + Word [Streamlit]
• Tab PDF  : Dịch PDF giữ nguyên layout (Gemini Flash)
• Tab Word : Dịch DOCX giữ nguyên định dạng (Gemini fallback chain)
• Password protection, live timer, retry với exponential backoff
"""

import os, io, json, time, tempfile, threading
import streamlit as st
import fitz
from docx import Document
from google import genai
from google.genai import types as gtypes
from datetime import datetime

# ══════════════════════════════════════════════════════════════════════════════
# CONFIG
# ══════════════════════════════════════════════════════════════════════════════
API_KEY      = st.secrets.get("GEMINI_API_KEY", "")
APP_PASSWORD = st.secrets["APP_PASSWORD"]

# PDF
PDF_MODEL    = "gemini-2.5-flash"
PRICE_INPUT  = 0.10
PRICE_OUTPUT = 0.40
USD_TO_VND   = 25400
PDF_DELAY    = 0.3
MAX_RETRIES  = 5
RETRY_CODES  = ("429", "resource_exhausted", "quota", "rate")
UNICODE_FONTS = [
    "Carlito-Regular.ttf",
    "/usr/share/fonts/truetype/crosextra/Carlito-Regular.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "C:/Windows/Fonts/Arial.ttf",
    "C:/Windows/Fonts/calibri.ttf",
]

# Word
WORD_MODELS = [
    "gemini-2.5-flash-lite",
    "gemini-2.5-flash",
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
]
CHUNK_SIZE         = 25
MAX_WORD_TOKENS    = 65536
NO_TRANSLATE_ROLES = {"header", "footer"}

LANGUAGES = ["Tiếng Việt", "Tiếng Anh", "Tiếng Nhật", "Tiếng Trung", "Tiếng Pháp", "Tiếng Đức"]
LANG_EN = {
    "Tiếng Việt": "Vietnamese",
    "Tiếng Anh":  "English",
    "Tiếng Nhật": "Japanese",
    "Tiếng Trung": "Chinese",
    "Tiếng Pháp": "French",
    "Tiếng Đức":  "German",
}

_word_working_model: list = [None]

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG & CSS
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(page_title="Translator — Vi Nguyen", page_icon="⬡", layout="centered")

st.markdown("""
<style>
    .stApp { background-color: #0f1117; color: #e2e8f0; }
    .block-container { max-width: 820px; padding-top: 2rem; }
    h1, h2, h3 { color: #e2e8f0 !important; }
    p, li, span { color: #e2e8f0; }

    div[data-testid="stFileUploader"] {
        border: 1.5px dashed #4a5080 !important; border-radius: 10px !important;
        padding: 8px !important; background: #1a1d27 !important;
    }
    div[data-testid="stFileUploader"] label { color: #e2e8f0 !important; }
    [data-testid="stFileUploaderDropzone"] { background-color: #242838 !important; border-color: #4a5080 !important; }
    [data-testid="stFileUploaderDropzone"] > div { background-color: #242838 !important; }
    [data-testid="stFileUploaderDropzone"] span { color: #e2e8f0 !important; }
    [data-testid="stFileUploaderDropzone"] button { background-color: #2d3149 !important; color: #e2e8f0 !important; border-color: #4a5080 !important; }
    [data-testid="stFileUploaderDropzone"] p { color: #94a3b8 !important; }

    label, .stSelectbox label, .stTextInput label,
    div[data-testid="stWidgetLabel"] p { color: #e2e8f0 !important; font-weight: 500; }
    .stSelectbox div[data-baseweb="select"] > div { background: #1a1d27 !important; color: #e2e8f0 !important; border-color: #4a5080 !important; }
    .stTextInput input { background: #1a1d27 !important; color: #e2e8f0 !important; border-color: #4a5080 !important; }
    input::placeholder { color: #64748b !important; }

    .stat-box { background: #1a1d27; border-radius: 10px; padding: 14px 18px; text-align: center; border: 1px solid #4a5080; }
    .stat-val { font-size: 1.4rem; font-weight: bold; color: #e2e8f0; }
    .stat-lbl { font-size: 0.78rem; color: #94a3b8; margin-top: 4px; }

    .timer-box { background: #1a1d27; border-radius: 10px; padding: 12px 18px; text-align: center; border: 1px solid #6c63ff; margin-bottom: 10px; }
    .timer-val { font-size: 2rem; font-weight: bold; color: #a78bfa; font-family: 'Courier New', monospace; }
    .timer-lbl { font-size: 0.78rem; color: #94a3b8; margin-top: 2px; }
    .timer-status { font-size: 0.85rem; color: #818cf8; margin-top: 6px; }

    .log-box { background: #1a1d27; border-radius: 8px; padding: 14px 18px; font-family: 'Courier New', monospace; font-size: 0.83rem; max-height: 340px; overflow-y: auto; border: 1px solid #4a5080; white-space: pre-wrap; color: #c4cde0; line-height: 1.6; }

    .stButton > button { background-color: #6c63ff !important; color: white !important; border: none !important; border-radius: 8px !important; padding: 10px 28px !important; font-weight: bold !important; font-size: 1rem !important; width: 100% !important; }
    .stButton > button:hover { background-color: #a78bfa !important; }
    .stButton > button:disabled { background-color: #2d3149 !important; color: #64748b !important; }

    .stDownloadButton > button { background-color: #059669 !important; color: white !important; border-radius: 8px !important; font-weight: bold !important; font-size: 1rem !important; width: 100% !important; }
    .stDownloadButton > button:hover { background-color: #10b981 !important; }

    div[data-testid="stProgressBar"] > div { background-color: #1a1d27 !important; }
    div[data-testid="stProgressBar"] > div > div { background-color: #6c63ff !important; }

    hr { border-color: #2d3149 !important; }
    div[data-testid="stAlert"] { background: #1a1d27 !important; border-color: #4a5080 !important; color: #e2e8f0 !important; }

    .stTabs [data-baseweb="tab-list"] { background: #1a1d27; border-radius: 10px; padding: 4px; gap: 4px; }
    .stTabs [data-baseweb="tab"] { color: #94a3b8 !important; background: transparent; border-radius: 8px; font-weight: 600; font-size: 1rem; padding: 8px 24px; }
    .stTabs [aria-selected="true"] { color: #e2e8f0 !important; background: #6c63ff !important; }

    .login-box { background: #1a1d27; border: 1px solid #4a5080; border-radius: 14px; padding: 2.5rem 2rem; max-width: 380px; margin: 4rem auto 0 auto; text-align: center; }
    .login-title { font-size: 1.6rem; font-weight: bold; color: #e2e8f0; margin-bottom: 0.3rem; }
    .login-sub   { font-size: 0.85rem; color: #64748b; margin-bottom: 1.8rem; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PASSWORD GATE
# ══════════════════════════════════════════════════════════════════════════════
def check_password():
    if st.session_state.get("authenticated"):
        return True
    st.markdown("""
    <div class='login-box'>
        <div class='login-title'>⬡ Translator</div>
        <div class='login-sub'>PDF & Word — Powered by Gemini — Vi Nguyen</div>
    </div>""", unsafe_allow_html=True)
    col_l, col_m, col_r = st.columns([1, 2, 1])
    with col_m:
        st.markdown("<div style='height:1.2rem'></div>", unsafe_allow_html=True)
        pwd = st.text_input("🔑 Nhập mật khẩu", type="password", key="pwd_input", placeholder="Password...")
        if st.button("Đăng nhập", use_container_width=True):
            if pwd == APP_PASSWORD:
                st.session_state["authenticated"] = True
                st.rerun()
            else:
                st.error("❌ Sai mật khẩu, thử lại!")
    return False

if not check_password():
    st.stop()


# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("## ⬡ Translator")
st.markdown("<span style='color:#64748b;font-size:0.9rem'>PDF & Word — Powered by Gemini — Vi Nguyen</span>",
            unsafe_allow_html=True)
_, col_lo = st.columns([6, 1])
with col_lo:
    if st.button("🚪 Logout"):
        st.session_state.pop("authenticated", None)
        st.rerun()
st.divider()


# ══════════════════════════════════════════════════════════════════════════════
# SHARED GEMINI CLIENT
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_resource
def get_client():
    return genai.Client(api_key=API_KEY)

client = get_client()


# ══════════════════════════════════════════════════════════════════════════════
# PDF BACKEND
# ══════════════════════════════════════════════════════════════════════════════
def find_font():
    for p in UNICODE_FONTS:
        if os.path.isfile(p):
            return p
    return None

def int_to_rgb(c):
    return ((c >> 16) & 0xFF) / 255, ((c >> 8) & 0xFF) / 255, (c & 0xFF) / 255

def extract_line_groups(pdf_path, page_nums=None):
    result = {}
    doc    = fitz.open(pdf_path)
    total  = len(doc)
    targets = page_nums if page_nums else list(range(total))
    for pi in targets:
        if not (0 <= pi < total):
            continue
        page = doc[pi]
        data = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_MEDIABOX_CLIP)
        groups = []
        for block in data["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                spans = [sp for sp in line["spans"] if sp["text"].strip()]
                if not spans:
                    continue
                merged = "".join(sp["text"] for sp in spans).strip()
                if not merged:
                    continue
                groups.append({
                    "bbox": (
                        min(sp["bbox"][0] for sp in spans),
                        min(sp["bbox"][1] for sp in spans),
                        max(sp["bbox"][2] for sp in spans),
                        max(sp["bbox"][3] for sp in spans),
                    ),
                    "text": merged,
                    "size": max(sp["size"] for sp in spans),
                    "rgb":  int_to_rgb(spans[0]["color"]),
                    "bold": any(bool(sp["flags"] & (1 << 4)) for sp in spans),
                })
        result[pi] = groups
    doc.close()
    return result, total

def _get_bold_font_path(font_path):
    if not font_path:
        return None
    for regular, bold in [
        ("Carlito-Regular.ttf", "Carlito-Bold.ttf"),
        ("Arial.ttf", "Arialbd.ttf"),
        ("calibri.ttf", "calibrib.ttf"),
        ("times.ttf", "timesbd.ttf"),
        ("DejaVuSans.ttf", "DejaVuSans-Bold.ttf"),
    ]:
        candidate = font_path.replace(regular, bold)
        if os.path.isfile(candidate):
            return candidate
    return None

def _insert_line(page, font_path, font_bold_path, bbox, text, fontsize, color, bold):
    x0, y0, x1, y1 = bbox
    pw      = page.rect.width
    ph      = page.rect.height
    x1_use  = pw - 30
    line_h  = max(y1 - y0, fontsize * 1.5)
    y1_use  = min(y0 + line_h * 3, ph - 10)
    if font_path:
        fontfile = font_bold_path if bold and font_bold_path else font_path
        fontname = "FBold" if bold else "FReg"
        try:
            page.insert_font(fontname=fontname, fontfile=fontfile)
        except Exception:
            pass
    else:
        fontname = "hebo" if bold else "helv"
    size = max(fontsize, 5.0)
    while size >= 4.0:
        rc = page.insert_textbox(
            fitz.Rect(x0, y0, x1_use, y1_use),
            text, fontsize=size, fontname=fontname, color=color, align=0,
        )
        if rc >= 0:
            break
        size -= 0.5

def write_translated_pdf(src, dst, all_groups, all_trans, font_path):
    doc = fitz.open(src)
    font_bold_path = _get_bold_font_path(font_path)
    for pi, groups in all_groups.items():
        trans = all_trans.get(pi, [])
        if not groups or not trans:
            continue
        page = doc[pi]
        for g in groups:
            r        = fitz.Rect(g["bbox"])
            expanded = fitz.Rect(r.x0 - 10, r.y0 - 5, r.x1 + 10, r.y1 + 5)
            page.add_redact_annot(expanded.intersect(page.rect))
        page.apply_redactions(images=0)
        for g, t in zip(groups, trans):
            text = t.strip() if t and t.strip() else g["text"]
            try:
                _insert_line(page, font_path, font_bold_path,
                             g["bbox"], text, g["size"], g["rgb"], g["bold"])
            except Exception:
                pass
    doc.save(dst, garbage=4, deflate=True)
    doc.close()

def _parse_json_pdf(raw):
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()
    try:
        return json.loads(raw)
    except Exception:
        pass
    for end_token in ["},\n", "}, \n", "},"]:
        last = raw.rfind(end_token)
        if last > 0:
            try:
                return json.loads(raw[:last + 1] + "]")
            except Exception:
                pass
    s, e = raw.find("["), raw.rfind("]")
    if s != -1 and e > s:
        try:
            return json.loads(raw[s:e + 1])
        except Exception:
            pass
    return None

def _pdf_gemini_worker(contents, result_holder):
    try:
        resp  = client.models.generate_content(
            model=PDF_MODEL,
            contents=contents,
            config=gtypes.GenerateContentConfig(max_output_tokens=65536, temperature=0.1),
        )
        meta  = getattr(resp, "usage_metadata", None)
        in_t  = getattr(meta, "prompt_token_count",     0) or 0
        out_t = getattr(meta, "candidates_token_count", 0) or 0
        result_holder["result"] = (resp.text.strip(), in_t, out_t)
    except Exception as e:
        result_holder["error"] = e

def call_pdf_gemini_live(contents, timer_ph, t0, page_start, label, log_lines, log_ph):
    for attempt in range(MAX_RETRIES):
        result_holder = {}
        t = threading.Thread(target=_pdf_gemini_worker, args=(contents, result_holder), daemon=True)
        t.start()
        dot = 0
        while t.is_alive():
            elapsed      = time.time() - t0
            page_elapsed = time.time() - page_start
            dots         = "." * (dot % 4)
            timer_ph.markdown(
                f"""<div class='timer-box'>
                    <div class='timer-val'>⏱ {elapsed:.0f}s</div>
                    <div class='timer-lbl'>Tổng thời gian đã chạy</div>
                    <div class='timer-status'>🔄 {label} — chờ Gemini API{dots} ({page_elapsed:.0f}s)</div>
                </div>""",
                unsafe_allow_html=True,
            )
            dot += 1
            time.sleep(1)

        if "result" in result_holder:
            return result_holder["result"]

        err     = result_holder.get("error", Exception("Unknown error"))
        err_str = str(err).lower()
        is_rate = any(c in err_str for c in RETRY_CODES)

        if is_rate and attempt < MAX_RETRIES - 1:
            wait = (2 ** attempt) * 5
            ts   = datetime.now().strftime("%H:%M:%S")
            log_lines.append(f"[{ts}] ⚠️  Rate limit! Chờ {wait}s rồi thử lại ({attempt+1}/{MAX_RETRIES})...")
            log_ph.markdown(
                f"<div class='log-box'>{'<br>'.join(log_lines[-40:])}</div>",
                unsafe_allow_html=True,
            )
            for remaining in range(wait, 0, -1):
                elapsed = time.time() - t0
                timer_ph.markdown(
                    f"""<div class='timer-box' style='border-color:#f59e0b'>
                        <div class='timer-val' style='color:#f59e0b'>⏱ {elapsed:.0f}s</div>
                        <div class='timer-lbl'>Tổng thời gian đã chạy</div>
                        <div class='timer-status' style='color:#f59e0b'>⚠️ Rate limit — thử lại sau {remaining}s</div>
                    </div>""",
                    unsafe_allow_html=True,
                )
                time.sleep(1)
        else:
            raise err

def translate_pdf_page(groups, target_lang, page_idx, timer_ph, t0, page_start, log_lines, log_ph):
    numbered = "\n".join(f"[{i}] {g['text']}" for i, g in enumerate(groups))
    prompt = (
        f"Dịch sang {target_lang}. Giữ nguyên số thứ tự [0]...[{len(groups)-1}].\n"
        f"Trả về JSON object, ĐẦY ĐỦ từ \"0\" đến \"{len(groups)-1}\", không bỏ sót:\n"
        f"{{\"0\": \"bản dịch\", \"1\": \"bản dịch\", ...}}\n"
        f"Chỉ JSON, không giải thích.\n\n"
        f"{numbered}"
    )
    label = f"Đang dịch trang {page_idx + 1}"
    raw, in_t, out_t = call_pdf_gemini_live(
        prompt, timer_ph, t0, page_start, label, log_lines, log_ph
    )
    parsed = _parse_json_pdf(raw)
    if isinstance(parsed, dict):
        return [str(parsed.get(str(i), groups[i]["text"])) for i in range(len(groups))], in_t, out_t
    if isinstance(parsed, list) and len(parsed) == len(groups):
        return [str(x) for x in parsed], in_t, out_t
    return [g["text"] for g in groups], in_t, out_t

def parse_page_range(s, total):
    pages = []
    for part in s.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            pages.extend(range(int(a) - 1, int(b)))
        elif part.isdigit():
            pages.append(int(part) - 1)
    return sorted(p for p in set(pages) if 0 <= p < total)


# ══════════════════════════════════════════════════════════════════════════════
# WORD BACKEND
# ══════════════════════════════════════════════════════════════════════════════
def _iter_paragraphs_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            for nested in cell.tables:
                yield from _iter_paragraphs_in_table(nested)

def iter_all_paragraphs(doc):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        yield from _iter_paragraphs_in_table(table)
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                yield para
            for tbl in section.header.tables:
                yield from _iter_paragraphs_in_table(tbl)
        except Exception:
            pass
        try:
            for para in section.footer.paragraphs:
                yield para
            for tbl in section.footer.tables:
                yield from _iter_paragraphs_in_table(tbl)
        except Exception:
            pass

def replace_paragraph_text_keep_format(paragraph, new_text: str):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    first_run = next((r for r in runs if r.text), None)
    if first_run is None:
        runs[0].text = new_text
        return
    first_run.text = new_text
    after = False
    for run in runs:
        if run is first_run:
            after = True
            continue
        if after and run.text:
            run.text = ""

def _detect_role(para, in_table: bool, in_header: bool, in_footer: bool) -> str:
    if in_header: return "header"
    if in_footer: return "footer"
    style = (para.style.name or "").lower()
    if "toc" in style: return "toc"
    if "heading" in style: return "section_heading"
    if in_table: return "table_cell"
    try:
        pPr = para._p.pPr
        if pPr is not None and pPr.numPr is not None:
            return "bullet"
    except Exception:
        pass
    if style in ("title", "subtitle"): return "title"
    if any(k in style for k in ("caption", "note", "warning", "caution")): return "note"
    return "paragraph"

def extract_docx_blocks(docx_bytes: bytes) -> list:
    doc = Document(io.BytesIO(docx_bytes))
    table_para_set: set  = set()
    header_para_set: set = set()
    footer_para_set: set = set()

    for tbl in doc.tables:
        for para in _iter_paragraphs_in_table(tbl):
            table_para_set.add(id(para._element))
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                header_para_set.add(id(para._element))
            for tbl in section.header.tables:
                for para in _iter_paragraphs_in_table(tbl):
                    header_para_set.add(id(para._element))
        except Exception:
            pass
        try:
            for para in section.footer.paragraphs:
                footer_para_set.add(id(para._element))
            for tbl in section.footer.tables:
                for para in _iter_paragraphs_in_table(tbl):
                    footer_para_set.add(id(para._element))
        except Exception:
            pass

    blocks = []
    for idx, para in enumerate(iter_all_paragraphs(doc)):
        text = para.text.strip()
        if not text:
            continue
        pid  = id(para._element)
        role = _detect_role(para, pid in table_para_set, pid in header_para_set, pid in footer_para_set)
        blocks.append({"id": f"p{idx}", "text": text, "role": role, "para_idx": idx})
    return blocks

def build_doc_context(blocks: list) -> str:
    titles   = [b["text"] for b in blocks if b["role"] == "title"][:2]
    headings = [b["text"] for b in blocks if b["role"] == "section_heading"][:10]
    tocs     = [b["text"] for b in blocks if b["role"] == "toc"][:8]
    lines = []
    if titles:   lines.append("Document title: " + " | ".join(titles))
    if headings: lines.append("Main sections: " + " | ".join(headings[:6]))
    if tocs:     lines.append("TOC: " + " | ".join(tocs[:5]))
    ctx      = "\n".join(lines) if lines else "Technical document."
    all_text = " ".join(b["text"] for b in blocks[:30]).lower()
    if any(k in all_text for k in ["elevator", "lift", "hoistway", "schindler", "inventio"]):
        ctx += "\nDomain: elevator/lift engineering."
    elif any(k in all_text for k in ["safety", "standard", "regulation", "iso", "en 81"]):
        ctx += "\nDomain: safety standards."
    elif any(k in all_text for k in ["software", "api", "code", "function"]):
        ctx += "\nDomain: software/IT."
    return ctx

def _call_gemini_word(prompt: str) -> str:
    models   = ([_word_working_model[0]] if _word_working_model[0] else []) + \
               [m for m in WORD_MODELS if m != _word_working_model[0]]
    last_err = "no models"
    for model in models:
        try:
            resp = client.models.generate_content(
                model=model,
                contents=prompt,
                config=gtypes.GenerateContentConfig(temperature=0.1, max_output_tokens=MAX_WORD_TOKENS),
            )
            text = (resp.text or "").strip()
            if not text:
                last_err = "empty response"
                continue
            _word_working_model[0] = model
            return text
        except Exception as e:
            last_err = str(e)
    raise RuntimeError(f"All Word models failed: {last_err}")

def _parse_json_word(raw: str) -> list:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()
    s, e = raw.find("["), raw.rfind("]")
    if s != -1 and e > s:
        candidate = raw[s:e+1]
        try:
            data = json.loads(candidate)
            if isinstance(data, list):
                return data
        except json.JSONDecodeError:
            last = candidate.rfind("},")
            if last == -1:
                last = candidate.rfind("}")
            if last > 0:
                try:
                    data = json.loads(candidate[:last+1].rstrip(",") + "]")
                    if isinstance(data, list):
                        return data
                except Exception:
                    pass
    return []

def translate_word_chunk(chunk: list, target_lang: str, doc_context: str) -> dict:
    payload = json.dumps(
        [{"id": b["id"], "text": b["text"], "role": b.get("role", "paragraph")} for b in chunk],
        ensure_ascii=False,
    )
    prompt = f"""Translate these Word document blocks into {target_lang}.

Document context:
{doc_context}

Rules:
- Return ONLY valid JSON array. No markdown, no explanation.
- Keep "id" exactly as given.
- Translate ALL words into {target_lang}. Leave nothing in the source language.
- Preserve: numbers, units, product codes, model names, document IDs, revision codes.
- Do NOT translate company names (e.g. INVENTIO AG, Schindler, Otis).
- Preserve ALL punctuation, whitespace, tabs (\\t), line breaks (\\n).
- Match source sentence count — do NOT merge or split sentences.
- For toc: translate heading text only; preserve numbering, dot leaders (...), page numbers.
- For bullet: concise imperative, keep bullet symbol.
- For heading/section_heading: concise technical, match source brevity.
- For table_cell: concise, consistent across row.
- For note: keep WARNING/NOTE/CAUTION label verbatim.

Format:
[{{"id":"...","text":"translated text"}}]

Blocks:
{payload}"""
    try:
        raw    = _call_gemini_word(prompt)
        parsed = _parse_json_word(raw)
        result = {}
        for item in parsed:
            if isinstance(item, dict) and item.get("id") and item.get("text") is not None:
                result[item["id"]] = str(item["text"])
        for b in chunk:
            if b["id"] not in result:
                result[b["id"]] = b["text"]
        return result
    except Exception:
        return {b["id"]: b["text"] for b in chunk}

def apply_word_translations(original_bytes: bytes, blocks: list, translations: dict) -> bytes:
    doc         = Document(io.BytesIO(original_bytes))
    idx_to_para = {idx: para for idx, para in enumerate(iter_all_paragraphs(doc))}
    for block in blocks:
        if block.get("role") in NO_TRANSLATE_ROLES:
            continue
        tr   = translations.get(block["id"])
        para = idx_to_para.get(block.get("para_idx"))
        if not tr or para is None:
            continue
        try:
            replace_paragraph_text_keep_format(para, tr)
        except Exception:
            pass
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# UI HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def make_log_adder(log_lines: list, log_ph):
    def add_log(msg):
        ts = datetime.now().strftime("%H:%M:%S")
        log_lines.append(f"[{ts}] {msg}")
        log_ph.markdown(
            f"<div class='log-box'>{'<br>'.join(log_lines[-40:])}</div>",
            unsafe_allow_html=True,
        )
    return add_log

def timer_box_html(elapsed, label, status, color="#6c63ff", val_color="#a78bfa", status_color="#818cf8"):
    return f"""<div class='timer-box' style='border-color:{color}'>
        <div class='timer-val' style='color:{val_color}'>⏱ {elapsed:.0f}s</div>
        <div class='timer-lbl'>Tổng thời gian đã chạy</div>
        <div class='timer-status' style='color:{status_color}'>{status}</div>
    </div>"""


# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════
tab_pdf, tab_word = st.tabs(["📄 Dịch PDF", "📝 Dịch Word"])


# ─── PDF TAB ──────────────────────────────────────────────────────────────────
with tab_pdf:
    uploaded_pdf = st.file_uploader("📄 Chọn file PDF cần dịch", type=["pdf"], key="pdf_uploader")
    col1, col2 = st.columns(2)
    with col1:
        lang_pdf = st.selectbox("🌐 Ngôn ngữ đích", LANGUAGES, key="pdf_lang")
    with col2:
        pages_s = st.text_input("📑 Trang cụ thể (tuỳ chọn)",
                                placeholder="Vd: 1-5,8  •  Để trống = tất cả", key="pdf_pages")
    st.divider()
    run_pdf = st.button("▶  Bắt đầu dịch PDF", disabled=(uploaded_pdf is None), key="pdf_run")

    if run_pdf and uploaded_pdf:
        for k in ("pdf_bytes", "pdf_out_name", "pdf_summary"):
            st.session_state.pop(k, None)

        st.markdown("### 📊 Tiến độ")
        timer_ph = st.empty()
        col_pg, col_ln, col_usd, col_vnd = st.columns(4)
        ph_pg  = col_pg.empty()
        ph_ln  = col_ln.empty()
        ph_usd = col_usd.empty()
        ph_vnd = col_vnd.empty()

        def render_pdf_stats(pages_done, total_pg, total_lines, tok_in, tok_out):
            usd = ((tok_in / 1e6) * PRICE_INPUT + (tok_out / 1e6) * PRICE_OUTPUT) * 10
            vnd = usd * USD_TO_VND
            ph_pg.markdown(f"<div class='stat-box'><div class='stat-val'>{pages_done}/{total_pg}</div><div class='stat-lbl'>Trang</div></div>", unsafe_allow_html=True)
            ph_ln.markdown(f"<div class='stat-box'><div class='stat-val'>{total_lines:,}</div><div class='stat-lbl'>Dòng text</div></div>", unsafe_allow_html=True)
            ph_usd.markdown(f"<div class='stat-box'><div class='stat-val'>${usd:.4f}</div><div class='stat-lbl'>USD</div></div>", unsafe_allow_html=True)
            ph_vnd.markdown(f"<div class='stat-box'><div class='stat-val'>{vnd:,.0f}₫</div><div class='stat-lbl'>VND</div></div>", unsafe_allow_html=True)

        render_pdf_stats(0, 0, 0, 0, 0)
        prog_pdf  = st.progress(0, text="Đang chuẩn bị...")
        st.markdown("### 📋 Nhật ký hoạt động")
        log_ph_pdf    = st.empty()
        log_lines_pdf = []
        add_log = make_log_adder(log_lines_pdf, log_ph_pdf)

        src_path = dst_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_in:
                tmp_in.write(uploaded_pdf.read())
                src_path = tmp_in.name
            tmp_out  = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            dst_path = tmp_out.name
            tmp_out.close()

            add_log(f"📄 Đã nhận file: {uploaded_pdf.name}")
            probe    = fitz.open(src_path)
            total_pg = len(probe)
            probe.close()

            targets = (parse_page_range(pages_s, total_pg)
                       if pages_s.strip() else list(range(total_pg)))
            add_log(f"📄 {total_pg} trang tổng, sẽ dịch {len(targets)} trang")

            prog_pdf.progress(5, text="Trích xuất text...")
            add_log("🔍 Trích xuất text từ PDF...")
            all_groups, _ = extract_line_groups(src_path, targets)
            total_lines   = sum(len(v) for v in all_groups.values())
            add_log(f"✅ {total_lines} dòng text")
            render_pdf_stats(0, total_pg, total_lines, 0, 0)
            add_log(f"🤖 Kết nối {PDF_MODEL}")

            all_trans = {}
            tok_in = tok_out = 0
            t0 = time.time()

            for idx, pi in enumerate(targets):
                groups     = all_groups.get(pi, [])
                page_start = time.time()
                pct        = int(10 + (idx / len(targets)) * 80)
                prog_pdf.progress(pct, text=f"Dịch trang {pi + 1}/{total_pg}...")
                add_log(f"📄 Trang {pi + 1}/{total_pg}: {len(groups)} dòng — gửi API...")

                if not groups:
                    all_trans[pi] = []
                    timer_ph.markdown(
                        timer_box_html(time.time()-t0, "", f"⏭ Trang {pi+1} trống, bỏ qua"),
                        unsafe_allow_html=True,
                    )
                    continue

                try:
                    trans, in_t, out_t = translate_pdf_page(
                        groups, lang_pdf, pi, timer_ph, t0, page_start, log_lines_pdf, log_ph_pdf
                    )
                    tok_in  += in_t
                    tok_out += out_t
                    add_log(f"   ✅ {len(trans)} dòng ({in_t:,}in/{out_t:,}out tok) — {time.time()-page_start:.1f}s")
                except Exception as e:
                    add_log(f"   ❌ Lỗi trang {pi + 1}: {e}")
                    trans = [g["text"] for g in groups]

                all_trans[pi] = trans
                render_pdf_stats(idx + 1, total_pg, total_lines, tok_in, tok_out)
                time.sleep(PDF_DELAY)

            prog_pdf.progress(92, text="Tạo PDF...")
            add_log("💾 Đang tạo PDF...")
            timer_ph.markdown(
                timer_box_html(time.time()-t0, "", "💾 Đang ghi file PDF..."),
                unsafe_allow_html=True,
            )
            font_path = find_font()
            add_log(f"🔤 Font: {os.path.basename(font_path) if font_path else 'built-in'}")
            write_translated_pdf(src_path, dst_path, all_groups, all_trans, font_path)

            elapsed = time.time() - t0
            usd     = ((tok_in / 1e6) * PRICE_INPUT + (tok_out / 1e6) * PRICE_OUTPUT) * 10
            vnd     = usd * USD_TO_VND
            render_pdf_stats(len(targets), total_pg, total_lines, tok_in, tok_out)
            prog_pdf.progress(100, text="✅ Hoàn thành!")
            timer_ph.markdown(
                f"""<div class='timer-box' style='border-color:#059669'>
                    <div class='timer-val' style='color:#10b981'>✅ {elapsed:.1f}s</div>
                    <div class='timer-lbl'>Tổng thời gian</div>
                    <div class='timer-status' style='color:#10b981'>Dịch xong {len(targets)} trang!</div>
                </div>""",
                unsafe_allow_html=True,
            )
            add_log("─" * 44)
            add_log(f"🎉 Xong {len(targets)} trang trong {elapsed:.1f}s")
            add_log(f"💰 Token: {tok_in:,} in + {tok_out:,} out")
            add_log(f"💵 Chi phí: ${usd:.4f} USD ≈ {vnd:,.0f} VND")

            with open(dst_path, "rb") as f:
                st.session_state["pdf_bytes"] = f.read()
            st.session_state["pdf_out_name"] = uploaded_pdf.name.replace(".pdf", f"_translated_{lang_pdf[:2]}.pdf")
            st.session_state["pdf_summary"]  = f"✅ Dịch xong {len(targets)} trang trong {elapsed:.1f}s  |  ${usd:.4f} USD ≈ {vnd:,.0f} VND"

        except Exception as e:
            add_log(f"❌ Lỗi: {e}")
            st.error(f"❌ Có lỗi xảy ra: {e}")
            timer_ph.markdown(
                f"""<div class='timer-box' style='border-color:#ef4444'>
                    <div class='timer-val' style='color:#ef4444'>❌ Lỗi</div>
                    <div class='timer-lbl'>Đã dừng</div>
                    <div class='timer-status' style='color:#ef4444'>{str(e)[:80]}</div>
                </div>""",
                unsafe_allow_html=True,
            )
        finally:
            for p in (src_path, dst_path):
                try:
                    if p:
                        os.unlink(p)
                except Exception:
                    pass

    if "pdf_bytes" in st.session_state:
        st.divider()
        st.success(st.session_state["pdf_summary"])
        st.download_button(
            label="⬇️  Tải PDF đã dịch",
            data=st.session_state["pdf_bytes"],
            file_name=st.session_state["pdf_out_name"],
            mime="application/pdf",
            use_container_width=True,
        )
    elif not uploaded_pdf:
        st.info("👆 Vui lòng upload file PDF để bắt đầu")


# ─── WORD TAB ─────────────────────────────────────────────────────────────────
with tab_word:
    uploaded_docx = st.file_uploader("📝 Chọn file Word cần dịch", type=["docx"], key="word_uploader")
    lang_word = st.selectbox("🌐 Ngôn ngữ đích", LANGUAGES, key="word_lang")
    st.divider()
    run_word = st.button("▶  Bắt đầu dịch Word", disabled=(uploaded_docx is None), key="word_run")

    if run_word and uploaded_docx:
        for k in ("word_bytes", "word_out_name", "word_summary"):
            st.session_state.pop(k, None)

        st.markdown("### 📊 Tiến độ")
        timer_ph_w = st.empty()
        col_bl, col_ck, col_pct = st.columns(3)
        ph_blocks = col_bl.empty()
        ph_chunks = col_ck.empty()
        ph_pct    = col_pct.empty()

        def render_word_stats(total_bl, num_chunks, done_chunks):
            pct = int(done_chunks / num_chunks * 100) if num_chunks > 0 else 0
            ph_blocks.markdown(f"<div class='stat-box'><div class='stat-val'>{total_bl}</div><div class='stat-lbl'>Đoạn văn</div></div>", unsafe_allow_html=True)
            ph_chunks.markdown(f"<div class='stat-box'><div class='stat-val'>{done_chunks}/{num_chunks}</div><div class='stat-lbl'>Chunk</div></div>", unsafe_allow_html=True)
            ph_pct.markdown(f"<div class='stat-box'><div class='stat-val'>{pct}%</div><div class='stat-lbl'>Hoàn thành</div></div>", unsafe_allow_html=True)

        render_word_stats(0, 0, 0)
        prog_word = st.progress(0, text="Đang chuẩn bị...")
        st.markdown("### 📋 Nhật ký hoạt động")
        log_ph_w    = st.empty()
        log_lines_w = []
        add_log_w   = make_log_adder(log_lines_w, log_ph_w)

        try:
            docx_bytes = uploaded_docx.read()
            add_log_w(f"📄 Đã nhận file: {uploaded_docx.name}")

            add_log_w("🔍 Phân tích cấu trúc tài liệu...")
            blocks       = extract_docx_blocks(docx_bytes)
            translatable = [b for b in blocks if b["role"] not in NO_TRANSLATE_ROLES]
            hf_count     = len(blocks) - len(translatable)
            toc_count    = sum(1 for b in blocks if b["role"] == "toc")
            add_log_w(f"✅ {len(blocks)} đoạn: {len(translatable)} cần dịch, {toc_count} TOC, {hf_count} header/footer")

            doc_context = build_doc_context(blocks)
            target_lang = LANG_EN[lang_word]
            chunks      = [translatable[i:i+CHUNK_SIZE] for i in range(0, len(translatable), CHUNK_SIZE)]
            num_chunks  = len(chunks)
            render_word_stats(len(translatable), num_chunks, 0)
            add_log_w(f"🤖 Bắt đầu dịch — {len(translatable)} đoạn / {num_chunks} chunk")
            add_log_w(f"📡 Model ưu tiên: {WORD_MODELS[0]} (fallback: {', '.join(WORD_MODELS[1:])})")

            result_holder = {
                "translations":  {},
                "done":          False,
                "error":         None,
                "chunk_done":    0,
                "current_chunk": 0,
            }

            def word_worker_fn(chunks, target_lang, doc_context, result_holder):
                try:
                    for i, chunk in enumerate(chunks):
                        result_holder["current_chunk"] = i + 1
                        r = translate_word_chunk(chunk, target_lang, doc_context)
                        result_holder["translations"].update(r)
                        result_holder["chunk_done"] = i + 1
                    result_holder["done"] = True
                except Exception as e:
                    result_holder["error"] = str(e)
                    result_holder["done"]  = True

            t0_w = time.time()
            threading.Thread(
                target=word_worker_fn,
                args=(chunks, target_lang, doc_context, result_holder),
                daemon=True,
            ).start()

            dot = 0
            while not result_holder["done"]:
                elapsed   = time.time() - t0_w
                done_c    = result_holder["chunk_done"]
                current_c = result_holder["current_chunk"]
                pct       = int(done_c / num_chunks * 90) if num_chunks > 0 else 0
                dots      = "." * (dot % 4)
                timer_ph_w.markdown(
                    timer_box_html(elapsed, "", f"🔄 Đang dịch chunk {current_c}/{num_chunks}{dots}"),
                    unsafe_allow_html=True,
                )
                prog_word.progress(pct, text=f"Chunk {done_c}/{num_chunks}...")
                render_word_stats(len(translatable), num_chunks, done_c)
                dot += 1
                time.sleep(1)

            if result_holder["error"]:
                raise RuntimeError(result_holder["error"])

            translations = result_holder["translations"]
            elapsed_w    = time.time() - t0_w

            add_log_w(f"✅ Dịch xong {len(translatable)} đoạn trong {elapsed_w:.1f}s")
            add_log_w(f"🤖 Model đã dùng: {_word_working_model[0] or WORD_MODELS[0]}")

            prog_word.progress(92, text="Tạo file DOCX...")
            add_log_w("💾 Đang tạo file DOCX...")
            timer_ph_w.markdown(
                timer_box_html(elapsed_w, "", "💾 Đang ghi file DOCX..."),
                unsafe_allow_html=True,
            )

            translated_bytes = apply_word_translations(docx_bytes, blocks, translations)
            render_word_stats(len(translatable), num_chunks, num_chunks)
            prog_word.progress(100, text="✅ Hoàn thành!")
            timer_ph_w.markdown(
                f"""<div class='timer-box' style='border-color:#059669'>
                    <div class='timer-val' style='color:#10b981'>✅ {elapsed_w:.1f}s</div>
                    <div class='timer-lbl'>Tổng thời gian</div>
                    <div class='timer-status' style='color:#10b981'>Dịch xong {len(translatable)} đoạn!</div>
                </div>""",
                unsafe_allow_html=True,
            )
            add_log_w("─" * 44)
            add_log_w(f"🎉 Xong {len(translatable)} đoạn trong {elapsed_w:.1f}s")

            st.session_state["word_bytes"]    = translated_bytes
            st.session_state["word_out_name"] = uploaded_docx.name.replace(".docx", f"_translated_{lang_word[:2]}.docx")
            st.session_state["word_summary"]  = f"✅ Dịch xong {len(translatable)} đoạn trong {elapsed_w:.1f}s"

        except Exception as e:
            add_log_w(f"❌ Lỗi: {e}")
            st.error(f"❌ Có lỗi xảy ra: {e}")
            timer_ph_w.markdown(
                f"""<div class='timer-box' style='border-color:#ef4444'>
                    <div class='timer-val' style='color:#ef4444'>❌ Lỗi</div>
                    <div class='timer-lbl'>Đã dừng</div>
                    <div class='timer-status' style='color:#ef4444'>{str(e)[:80]}</div>
                </div>""",
                unsafe_allow_html=True,
            )

    if "word_bytes" in st.session_state:
        st.divider()
        st.success(st.session_state["word_summary"])
        st.download_button(
            label="⬇️  Tải Word đã dịch",
            data=st.session_state["word_bytes"],
            file_name=st.session_state["word_out_name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    elif not uploaded_docx:
        st.info("👆 Vui lòng upload file Word (.docx) để bắt đầu")

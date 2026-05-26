"""
So sánh bản dịch DOCX với bản gốc → đánh giá chất lượng + đề xuất.

Flow:
  extract original + translated → align (positional + smart resync) → chunk pairs →
  LLM review (parallel) → list issues with severity/suggestion/reason.
"""
import io
import json
import re
from concurrent.futures import ThreadPoolExecutor, as_completed

from google.genai import types as gtypes

from word_backend import get_working_model
from config import NO_TRANSLATE_ROLES, WORD_MODELS

REVIEW_CHUNK_SIZE = 12       # pairs per LLM call
MAX_REVIEW_WORKERS = 6

_NUM_PATTERN = re.compile(r"\b\d[\d,\.]*\b")
_URL_PATTERN = re.compile(r"https?://\S+|www\.\S+")
_CODE_PATTERN = re.compile(r"\b[A-Z]{2,}[A-Z0-9_-]*\b")


def _pair_similarity(orig: dict, trans: dict) -> float:
    """
    Heuristic similarity score in [0, 1] for cross-language pair matching.
    Combines: length ratio, number overlap, URL overlap, code/acronym overlap, role match.
    """
    o_text, t_text = orig["text"], trans["text"]
    if not o_text or not t_text:
        return 0.0

    # 1. Length ratio (translation usually 0.7-1.5x source length)
    o_len, t_len = max(len(o_text), 1), max(len(t_text), 1)
    ratio = min(o_len, t_len) / max(o_len, t_len)
    len_score = ratio if 0.3 <= ratio <= 1.0 else 0.1

    # 2. Number overlap (numbers should be preserved across languages)
    o_nums = set(_NUM_PATTERN.findall(o_text))
    t_nums = set(_NUM_PATTERN.findall(t_text))
    if o_nums or t_nums:
        num_score = len(o_nums & t_nums) / max(len(o_nums | t_nums), 1)
    else:
        num_score = 0.5  # neutral

    # 3. URL / code overlap (technical identifiers preserved)
    o_codes = set(_URL_PATTERN.findall(o_text)) | set(_CODE_PATTERN.findall(o_text))
    t_codes = set(_URL_PATTERN.findall(t_text)) | set(_CODE_PATTERN.findall(t_text))
    if o_codes or t_codes:
        code_score = len(o_codes & t_codes) / max(len(o_codes | t_codes), 1)
    else:
        code_score = 0.5

    # 4. Role match bonus
    role_bonus = 0.1 if orig.get("role") == trans.get("role") else 0.0

    return min(1.0, 0.4 * len_score + 0.3 * num_score + 0.2 * code_score + role_bonus)


def align_blocks(orig_blocks: list[dict],
                 trans_blocks: list[dict],
                 smart: bool = True,
                 ) -> tuple[list[tuple], list[str]]:
    """
    Align original ↔ translated blocks.
    - smart=False: pure positional alignment (legacy behavior)
    - smart=True: positional + drift resync (sliding window when content mismatches)
    Returns (pairs, warnings).
    """
    def keep(b):
        return (b["text"].strip()
                and b["role"] not in NO_TRANSLATE_ROLES
                and b["role"] != "image_alt")

    o = [b for b in orig_blocks if keep(b)]
    t = [b for b in trans_blocks if keep(b)]
    warnings: list[str] = []
    if len(o) != len(t):
        warnings.append(
            f"⚠️ Số đoạn không khớp: gốc={len(o)} vs dịch={len(t)}."
        )

    if not smart:
        n = min(len(o), len(t))
        pairs = [(o[i], t[i]) for i in range(n)]
        return pairs, warnings

    # Smart alignment: greedy with local resync window
    pairs: list[tuple] = []
    skipped_orig: list[int] = []
    skipped_trans: list[int] = []
    i = j = 0
    SIM_THRESHOLD = 0.35
    LOOKAHEAD = 3  # try shifting up to ±3 positions to resync

    while i < len(o) and j < len(t):
        sim = _pair_similarity(o[i], t[j])
        if sim >= SIM_THRESHOLD:
            pairs.append((o[i], t[j]))
            i += 1
            j += 1
            continue
        # Drift: probe ahead in both streams to find best match
        best = (sim, 0, 0)
        for di in range(LOOKAHEAD + 1):
            for dj in range(LOOKAHEAD + 1):
                if di == 0 and dj == 0:
                    continue
                if i + di >= len(o) or j + dj >= len(t):
                    continue
                s = _pair_similarity(o[i + di], t[j + dj])
                if s > best[0]:
                    best = (s, di, dj)
        _, di, dj = best
        if di > 0:
            skipped_orig.extend(range(i, i + di))
        if dj > 0:
            skipped_trans.extend(range(j, j + dj))
        i += di
        j += dj
        if i < len(o) and j < len(t):
            pairs.append((o[i], t[j]))
            i += 1
            j += 1

    if skipped_orig:
        warnings.append(f"⚠️ Smart align: skip {len(skipped_orig)} đoạn gốc lệch nhịp")
    if skipped_trans:
        warnings.append(f"⚠️ Smart align: skip {len(skipped_trans)} đoạn dịch lệch nhịp")
    warnings.append(f"🔗 Smart align: matched {len(pairs)} pair")
    return pairs, warnings


def _build_review_prompt(pairs_chunk: list[tuple], target_lang: str,
                         doc_context: str, strictness: str) -> str:
    strictness_desc = {
        "strict":  "Strict — flag every imperfection, even minor style issues.",
        "balanced": "Balanced — flag accuracy, completeness, and notable style issues. Ignore trivial wording preferences.",
        "permissive": "Permissive — only flag clear errors: wrong meaning, missing content, severe style/register problems.",
    }.get(strictness, "Balanced")

    pairs_text = []
    for i, (ob, tb) in enumerate(pairs_chunk):
        pairs_text.append(
            f"[{i}] id={ob['id']} role={ob['role']}\n"
            f"  ORIGINAL: {ob['text']}\n"
            f"  TRANSLATED: {tb['text']}"
        )

    return (
        f"You are a professional translation reviewer. The target language is: {target_lang}.\n"
        f"Document context: {doc_context}\n\n"
        f"Strictness: {strictness_desc}\n\n"
        f"For EACH pair below, evaluate:\n"
        f"1. ACCURACY — does the translation convey the original meaning faithfully?\n"
        f"2. COMPLETENESS — is any content missing or hallucinated?\n"
        f"3. STYLE — tone, register, fluency, naturalness in {target_lang}?\n"
        f"4. TERMINOLOGY — consistency and correctness of domain terms?\n\n"
        f"Return ONLY a JSON object:\n"
        f"{{\"reviews\": [\n"
        f"  {{\"id\": \"<id>\", \"severity\": \"critical|major|minor|ok\",\n"
        f"   \"issues\": [\"short label\", ...], \"suggested\": \"improved translation\",\n"
        f"   \"reason\": \"concise explanation\"}}, ...\n"
        f"]}}\n\n"
        f"Rules:\n"
        f"- severity=ok → omit \"suggested\" and \"reason\" (or leave empty)\n"
        f"- severity=critical → wrong meaning, missing key content, hallucination\n"
        f"- severity=major → significant style/terminology issues\n"
        f"- severity=minor → small improvements only\n"
        f"- Include EVERY pair in your response, even if severity=ok\n"
        f"- \"suggested\" must be in {target_lang}\n"
        f"- \"reason\" must be in Vietnamese\n\n"
        f"Pairs:\n" + "\n\n".join(pairs_text)
    )


def review_chunk(client, pairs_chunk, target_lang, doc_context, strictness,
                 chunk_idx: int) -> tuple[list[dict], int, int]:
    """Return (reviews, in_tokens, out_tokens)."""
    prompt = _build_review_prompt(pairs_chunk, target_lang, doc_context, strictness)
    model = get_working_model() or WORD_MODELS[0]
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=gtypes.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.2,
        ),
    )
    meta  = getattr(response, "usage_metadata", None)
    in_t  = getattr(meta, "prompt_token_count",     0) or 0
    out_t = getattr(meta, "candidates_token_count", 0) or 0
    try:
        data = json.loads(response.text)
        reviews = data.get("reviews", [])
        if not isinstance(reviews, list):
            reviews = []
    except Exception:
        reviews = []
    return reviews, in_t, out_t


def chunk_pairs(pairs: list[tuple], size: int = REVIEW_CHUNK_SIZE) -> list[list[tuple]]:
    return [pairs[i:i+size] for i in range(0, len(pairs), size)]


def review_parallel(holder, client, chunks, target_lang, doc_context, strictness,
                    max_workers: int = MAX_REVIEW_WORKERS):
    """Background runner. holder is shared dict updated as chunks finish."""
    try:
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {
                ex.submit(review_chunk, client, c, target_lang, doc_context,
                          strictness, i): i
                for i, c in enumerate(chunks)
            }
            for fut in as_completed(futures):
                idx = futures[fut]
                try:
                    reviews, in_t, out_t = fut.result()
                    holder["reviews"].extend(reviews)
                    holder["tok_in"]  += in_t
                    holder["tok_out"] += out_t
                    holder["chunk_done"] += 1
                    holder["chunk_log"].append({
                        "idx": idx, "size": len(chunks[idx]),
                        "in_t": in_t, "out_t": out_t, "error": None,
                    })
                except Exception as e:
                    holder["chunk_done"] += 1
                    holder["chunk_log"].append({
                        "idx": idx, "size": len(chunks[idx]),
                        "in_t": 0, "out_t": 0, "error": str(e)[:200],
                    })
    except Exception as e:
        holder["error"] = str(e)
    finally:
        holder["done"] = True


def summarize_reviews(reviews: list[dict]) -> dict:
    """Count by severity."""
    counts = {"critical": 0, "major": 0, "minor": 0, "ok": 0}
    for r in reviews:
        sev = r.get("severity", "ok")
        if sev in counts:
            counts[sev] += 1
    return counts


def build_review_report_docx(pairs_by_id: dict, reviews: list[dict]) -> bytes:
    """Generate DOCX report: only entries with severity != ok."""
    from docx import Document
    from docx.shared import RGBColor, Inches

    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)

    doc.add_heading("Báo cáo đánh giá bản dịch", level=0)

    counts = summarize_reviews(reviews)
    p = doc.add_paragraph()
    p.add_run(f"Tổng: {len(reviews)} đoạn  •  ").bold = True
    p.add_run(f"🔴 Critical: {counts['critical']}  ").font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(f"🟠 Major: {counts['major']}  ").font.color.rgb = RGBColor(0xD0, 0x70, 0x00)
    p.add_run(f"🟡 Minor: {counts['minor']}  ").font.color.rgb = RGBColor(0x80, 0x80, 0x00)
    p.add_run(f"✅ OK: {counts['ok']}").font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    sev_order = {"critical": 0, "major": 1, "minor": 2, "ok": 3}
    sorted_revs = sorted(reviews, key=lambda r: sev_order.get(r.get("severity"), 4))

    for r in sorted_revs:
        sev = r.get("severity", "ok")
        if sev == "ok":
            continue
        pair = pairs_by_id.get(r.get("id"))
        if not pair:
            continue
        ob, tb = pair

        h = doc.add_heading(f"[{sev.upper()}] {r['id']}", level=2)
        if sev == "critical":
            h.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif sev == "major":
            h.runs[0].font.color.rgb = RGBColor(0xD0, 0x70, 0x00)

        issues = r.get("issues") or []
        if issues:
            doc.add_paragraph("Vấn đề: " + ", ".join(issues)).runs[0].italic = True

        doc.add_paragraph().add_run("Gốc:").bold = True
        doc.add_paragraph(ob["text"])
        doc.add_paragraph().add_run("Bản dịch hiện tại:").bold = True
        doc.add_paragraph(tb["text"])
        if r.get("suggested"):
            doc.add_paragraph().add_run("Đề xuất:").bold = True
            doc.add_paragraph(r["suggested"])
        if r.get("reason"):
            doc.add_paragraph().add_run("Lý do:").bold = True
            doc.add_paragraph(r["reason"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# DIFF HIGHLIGHTING — word-level inline diff between two strings
# ══════════════════════════════════════════════════════════════════════════════
def word_diff_html(old: str, new: str) -> tuple[str, str]:
    """
    Compute word-level diff between two strings.
    Returns (old_html, new_html) with <span> markup:
      - removed words (in old): red strikethrough background
      - added words   (in new): green background
      - unchanged    : default text
    """
    import difflib
    o_words = old.split()
    n_words = new.split()
    matcher = difflib.SequenceMatcher(None, o_words, n_words, autojunk=False)

    old_parts: list[str] = []
    new_parts: list[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        o_chunk = " ".join(o_words[i1:i2])
        n_chunk = " ".join(n_words[j1:j2])
        if tag == "equal":
            old_parts.append(_esc(o_chunk))
            new_parts.append(_esc(n_chunk))
        elif tag == "delete":
            old_parts.append(f'<span style="background:#ffcccc;text-decoration:line-through;color:#cc0000;font-weight:600;padding:1px 3px;border-radius:3px">{_esc(o_chunk)}</span>')
        elif tag == "insert":
            new_parts.append(f'<span style="background:#b6f5b6;color:#006600;font-weight:600;padding:1px 3px;border-radius:3px">{_esc(n_chunk)}</span>')
        elif tag == "replace":
            old_parts.append(f'<span style="background:#ffcccc;text-decoration:line-through;color:#cc0000;font-weight:600;padding:1px 3px;border-radius:3px">{_esc(o_chunk)}</span>')
            new_parts.append(f'<span style="background:#b6f5b6;color:#006600;font-weight:600;padding:1px 3px;border-radius:3px">{_esc(n_chunk)}</span>')
    return " ".join(old_parts), " ".join(new_parts)


def _esc(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


# ══════════════════════════════════════════════════════════════════════════════
# QUALITY SCORE — overall + per-page/section breakdown
# ══════════════════════════════════════════════════════════════════════════════
def compute_quality_score(reviews: list[dict],
                          pairs_by_id: dict | None = None) -> dict:
    """
    Compute quality metrics:
      - overall: 0-100 score (penalty 15/critical, 5/major, 1/minor per 100 pairs)
      - severity_pct: % distribution
      - per_page: page → {score, critical, major, minor, ok, total} (only if PDF blocks)
    """
    total = len(reviews) or 1
    counts = {"critical": 0, "major": 0, "minor": 0, "ok": 0}
    for r in reviews:
        s = r.get("severity", "ok")
        if s in counts:
            counts[s] += 1

    penalty = (counts["critical"] * 15 + counts["major"] * 5 + counts["minor"] * 1)
    overall = max(0, round(100 - penalty * 100 / total))

    severity_pct = {k: round(v * 100 / total, 1) for k, v in counts.items()}

    per_page: dict[int, dict] = {}
    if pairs_by_id:
        for r in reviews:
            pair = pairs_by_id.get(r.get("id"))
            if not pair:
                continue
            ob, _tb = pair
            page = ob.get("page")
            if page is None:
                continue
            slot = per_page.setdefault(page + 1, {
                "critical": 0, "major": 0, "minor": 0, "ok": 0, "total": 0,
            })
            sev = r.get("severity", "ok")
            if sev in slot:
                slot[sev] += 1
            slot["total"] += 1

        for page, slot in per_page.items():
            t = slot["total"] or 1
            pen = slot["critical"] * 15 + slot["major"] * 5 + slot["minor"] * 1
            slot["score"] = max(0, round(100 - pen * 100 / t))

    return {
        "overall": overall,
        "total":   total,
        "counts":  counts,
        "severity_pct": severity_pct,
        "per_page": per_page,
    }


# ══════════════════════════════════════════════════════════════════════════════
# APPLY EDITS — write user-edited translations back to original DOCX/PDF
# ══════════════════════════════════════════════════════════════════════════════
def apply_edits_to_docx(translated_docx_bytes: bytes,
                        edits: dict[str, str]) -> bytes:
    """
    Re-write translated DOCX with edits.
    edits = {block_id: new_translation}.
    Reuses word_backend.apply_translations() which iterates blocks by para_idx.
    """
    from word_backend import extract_docx_blocks, apply_translations
    blocks = extract_docx_blocks(translated_docx_bytes)
    translations = {b["id"]: edits[b["id"]] for b in blocks if b["id"] in edits}
    return apply_translations(translated_docx_bytes, blocks, translations)


def apply_edits_to_pdf(translated_pdf_bytes: bytes, orig_pdf_bytes: bytes,
                       edits: dict[str, str]) -> bytes:
    """
    Re-render PDF: extract line groups from ORIGINAL PDF, apply edits keyed by
    pdf_p{page}_{idx} convention, render translated PDF.
    edits = {block_id: new_translation} where block_id format is 'pdf_p{P}_{IDX}'.
    """
    import tempfile, os as _os
    from pdf_backend import extract_line_groups, write_translated_pdf, find_font

    src_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
    dst_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
    try:
        with open(src_path, "wb") as f:
            f.write(orig_pdf_bytes)
        all_groups, _, _ = extract_line_groups(src_path)

        # Re-build all_trans: use edits where available, else keep original text
        all_trans: dict = {}
        running_idx = 0
        for pi in sorted(all_groups.keys()):
            page_trans = []
            for li, g in enumerate(all_groups[pi]):
                bid = f"pdf_p{pi}_{running_idx}"
                page_trans.append(edits.get(bid, g["text"]))
                running_idx += 1
            all_trans[pi] = page_trans

        write_translated_pdf(src_path, dst_path, all_groups, all_trans, find_font())
        with open(dst_path, "rb") as f:
            return f.read()
    finally:
        for p in (src_path, dst_path):
            try:
                _os.unlink(p)
            except Exception:
                pass

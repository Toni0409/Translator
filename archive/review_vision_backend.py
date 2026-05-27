"""
Vision-based document comparison.

Renders each page of both documents as an image, sends each (original page,
translated page) pair to Gemini Vision, and aggregates per-page reports.

Flow:
  render_doc_as_page_images(orig)   → list[bytes]   (PNG per page)
  render_doc_as_page_images(trans)  → list[bytes]
  compare_pages_parallel(...)       → fills `holder` with per-page results
  vision_quality_score(pages)       → overall score + counts
  build_vision_report_docx(pages)   → DOCX export
"""
import io
import os
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz
from google.genai import types as gtypes

from config import WORD_MODELS
from gemini import parse_json_loose
from word_backend import get_working_model

# 2.0x ≈ 144 DPI — text remains crisp without exploding token cost
VISION_DPI_SCALE = 2.0
MAX_VISION_WORKERS = 6
VISION_PAGE_RETRIES = 3
SOFFICE_TIMEOUT = 180


# ══════════════════════════════════════════════════════════════════════════════
# RENDERING — DOCX/PDF → list of PNG bytes (one per page)
# ══════════════════════════════════════════════════════════════════════════════
def _docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convert DOCX → PDF using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "in.docx")
        with open(src, "wb") as f:
            f.write(docx_bytes)
        proc = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", tmpdir, src],
            capture_output=True, timeout=SOFFICE_TIMEOUT,
        )
        if proc.returncode != 0:
            err = proc.stderr.decode(errors="ignore")[:300]
            raise RuntimeError(f"LibreOffice convert failed: {err}")
        out_pdf = os.path.join(tmpdir, "in.pdf")
        if not os.path.exists(out_pdf):
            raise RuntimeError("LibreOffice produced no PDF output")
        with open(out_pdf, "rb") as f:
            return f.read()


def render_doc_as_page_images(file_bytes: bytes, file_name: str,
                              on_progress=None) -> list[bytes]:
    """
    Render every page of the document as a PNG. DOCX goes through soffice.
    on_progress(done, total, stage) optionally called as rendering proceeds.
    """
    name_lower = (file_name or "").lower()
    if name_lower.endswith(".docx"):
        if on_progress:
            on_progress(0, 0, "converting")
        pdf_bytes = _docx_to_pdf_bytes(file_bytes)
    else:
        pdf_bytes = file_bytes

    images: list[bytes] = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        total = doc.page_count
        mat = fitz.Matrix(VISION_DPI_SCALE, VISION_DPI_SCALE)
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            images.append(pix.tobytes("png"))
            if on_progress:
                on_progress(i + 1, total, "rendering")
    return images


def parse_page_spec(spec: str, max_pages: int) -> list[int]:
    """
    Parse a 1-indexed page spec like "1-5,7,9-12" → sorted list of ints.
    Empty / "all" / "*" returns all pages.
    Out-of-range values are clipped; invalid tokens are skipped silently.
    """
    spec = (spec or "").strip().lower()
    if not spec or spec in ("all", "*", "tat ca", "tất cả"):
        return list(range(1, max_pages + 1))
    out: set[int] = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, _, b = part.partition("-")
            try:
                start = max(1, int(a))
                end = min(max_pages, int(b))
                if end < start:
                    start, end = end, start
                for v in range(start, end + 1):
                    out.add(v)
            except ValueError:
                continue
        else:
            try:
                v = int(part)
                if 1 <= v <= max_pages:
                    out.add(v)
            except ValueError:
                continue
    return sorted(out)


# ══════════════════════════════════════════════════════════════════════════════
# COMPARISON — send (orig, trans) page pair to Gemini Vision
# ══════════════════════════════════════════════════════════════════════════════
def _build_vision_compare_prompt(target_lang: str, page_num: int,
                                 focus_areas: str = "") -> str:
    focus_block = ""
    if (focus_areas or "").strip():
        focus_block = (
            f"\nUSER FOCUS — pay extra attention to:\n"
            f"  {focus_areas.strip()}\n"
        )
    return (
        f"You are comparing TWO pages from the SAME document:\n"
        f"  • Image 1 = ORIGINAL page {page_num}\n"
        f"  • Image 2 = TRANSLATION into {target_lang}, page {page_num}\n"
        f"{focus_block}\n"
        f"Compare them as a professional translation reviewer. Focus ONLY on what matters:\n"
        f"  ✅ Translation accuracy & completeness (meaning preserved? content missing/added?)\n"
        f"  ✅ Terminology consistency for domain terms\n"
        f"  ✅ Style / tone / fluency in {target_lang}\n"
        f"  ✅ Major layout differences that change meaning or readability\n\n"
        f"IGNORE the following — do NOT report them as issues:\n"
        f"  ❌ Pure numbers, dates, codes, URLs, email addresses — these are normally preserved\n"
        f"  ❌ Header / footer text (page numbers, doc titles, running headers, watermarks)\n"
        f"  ❌ Minor font / spacing / alignment differences caused by re-rendering\n"
        f"  ❌ Image content (only flag if a whole image is missing or replaced)\n\n"
        f"Return ONLY a JSON object:\n"
        f"{{\n"
        f"  \"page\": {page_num},\n"
        f"  \"severity\": \"critical|major|minor|ok\",\n"
        f"  \"summary\": \"<1-3 sentences in Vietnamese summarizing the page>\",\n"
        f"  \"issues\": [\n"
        f"    {{\n"
        f"      \"severity\": \"critical|major|minor\",\n"
        f"      \"category\": \"missing|added|mistranslation|terminology|style|layout\",\n"
        f"      \"location\": \"<where on page, e.g. 'Tiêu đề', 'Đoạn 2', 'Bảng - hàng 3'>\",\n"
        f"      \"original\": \"<short verbatim quote from original, ≤120 chars>\",\n"
        f"      \"translated\": \"<short verbatim quote from translation, ≤120 chars>\",\n"
        f"      \"description\": \"<concise explanation in Vietnamese>\",\n"
        f"      \"suggested\": \"<optional improved translation in {target_lang}, empty if none>\"\n"
        f"    }}\n"
        f"  ]\n"
        f"}}\n\n"
        f"Rules:\n"
        f"- If the page matches well: severity=\"ok\", issues=[], summary briefly confirms correctness.\n"
        f"- severity=critical → wrong meaning, missing key content, hallucination\n"
        f"- severity=major    → significant terminology / style problems\n"
        f"- severity=minor    → small improvements only\n"
        f"- Keep ≤ 10 issues per page; pick the most impactful.\n"
        f"- summary + description must be in Vietnamese.\n"
        f"- original / translated must be verbatim quotes from the images."
    )


def _compare_pair(client, orig_png: bytes, trans_png: bytes,
                  page_num: int, target_lang: str,
                  focus_areas: str = "") -> tuple[dict, int, int]:
    """Send one page pair to Gemini Vision; return (result, in_tok, out_tok)."""
    model = get_working_model() or WORD_MODELS[0]
    prompt = _build_vision_compare_prompt(target_lang, page_num, focus_areas)

    last_err: Exception | None = None
    for _ in range(VISION_PAGE_RETRIES):
        try:
            response = client.models.generate_content(
                model=model,
                contents=[
                    gtypes.Part.from_bytes(data=orig_png, mime_type="image/png"),
                    gtypes.Part.from_bytes(data=trans_png, mime_type="image/png"),
                    prompt,
                ],
                config=gtypes.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0.2,
                ),
            )
            meta = getattr(response, "usage_metadata", None)
            in_t  = getattr(meta, "prompt_token_count",     0) or 0
            out_t = getattr(meta, "candidates_token_count", 0) or 0
            data = parse_json_loose(response.text) or {}
            if not isinstance(data, dict):
                data = {}
            data.setdefault("page", page_num)
            data.setdefault("severity", "ok")
            data.setdefault("summary", "")
            issues = data.get("issues")
            if not isinstance(issues, list):
                issues = []
            data["issues"] = issues
            return data, in_t, out_t
        except Exception as e:
            last_err = e
    raise last_err if last_err else RuntimeError("Unknown vision error")


def compare_pages_parallel(holder, client, orig_images: list[bytes],
                           trans_images: list[bytes], target_lang: str,
                           page_nums: list[int] | None = None,
                           focus_areas: str = "",
                           max_workers: int = MAX_VISION_WORKERS):
    """
    Background runner. holder is a shared dict updated as page pairs finish.
    page_nums (1-indexed) limits which page pairs to send; default = every pair.
    focus_areas is appended to each per-page prompt as USER FOCUS.
    """
    try:
        n_pairs = min(len(orig_images), len(trans_images))
        if page_nums is None:
            page_nums = list(range(1, n_pairs + 1))
        else:
            page_nums = [p for p in page_nums if 1 <= p <= n_pairs]

        if len(orig_images) != len(trans_images):
            holder["warnings"].append(
                f"⚠️ Số trang khác nhau: gốc={len(orig_images)}, "
                f"dịch={len(trans_images)}. So sánh {n_pairs} cặp đầu tiên."
            )

        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {
                ex.submit(_compare_pair, client, orig_images[p - 1],
                          trans_images[p - 1], p, target_lang, focus_areas): p
                for p in page_nums
            }
            for fut in as_completed(futures):
                page = futures[fut]
                try:
                    result, in_t, out_t = fut.result()
                    holder["pages"].append(result)
                    holder["tok_in"]  += in_t
                    holder["tok_out"] += out_t
                    holder["page_done"] += 1
                    holder["page_log"].append({
                        "page": page,
                        "severity": result.get("severity", "ok"),
                        "issues": len(result.get("issues") or []),
                        "in_t": in_t, "out_t": out_t, "error": None,
                    })
                except Exception as e:
                    holder["page_done"] += 1
                    holder["page_log"].append({
                        "page": page, "severity": "error", "issues": 0,
                        "in_t": 0, "out_t": 0, "error": str(e)[:200],
                    })
    except Exception as e:
        holder["error"] = str(e)
    finally:
        holder["done"] = True


def synthesize_overall_summary(client, pages: list[dict],
                               target_lang: str) -> tuple[str, int, int]:
    """
    Build an executive summary in Vietnamese across all per-page results.
    Returns (summary_text, in_tokens, out_tokens). On error returns ("", 0, 0).
    """
    if not pages:
        return "", 0, 0

    lines: list[str] = []
    for p in sorted(pages, key=lambda x: x.get("page", 0)):
        page_num = p.get("page", "?")
        sev = p.get("severity", "ok")
        summary = (p.get("summary") or "").strip()
        issues = p.get("issues") or []
        lines.append(f"- Page {page_num} [{sev}]: {summary} ({len(issues)} issue)")
        for iss in issues[:3]:
            lines.append(
                f"    • {iss.get('severity','')}/{iss.get('category','')} "
                f"@ {iss.get('location','')}: "
                f"{(iss.get('description','') or '')[:160]}"
            )

    prompt = (
        f"You are a senior translation reviewer. The translation target language is {target_lang}.\n"
        f"Below are per-page review summaries of a translated document.\n\n"
        f"Write an EXECUTIVE SUMMARY in Vietnamese, 3-5 short paragraphs:\n"
        f"1. Đánh giá chung chất lượng bản dịch (1 đoạn)\n"
        f"2. Các lỗi/mẫu lặp lại nổi bật ở nhiều trang (1 đoạn)\n"
        f"3. Trang/vấn đề cần ưu tiên sửa nhất (1 đoạn, liệt kê số trang)\n"
        f"4. Khuyến nghị cho người dịch (1 đoạn)\n\n"
        f"Return ONLY a JSON object: {{\"summary\": \"<văn bản markdown, dùng \\n\\n giữa các đoạn>\"}}\n\n"
        f"Per-page data:\n" + "\n".join(lines)
    )

    model = get_working_model() or WORD_MODELS[0]
    try:
        response = client.models.generate_content(
            model=model,
            contents=prompt,
            config=gtypes.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.3,
            ),
        )
        meta = getattr(response, "usage_metadata", None)
        in_t  = getattr(meta, "prompt_token_count",     0) or 0
        out_t = getattr(meta, "candidates_token_count", 0) or 0
        data = parse_json_loose(response.text) or {}
        if isinstance(data, dict):
            return str(data.get("summary", "") or "").strip(), in_t, out_t
        return "", in_t, out_t
    except Exception:
        return "", 0, 0


# ══════════════════════════════════════════════════════════════════════════════
# AGGREGATE — quality score + DOCX report
# ══════════════════════════════════════════════════════════════════════════════
def aggregate_issues_by_category(pages: list[dict]) -> list[dict]:
    """
    Aggregate all issues across pages by (category, severity).
    Returns sorted list of {category, total, critical, major, minor, pages}.
    """
    from collections import defaultdict
    bucket: dict[str, dict] = defaultdict(
        lambda: {"critical": 0, "major": 0, "minor": 0, "pages": set()}
    )
    for p in pages:
        page_num = p.get("page")
        for iss in p.get("issues") or []:
            cat = (iss.get("category") or "other").strip().lower() or "other"
            sev = iss.get("severity", "minor")
            slot = bucket[cat]
            if sev in slot:
                slot[sev] += 1
            if page_num is not None:
                slot["pages"].add(page_num)

    rows = []
    for cat, slot in bucket.items():
        total = slot["critical"] + slot["major"] + slot["minor"]
        rows.append({
            "category": cat,
            "total":    total,
            "critical": slot["critical"],
            "major":    slot["major"],
            "minor":    slot["minor"],
            "pages":    sorted(slot["pages"]),
        })
    rows.sort(key=lambda r: (-r["critical"], -r["major"], -r["minor"]))
    return rows


def flatten_issues(pages: list[dict]) -> list[dict]:
    """Flatten all per-page issues into a single sorted list with `page` field."""
    sev_order = {"critical": 0, "major": 1, "minor": 2}
    out = []
    for p in pages:
        page_num = p.get("page")
        for iss in p.get("issues") or []:
            row = dict(iss)
            row["page"] = page_num
            out.append(row)
    out.sort(key=lambda r: (sev_order.get(r.get("severity"), 9), r.get("page", 0)))
    return out


def vision_quality_score(pages: list[dict]) -> dict:
    """Compute overall score + issue/page distribution."""
    counts = {"critical": 0, "major": 0, "minor": 0, "ok": 0}
    for p in pages:
        for iss in p.get("issues") or []:
            sev = iss.get("severity", "minor")
            if sev in counts:
                counts[sev] += 1

    total_pages = len(pages) or 1
    penalty = counts["critical"] * 15 + counts["major"] * 5 + counts["minor"] * 1
    overall = max(0, round(100 - penalty / total_pages))
    total_issues = counts["critical"] + counts["major"] + counts["minor"]

    page_sev = {"critical": 0, "major": 0, "minor": 0, "ok": 0, "error": 0}
    for p in pages:
        s = p.get("severity", "ok")
        if s in page_sev:
            page_sev[s] += 1

    return {
        "overall": overall,
        "total_pages": len(pages),
        "total_issues": total_issues,
        "issue_counts": counts,
        "page_severity_counts": page_sev,
    }


def build_vision_report_docx(pages: list[dict], orig_name: str,
                             trans_name: str, target_lang: str,
                             overall_summary: str = "") -> bytes:
    from docx import Document
    from docx.shared import RGBColor, Inches

    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)

    doc.add_heading("Báo cáo so sánh tài liệu — Vision mode", level=0)
    doc.add_paragraph(f"Gốc: {orig_name}")
    doc.add_paragraph(f"Dịch: {trans_name}")
    doc.add_paragraph(f"Ngôn ngữ đích: {target_lang}")

    metrics = vision_quality_score(pages)
    ic = metrics["issue_counts"]
    p = doc.add_paragraph()
    p.add_run(f"Tổng số trang: {metrics['total_pages']}  •  ").bold = True
    p.add_run(f"Quality score: {metrics['overall']}/100").bold = True

    p = doc.add_paragraph()
    p.add_run(f"🔴 Critical: {ic['critical']}  ").font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(f"🟠 Major: {ic['major']}  ").font.color.rgb = RGBColor(0xD0, 0x70, 0x00)
    p.add_run(f"🟡 Minor: {ic['minor']}").font.color.rgb = RGBColor(0x80, 0x80, 0x00)

    if overall_summary:
        doc.add_heading("Tóm tắt tổng quan", level=1)
        for para in overall_summary.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    for page in sorted(pages, key=lambda x: x.get("page", 0)):
        page_num = page.get("page", "?")
        sev = page.get("severity", "ok")
        sev_emoji = {"critical": "🔴", "major": "🟠",
                     "minor": "🟡", "ok": "✅"}.get(sev, "⚪")
        h = doc.add_heading(f"{sev_emoji} Trang {page_num} — {sev.upper()}", level=2)
        if sev == "critical":
            h.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif sev == "major":
            h.runs[0].font.color.rgb = RGBColor(0xD0, 0x70, 0x00)

        if page.get("summary"):
            p = doc.add_paragraph()
            p.add_run("Tổng quan: ").bold = True
            p.add_run(page["summary"])

        for iss in page.get("issues") or []:
            isev = iss.get("severity", "minor")
            isev_emoji = {"critical": "🔴", "major": "🟠",
                          "minor": "🟡"}.get(isev, "⚪")
            doc.add_heading(
                f"{isev_emoji} [{isev.upper()}] "
                f"{iss.get('category', '')} — {iss.get('location', '')}",
                level=3,
            )
            if iss.get("original"):
                doc.add_paragraph().add_run("Gốc: ").bold = True
                doc.add_paragraph(iss["original"])
            if iss.get("translated"):
                doc.add_paragraph().add_run("Dịch: ").bold = True
                doc.add_paragraph(iss["translated"])
            if iss.get("description"):
                doc.add_paragraph().add_run("Mô tả: ").bold = True
                doc.add_paragraph(iss["description"])
            if iss.get("suggested"):
                doc.add_paragraph().add_run("Đề xuất: ").bold = True
                doc.add_paragraph(iss["suggested"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

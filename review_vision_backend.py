"""
Vision-based document comparison.

Renders each page of both documents as an image, sends each (original page,
translated page) pair to Gemini Vision, and aggregates per-page reports.

Flow:
  render_doc_as_page_images(orig)   → list[bytes]   (PNG per page)
  render_doc_as_page_images(trans)  → list[bytes]
  compare_pages_parallel(...)       → fills `holder` with per-page results
  vision_quality_score(pages)       → overall score + counts
  build_vision_report_docx(pages)   → DOCX export
"""
import io
import os
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz
from google.genai import types as gtypes

from config import WORD_MODELS
from gemini import parse_json_loose
from word_backend import get_working_model

# 2.0x ≈ 144 DPI — text remains crisp without exploding token cost
VISION_DPI_SCALE = 2.0
MAX_VISION_WORKERS = 6
VISION_PAGE_RETRIES = 3
SOFFICE_TIMEOUT = 180


# ══════════════════════════════════════════════════════════════════════════════
# RENDERING — DOCX/PDF → list of PNG bytes (one per page)
# ══════════════════════════════════════════════════════════════════════════════
def _docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convert DOCX → PDF using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "in.docx")
        with open(src, "wb") as f:
            f.write(docx_bytes)
        proc = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", tmpdir, src],
            capture_output=True, timeout=SOFFICE_TIMEOUT,
        )
        if proc.returncode != 0:
            err = proc.stderr.decode(errors="ignore")[:300]
            raise RuntimeError(f"LibreOffice convert failed: {err}")
        out_pdf = os.path.join(tmpdir, "in.pdf")
        if not os.path.exists(out_pdf):
            raise RuntimeError("LibreOffice produced no PDF output")
        with open(out_pdf, "rb") as f:
            return f.read()


def render_doc_as_page_images(file_bytes: bytes, file_name: str) -> list[bytes]:
    """Render every page of the document as a PNG. DOCX goes through soffice."""
    name_lower = (file_name or "").lower()
    if name_lower.endswith(".docx"):
        pdf_bytes = _docx_to_pdf_bytes(file_bytes)
    else:
        pdf_bytes = file_bytes

    images: list[bytes] = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        mat = fitz.Matrix(VISION_DPI_SCALE, VISION_DPI_SCALE)
        for page in doc:
            pix = page.get_pixmap(matrix=mat, alpha=False)
            images.append(pix.tobytes("png"))
    return images


# ══════════════════════════════════════════════════════════════════════════════
# COMPARISON — send (orig, trans) page pair to Gemini Vision
# ══════════════════════════════════════════════════════════════════════════════
def _build_vision_compare_prompt(target_lang: str, page_num: int) -> str:
    return (
        f"You are comparing TWO pages from the SAME document:\n"
        f"  • Image 1 = ORIGINAL page {page_num}\n"
        f"  • Image 2 = TRANSLATION into {target_lang}, page {page_num}\n\n"
        f"Compare them as a professional translation reviewer. Focus ONLY on what matters:\n"
        f"  ✅ Translation accuracy & completeness (meaning preserved? content missing/added?)\n"
        f"  ✅ Terminology consistency for domain terms\n"
        f"  ✅ Style / tone / fluency in {target_lang}\n"
        f"  ✅ Major layout differences that change meaning or readability\n\n"
        f"IGNORE the following — do NOT report them as issues:\n"
        f"  ❌ Pure numbers, dates, codes, URLs, email addresses — these are normally preserved\n"
        f"  ❌ Header / footer text (page numbers, doc titles, running headers, watermarks)\n"
        f"  ❌ Minor font / spacing / alignment differences caused by re-rendering\n"
        f"  ❌ Image content (only flag if a whole image is missing or replaced)\n\n"
        f"Return ONLY a JSON object:\n"
        f"{{\n"
        f"  \"page\": {page_num},\n"
        f"  \"severity\": \"critical|major|minor|ok\",\n"
        f"  \"summary\": \"<1-3 sentences in Vietnamese summarizing the page>\",\n"
        f"  \"issues\": [\n"
        f"    {{\n"
        f"      \"severity\": \"critical|major|minor\",\n"
        f"      \"category\": \"missing|added|mistranslation|terminology|style|layout\",\n"
        f"      \"location\": \"<where on page, e.g. 'Tiêu đề', 'Đoạn 2', 'Bảng - hàng 3'>\",\n"
        f"      \"original\": \"<short verbatim quote from original, ≤120 chars>\",\n"
        f"      \"translated\": \"<short verbatim quote from translation, ≤120 chars>\",\n"
        f"      \"description\": \"<concise explanation in Vietnamese>\",\n"
        f"      \"suggested\": \"<optional improved translation in {target_lang}, empty if none>\"\n"
        f"    }}\n"
        f"  ]\n"
        f"}}\n\n"
        f"Rules:\n"
        f"- If the page matches well: severity=\"ok\", issues=[], summary briefly confirms correctness.\n"
        f"- severity=critical → wrong meaning, missing key content, hallucination\n"
        f"- severity=major    → significant terminology / style problems\n"
        f"- severity=minor    → small improvements only\n"
        f"- Keep ≤ 10 issues per page; pick the most impactful.\n"
        f"- summary + description must be in Vietnamese.\n"
        f"- original / translated must be verbatim quotes from the images."
    )


def _compare_pair(client, orig_png: bytes, trans_png: bytes,
                  page_num: int, target_lang: str) -> tuple[dict, int, int]:
    """Send one page pair to Gemini Vision; return (result, in_tok, out_tok)."""
    model = get_working_model() or WORD_MODELS[0]
    prompt = _build_vision_compare_prompt(target_lang, page_num)

    last_err: Exception | None = None
    for _ in range(VISION_PAGE_RETRIES):
        try:
            response = client.models.generate_content(
                model=model,
                contents=[
                    gtypes.Part.from_bytes(data=orig_png, mime_type="image/png"),
                    gtypes.Part.from_bytes(data=trans_png, mime_type="image/png"),
                    prompt,
                ],
                config=gtypes.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0.2,
                ),
            )
            meta = getattr(response, "usage_metadata", None)
            in_t  = getattr(meta, "prompt_token_count",     0) or 0
            out_t = getattr(meta, "candidates_token_count", 0) or 0
            data = parse_json_loose(response.text) or {}
            if not isinstance(data, dict):
                data = {}
            data.setdefault("page", page_num)
            data.setdefault("severity", "ok")
            data.setdefault("summary", "")
            issues = data.get("issues")
            if not isinstance(issues, list):
                issues = []
            data["issues"] = issues
            return data, in_t, out_t
        except Exception as e:
            last_err = e
    raise last_err if last_err else RuntimeError("Unknown vision error")


def compare_pages_parallel(holder, client, orig_images: list[bytes],
                           trans_images: list[bytes], target_lang: str,
                           max_workers: int = MAX_VISION_WORKERS):
    """Background runner. holder is shared dict updated as page pairs finish."""
    try:
        n_pairs = min(len(orig_images), len(trans_images))
        if len(orig_images) != len(trans_images):
            holder["warnings"].append(
                f"⚠️ Số trang khác nhau: gốc={len(orig_images)}, "
                f"dịch={len(trans_images)}. So sánh {n_pairs} cặp đầu tiên."
            )

        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {
                ex.submit(_compare_pair, client, orig_images[i],
                          trans_images[i], i + 1, target_lang): i
                for i in range(n_pairs)
            }
            for fut in as_completed(futures):
                idx = futures[fut]
                try:
                    result, in_t, out_t = fut.result()
                    holder["pages"].append(result)
                    holder["tok_in"]  += in_t
                    holder["tok_out"] += out_t
                    holder["page_done"] += 1
                    holder["page_log"].append({
                        "page": idx + 1,
                        "severity": result.get("severity", "ok"),
                        "issues": len(result.get("issues") or []),
                        "in_t": in_t, "out_t": out_t, "error": None,
                    })
                except Exception as e:
                    holder["page_done"] += 1
                    holder["page_log"].append({
                        "page": idx + 1, "severity": "error", "issues": 0,
                        "in_t": 0, "out_t": 0, "error": str(e)[:200],
                    })
    except Exception as e:
        holder["error"] = str(e)
    finally:
        holder["done"] = True


# ══════════════════════════════════════════════════════════════════════════════
# AGGREGATE — quality score + DOCX report
# ══════════════════════════════════════════════════════════════════════════════
def vision_quality_score(pages: list[dict]) -> dict:
    """Compute overall score + issue/page distribution."""
    counts = {"critical": 0, "major": 0, "minor": 0, "ok": 0}
    for p in pages:
        for iss in p.get("issues") or []:
            sev = iss.get("severity", "minor")
            if sev in counts:
                counts[sev] += 1

    total_pages = len(pages) or 1
    penalty = counts["critical"] * 15 + counts["major"] * 5 + counts["minor"] * 1
    overall = max(0, round(100 - penalty / total_pages))
    total_issues = counts["critical"] + counts["major"] + counts["minor"]

    page_sev = {"critical": 0, "major": 0, "minor": 0, "ok": 0, "error": 0}
    for p in pages:
        s = p.get("severity", "ok")
        if s in page_sev:
            page_sev[s] += 1

    return {
        "overall": overall,
        "total_pages": len(pages),
        "total_issues": total_issues,
        "issue_counts": counts,
        "page_severity_counts": page_sev,
    }


def build_vision_report_docx(pages: list[dict], orig_name: str,
                             trans_name: str, target_lang: str) -> bytes:
    from docx import Document
    from docx.shared import RGBColor, Inches

    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)

    doc.add_heading("Báo cáo so sánh tài liệu — Vision mode", level=0)
    doc.add_paragraph(f"Gốc: {orig_name}")
    doc.add_paragraph(f"Dịch: {trans_name}")
    doc.add_paragraph(f"Ngôn ngữ đích: {target_lang}")

    metrics = vision_quality_score(pages)
    ic = metrics["issue_counts"]
    p = doc.add_paragraph()
    p.add_run(f"Tổng số trang: {metrics['total_pages']}  •  ").bold = True
    p.add_run(f"Quality score: {metrics['overall']}/100").bold = True

    p = doc.add_paragraph()
    p.add_run(f"🔴 Critical: {ic['critical']}  ").font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(f"🟠 Major: {ic['major']}  ").font.color.rgb = RGBColor(0xD0, 0x70, 0x00)
    p.add_run(f"🟡 Minor: {ic['minor']}").font.color.rgb = RGBColor(0x80, 0x80, 0x00)

    for page in sorted(pages, key=lambda x: x.get("page", 0)):
        page_num = page.get("page", "?")
        sev = page.get("severity", "ok")
        sev_emoji = {"critical": "🔴", "major": "🟠",
                     "minor": "🟡", "ok": "✅"}.get(sev, "⚪")
        h = doc.add_heading(f"{sev_emoji} Trang {page_num} — {sev.upper()}", level=2)
        if sev == "critical":
            h.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif sev == "major":
            h.runs[0].font.color.rgb = RGBColor(0xD0, 0x70, 0x00)

        if page.get("summary"):
            p = doc.add_paragraph()
            p.add_run("Tổng quan: ").bold = True
            p.add_run(page["summary"])

        for iss in page.get("issues") or []:
            isev = iss.get("severity", "minor")
            isev_emoji = {"critical": "🔴", "major": "🟠",
                          "minor": "🟡"}.get(isev, "⚪")
            doc.add_heading(
                f"{isev_emoji} [{isev.upper()}] "
                f"{iss.get('category', '')} — {iss.get('location', '')}",
                level=3,
            )
            if iss.get("original"):
                doc.add_paragraph().add_run("Gốc: ").bold = True
                doc.add_paragraph(iss["original"])
            if iss.get("translated"):
                doc.add_paragraph().add_run("Dịch: ").bold = True
                doc.add_paragraph(iss["translated"])
            if iss.get("description"):
                doc.add_paragraph().add_run("Mô tả: ").bold = True
                doc.add_paragraph(iss["description"])
            if iss.get("suggested"):
                doc.add_paragraph().add_run("Đề xuất: ").bold = True
                doc.add_paragraph(iss["suggested"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
